"""Парсер данных ОПО из Word-документов."""
import io
import re
from datetime import datetime
from docx import Document


class WordImportService:
    """Извлечение полей ОПО из .docx файла."""

    FIELD_PATTERNS = {
        'f1_1': [
            r'Наименование\s+ОПО[:\s]+(.+)',
            r'Полное\s+наименование[:\s]+(.+)',
            r'^(\S.+(?:станция|компрессорная|котельная|газопровод|насосная|хранилище).+)$',
        ],
        'f1_4': [
            r'Адрес[:\s]+(.+)',
            r'Местоположение[:\s]+(.+)',
            r'Место\s+нахождения[:\s]+(.+)',
        ],
        'danger_class': [
            r'Класс\s+опасности[:\s]+([IIV]+)',
            r'Класс[:\s]+([IIV]+)',
        ],
        'f1_5': [
            r'ОКТМО[:\s]+(\d+)',
        ],
        'f1_7_1': [
            r'(?:Наименование|Собственник)[:\s]+(.+)',
        ],
        'f1_7_2': [
            r'ИНН\s+собственника[:\s]+(\d+)',
            r'ИНН[:\s]+(\d{10,12})',
        ],
        'f1_2': [
            r'Типовое\s+наименование[:\s]+(.+)',
            r'Тип\s+объекта[:\s]+(.+)',
        ],
        'f1_3': [
            r'Цифровое\s+обозначение[:\s]+(.+)',
            r'Код\s+отрасли[:\s]+(.+)',
        ],
        'f1_6': [
            r'Дата\s+ввода\s+в\s+эксплуатацию[:\s]+(\d{2}\.\d{2}\.\d{4})',
        ],
    }

    def import_from_word(self, file_content: bytes) -> dict:
        """Парсит .docx и возвращает dict с полями формы."""
        doc = Document(io.BytesIO(file_content))

        lines = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(' '.join(cells))

        full_text = '\n'.join(lines)

        extracted = {}
        for field, patterns in self.FIELD_PATTERNS.items():
            for pat in patterns:
                m = re.search(pat, full_text, re.IGNORECASE | re.MULTILINE)
                if m:
                    val = m.group(1).strip()
                    if field == 'f1_6':
                        try:
                            val = datetime.strptime(val, '%d.%m.%Y').strftime('%Y-%m-%d')
                        except ValueError:
                            continue
                    if field == 'danger_class':
                        cmap = {'1': 'I', '2': 'II', '3': 'III', '4': 'IV'}
                        val = cmap.get(val, val.upper())
                    extracted[field] = val
                    break

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    key = cells[0]
                    val = cells[1]
                    if 'Класс' in key and 'опасности' in key:
                        extracted['danger_class'] = val.upper()
                    elif 'ОКТМО' in key:
                        extracted['f1_5'] = val
                    elif 'ИНН' in key:
                        extracted['f1_7_2'] = val

        return {
            'success': bool(extracted),
            'data': extracted,
            'warnings': [] if extracted.get('f1_1') else ['Не удалось извлечь наименование ОПО'],
        }
