"""
Генератор документов (ADR-003).
Читает structure.json → рендерит Jinja2-шаблоны → собирает DOCX.
"""
from __future__ import annotations

import io
import json
from pathlib import Path
from uuid import UUID

import re

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from jinja2 import Environment, FileSystemLoader

# Значения подобраны по эталонному документу ПМЛА (styles.xml / sectPr)
BODY_FONT_NAME = "Times New Roman"
BODY_FONT_SIZE_PT = 12
HEADING_FONT_SIZE_PT = 14
PAGE_MARGINS_CM = {"top": 2.0, "bottom": 2.0, "left": 3.0, "right": 1.5}
FIRST_LINE_INDENT_CM = 1.25

# PII-защита
PII_FIELD_NAMES = {"full_name", "phone", "inn", "snils", "email", "address", "birth_date"}


class PiiRoutingError(Exception):
    """Ошибка: pii=true секция отправлена без локального LLM."""


def strip_pii(data):
    """Удаляет персональные данные из контекста (deep copy, оригинал не мутирует)."""
    if isinstance(data, dict):
        return {
            k: "[скрыто]" if k in PII_FIELD_NAMES else strip_pii(v)
            for k, v in data.items()
        }
    if isinstance(data, list):
        return [strip_pii(item) for item in data]
    if hasattr(data, "__dict__"):
        return strip_pii(vars(data))
    return data

from src.infrastructure.llm.providers import LLMProvider, LLMMessage
from src.infrastructure.rag.pipeline import Retriever
from src.core.settings import settings


TEMPLATES_DIR = Path(__file__).parent.parent.parent.parent / "templates"


class DocumentGenerator:
    """
    Генерирует документ DOCX по типу документа и ID ОПО.

    Алгоритм:
    1. Загрузить structure.json для типа документа
    2. Для каждого раздела:
       - content_type=data  → рендер Jinja2 с данными из БД
       - content_type=llm   → получить RAG-контекст + вызвать LLM → рендер Jinja2
    3. Собрать DOCX из отрендеренных секций
    """

    def __init__(self, llm: LLMProvider | None = None, retriever: Retriever = None,
                 local_llm: LLMProvider | None = None, external_llm: LLMProvider | None = None):
        self._llm = llm
        self._local_llm = local_llm
        self._external_llm = external_llm or llm
        self._retriever = retriever

    async def generate(
        self,
        document_type: str,
        context: dict,
    ) -> bytes:
        """
        context — словарь с данными:
        {
          "organization": {...},
          "facility": {...},
          "equipment": [...],
          "substances": [...],
          "responsible_persons": [...]
        }
        Возвращает DOCX как байты.
        """
        structure = self._load_structure(document_type)
        jinja_env = Environment(
            loader=FileSystemLoader(str(TEMPLATES_DIR / document_type)),
            autoescape=False,
        )

        rendered_sections: list[tuple[str, str]] = []

        for section in structure["sections"]:
            title = section["title"]
            template_name = section["template"]
            content_type = section["content_type"]

            if content_type == "data":
                # Только данные из БД — LLM не вызывается
                rendered = self._render_template(jinja_env, template_name, context)

            elif content_type == "llm":
                # RAG + LLM генерация
                rag_context = ""
                try:
                    rag_query = self._build_rag_query(section.get("rag_query", ""), context)
                    rag_chunks = await self._retriever.retrieve(rag_query)
                    rag_context = "\n\n".join(c.content for c in rag_chunks)
                except Exception:
                    pass  # RAG недоступен — продолжаем без него

                llm_content = await self._generate_section(
                    section_title=title,
                    context=context,
                    rag_context=rag_context,
                )
                render_ctx = {**context, "llm_content": llm_content, "rag_context": rag_context}
                rendered = self._render_template(jinja_env, template_name, render_ctx)

            else:
                rendered = f"[Неизвестный content_type: {content_type}]"

            rendered_sections.append((title, rendered))

        return self._build_docx(structure["title"], rendered_sections)

    def _load_structure(self, document_type: str) -> dict:
        structure_path = TEMPLATES_DIR / document_type / "structure.json"
        if not structure_path.exists():
            raise ValueError(f"Шаблон не найден: {document_type}")
        return json.loads(structure_path.read_text(encoding="utf-8"))

    def _render_template(self, env: Environment, template_name: str, context: dict) -> str:
        template = env.get_template(template_name)
        return template.render(**context)

    def _build_rag_query(self, query_template: str, context: dict) -> str:
        """Подставляет переменные в rag_query из structure.json."""
        facility = context.get("facility", {})
        substances = context.get("substances", [])
        substance_names = " ".join(
            s.get("name", "") if isinstance(s, dict) else getattr(s, "name", "")
            for s in substances
        )
        return query_template.format(
            facility_type=getattr(facility, "facility_type", "") or "",
            hazard_class=getattr(facility, "hazard_class", "") or "",
            substance_names=substance_names,
        )

    async def _generate_section(
        self,
        section_title: str,
        context: dict,
        rag_context: str,
    ) -> str:
        """Вызывает LLM для генерации текста раздела. Если LLM недоступна — заглушка."""
        if self._llm is None:
            return (
                f"[Раздел «{section_title}» — генерация ИИ недоступна]\n"
                f"Для генерации данного раздела необходимо настроить LLM-провайдер "
                f"(задайте OPENAI_API_KEY в .env)."
            )

        facility = context.get("facility")
        organization = context.get("organization")
        substances = context.get("substances", [])
        equipment = context.get("equipment", [])

        system_prompt = (
            "Ты — эксперт по промышленной безопасности с 10-летним стажем. "
            "Твоя задача — разработать раздел Плана мероприятий по локализации "
            "и ликвидации последствий аварий (ПМЛА) в строгом соответствии с "
            "Постановлением Правительства РФ №1437 от 15.09.2020 и ФЗ №116-ФЗ. "
            "Используй официально-деловой стиль. "
            "Не выдумывай факты — только на основе предоставленных данных. "
            "Если данных недостаточно — укажи [Данные не предоставлены]."
        )

        user_prompt = f"""
Раздел: {section_title}

--- ДАННЫЕ ОБ ОПО ---
Объект: {getattr(facility, 'name', '')}
Тип: {getattr(facility, 'facility_type', '')}
Класс опасности: {getattr(facility, 'hazard_class', '')}
Организация: {getattr(organization, 'name', '')}

Опасные вещества:
{self._format_substances(substances)}

Оборудование:
{self._format_equipment(equipment)}

--- ФРАГМЕНТЫ ИЗ НОРМАТИВНОЙ БАЗЫ ---
{rag_context if rag_context else "Нормативные фрагменты не найдены."}

--- ЗАДАНИЕ ---
Напиши текст раздела «{section_title}» для данного ОПО.
Текст должен быть конкретным, ссылаться на указанное оборудование и вещества.
"""

        response = await self._llm.complete(
            messages=[
                LLMMessage(role="system", content=system_prompt),
                LLMMessage(role="user", content=user_prompt),
            ]
        )
        return response.content

    def _format_substances(self, substances: list) -> str:
        lines = []
        for s in substances:
            if isinstance(s, dict):
                lines.append(f"- {s.get('name', '—')}, {s.get('quantity_kg', '—')} кг")
            else:
                lines.append(f"- {getattr(s, 'name', '—')}, {getattr(s, 'quantity_kg', '—')} кг")
        return "\n".join(lines) if lines else "Не указаны"

    def _format_equipment(self, equipment: list) -> str:
        lines = []
        for eq in equipment:
            if isinstance(eq, dict):
                lines.append(f"- {eq.get('name', '—')} ({eq.get('equipment_type', '—')})")
            else:
                lines.append(f"- {getattr(eq, 'name', '—')} ({getattr(eq, 'equipment_type', '—')})")
        return "\n".join(lines) if lines else "Не указано"

    def _setup_document_defaults(self, doc: DocxDocument) -> None:
        """Настраивает страницу и базовый стиль по образцу эталонного ПМЛА (ГОСТ)."""
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(PAGE_MARGINS_CM["top"])
        section.bottom_margin = Cm(PAGE_MARGINS_CM["bottom"])
        section.left_margin = Cm(PAGE_MARGINS_CM["left"])
        section.right_margin = Cm(PAGE_MARGINS_CM["right"])

        normal = doc.styles["Normal"]
        normal.font.name = BODY_FONT_NAME
        normal.font.size = Pt(BODY_FONT_SIZE_PT)
        # прописываем шрифт и для восточноазиатского/complex-script прогона —
        # иначе Word может подставить дефолтный Calibri для кириллицы
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), BODY_FONT_NAME)

        pf = normal.paragraph_format
        pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.0
        pf.space_after = Pt(0)

    def _add_heading(self, doc: DocxDocument, text: str, *, level: int, center: bool = True) -> None:
        """Заголовок в стиле эталона: Times New Roman, жирный, без синего цвета Word-стиля."""
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(text)
        run.font.name = BODY_FONT_NAME
        run.font.bold = True
        run.font.size = Pt(HEADING_FONT_SIZE_PT if level == 0 else HEADING_FONT_SIZE_PT - 2)
        run.font.color.rgb = RGBColor(0, 0, 0)

    _BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

    def _add_body_paragraph(self, doc: DocxDocument, line: str) -> None:
        """
        Добавляет абзац содержимого, конвертируя markdown-разметку **жирный**
        (частый артефакт LLM-ответов) в реальное жирное форматирование —
        вместо буквальных звёздочек в тексте.
        """
        paragraph = doc.add_paragraph()
        pos = 0
        for match in self._BOLD_RE.finditer(line):
            if match.start() > pos:
                paragraph.add_run(line[pos:match.start()])
            bold_run = paragraph.add_run(match.group(1))
            bold_run.font.bold = True
            pos = match.end()
        if pos < len(line):
            paragraph.add_run(line[pos:])

    def _build_docx(self, title: str, sections: list[tuple[str, str]]) -> bytes:
        """Собирает итоговый DOCX из отрендеренных секций."""
        doc = DocxDocument()
        self._setup_document_defaults(doc)

        # Заголовок документа
        self._add_heading(doc, title, level=0, center=True)

        for section_title, content in sections:
            # Заголовок раздела
            self._add_heading(doc, section_title, level=1, center=False)
            # Содержимое
            for line in content.strip().split("\n"):
                if line.strip():
                    self._add_body_paragraph(doc, line.strip())

            doc.add_paragraph()  # отступ между разделами

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()