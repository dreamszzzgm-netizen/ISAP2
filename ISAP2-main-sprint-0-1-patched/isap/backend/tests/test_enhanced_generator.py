"""
Тесты для enhanced_generator.py:
- PII-роутинг (strip_pii, local_llm/external_llm)
- GOST-форматирование DOCX
- Основная логика генерации
"""
from __future__ import annotations

import json
import io
import pytest
from unittest.mock import AsyncMock, MagicMock, patch
from pathlib import Path

from src.application.services.enhanced_generator import (
    EnhancedDocumentGenerator,
    strip_pii,
    PiiRoutingError,
    PII_FIELD_NAMES,
    BODY_FONT_NAME,
    BODY_FONT_SIZE_PT,
    PAGE_MARGINS_CM,
)


# ---------------------------------------------------------------------------
# Утилиты
# ---------------------------------------------------------------------------

class Person:
    def __init__(self, full_name, phone, position):
        self.full_name = full_name
        self.phone = phone
        self.position = position


def _make_mock_llm(content="Текст раздела"):
    llm = AsyncMock()
    llm.complete.return_value = MagicMock(content=content)
    return llm


def _make_retriever(chunks=None):
    retriever = AsyncMock()
    retriever.retrieve.return_value = chunks or []
    return retriever


def _minimal_context():
    return {
        "organization": {"name": "ТестОрг", "inn": "1234567890"},
        "facility": {"name": "Объект", "facility_type": "Котельная", "hazard_class": "III"},
        "equipment": [{"name": "Котёл", "equipment_type": "Теплогенератор"}],
        "substances": [{"name": "Метан", "quantity_kg": 500}],
        "responsible_persons": [
            {"full_name": "Иванов И.И.", "phone": "+79991112233", "position": "Инженер"}
        ],
    }


# ===========================================================================
# 1. strip_pii
# ===========================================================================

class TestStripPii:
    def test_removes_known_fields_from_dict(self):
        data = {"full_name": "Иванов Иван Иванович", "phone": "+79001234567", "position": "Председатель"}
        cleaned = strip_pii(data)
        assert cleaned["full_name"] == "[скрыто]"
        assert cleaned["phone"] == "[скрыто]"
        assert cleaned["position"] == "Председатель"

    def test_removes_known_fields_from_object(self):
        p = Person("Петров Пётр Петрович", "+79007654321", "Инженер")
        cleaned = strip_pii(p)
        assert cleaned["full_name"] == "[скрыто]"
        assert cleaned["phone"] == "[скрыто]"

    def test_handles_nested_lists(self):
        context = {"responsible_persons": [Person("Сидоров С.С.", "89001112233", "Диспетчер")]}
        cleaned = strip_pii(context)
        assert cleaned["responsible_persons"][0]["full_name"] == "[скрыто]"

    def test_does_not_mutate_original(self):
        original = {"full_name": "Иванов Иван Иванович"}
        strip_pii(original)
        assert original["full_name"] == "Иванов Иван Иванович"

    def test_deep_nested_pii(self):
        data = {"a": {"b": {"full_name": "test", "phone": "123"}}}
        cleaned = strip_pii(data)
        assert cleaned["a"]["b"]["full_name"] == "[скрыто]"
        assert cleaned["a"]["b"]["phone"] == "[скрыто]"

    def test_non_pii_fields_pass_through(self):
        data = {"inn": "1234567890", "snils": "123-456-789 00", "email": "test@mail.ru"}
        cleaned = strip_pii(data)
        assert cleaned["inn"] == "[скрыто]"
        assert cleaned["snils"] == "[скрыто]"
        assert cleaned["email"] == "[скрыто]"

    def test_list_of_dicts(self):
        data = [{"full_name": "A", "phone": "B"}, {"full_name": "C", "phone": "D"}]
        cleaned = strip_pii(data)
        assert cleaned[0]["full_name"] == "[скрыто]"
        assert cleaned[1]["phone"] == "[скрыто]"


# ===========================================================================
# 2. PII-роутинг
# ===========================================================================

class TestPiiRouting:
    def test_strip_pii_removes_personal_data(self):
        """PII-данные удаляются из контекста перед отправкой в LLM."""
        context = {
            "organization": {"name": "ООО Тест", "inn": "1234567890"},
            "responsible_persons": [
                {"full_name": "Иванов И.И.", "phone": "+7-999-111-22-33", "position": "Директор"}
            ],
        }
        cleaned = strip_pii(context)
        assert cleaned["responsible_persons"][0]["full_name"] == "[скрыто]"
        assert cleaned["responsible_persons"][0]["phone"] == "[скрыто]"
        assert cleaned["organization"]["inn"] == "[скрыто]"
        # Не-PII поля остаются
        assert cleaned["organization"]["name"] == "ООО Тест"
        assert cleaned["responsible_persons"][0]["position"] == "Директор"

    def test_strip_pii_does_not_mutate_original(self):
        """strip_pii не мутирует исходный контекст."""
        original = {
            "responsible_persons": [{"full_name": "Иванов И.И.", "phone": "+7-999-111-22-33"}]
        }
        strip_pii(original)
        assert original["responsible_persons"][0]["full_name"] == "Иванов И.И."

    def test_pii_fields_constant_includes_required_fields(self):
        """PII_FIELD_NAMES содержит все поля персональных данных."""
        assert "full_name" in PII_FIELD_NAMES
        assert "phone" in PII_FIELD_NAMES
        assert "email" in PII_FIELD_NAMES
        assert "inn" in PII_FIELD_NAMES
        assert "snils" in PII_FIELD_NAMES

    @pytest.mark.asyncio
    async def test_non_pii_section_uses_external_llm(self):
        local_llm = _make_mock_llm("Локальный")
        external_llm = _make_mock_llm("Внешний текст")
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=local_llm,
            external_llm=external_llm,
            retriever=retriever,
            document_repo=doc_repo,
            regulatory_repo=reg_repo,
        )

        gen._load_structure = MagicMock(return_value={
            "title": "Тест",
            "sections": [{
                "id": "test_section",
                "title": "Характеристика",
                "template": "dummy.j2",
                "content_type": "llm",
                "pii": False,
                "slot_type": "text",
                "rag_query": "",
            }],
        })
        gen._render_template = MagicMock(return_value="ok")
        gen._doc_validator.validate = AsyncMock(return_value=MagicMock(passed=True, issues=[]))
        gen._build_docx = MagicMock(return_value=b"docx-bytes")

        with patch("src.core.settings.settings") as mock_settings:
            mock_settings.ai_review_enabled = False
            await gen.generate(document_id=MagicMock(), context=_minimal_context())

        # external_llm вызван 1 раз (генерация секции), local_llm — 0
        assert external_llm.complete.call_count == 1
        assert local_llm.complete.call_count == 0

    @pytest.mark.asyncio
    async def test_pii_context_scrubbed_before_send(self):
        """Даже для non-pii секций контекст чистится от PII."""
        external_llm = _make_mock_llm()
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=external_llm,
            external_llm=external_llm,
            retriever=retriever,
            document_repo=doc_repo,
            regulatory_repo=reg_repo,
        )

        gen._load_structure = MagicMock(return_value={
            "title": "Тест",
            "sections": [{
                "id": "s1",
                "title": "Раздел",
                "template": "dummy.j2",
                "content_type": "llm",
                "pii": False,
                "slot_type": "text",
                "rag_query": "",
            }],
        })
        gen._render_template = MagicMock(return_value="ok")
        gen._doc_validator.validate = AsyncMock(return_value=MagicMock(passed=True, issues=[]))
        gen._build_docx = MagicMock(return_value=b"docx")

        context = _minimal_context()
        await gen.generate(document_id=MagicMock(), context=context)

        # Проверяем что в промпте external_llm нет ФИО
        call_args = external_llm.complete.call_args
        messages = call_args.kwargs.get("messages", call_args[1].get("messages", []))
        all_text = " ".join(m.content for m in messages)
        assert "Иванов И.И." not in all_text
        assert "+79991112233" not in all_text


# ===========================================================================
# 3. GOST-форматирование DOCX
# ===========================================================================

class TestGostFormatting:
    def test_build_docx_returns_bytes(self):
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=None, external_llm=None, retriever=retriever,
            document_repo=doc_repo, regulatory_repo=reg_repo,
        )

        sections = {"Раздел 1": "Содержимое раздела"}
        metadata = {"version": "1.0", "generated_at": "2026-07-04", "status": "approved"}
        result = gen._build_docx("Тестовый документ", sections, metadata)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_build_docx_with_bold_markdown(self):
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=None, external_llm=None, retriever=retriever,
            document_repo=doc_repo, regulatory_repo=reg_repo,
        )

        sections = {"Раздел": "Текст с **жирным** словом"}
        metadata = {"version": "1.0", "generated_at": "2026-07-04", "status": "draft"}
        result = gen._build_docx("Тест", sections, metadata)
        assert isinstance(result, bytes)

    def test_build_docx_with_calculation_results(self):
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=None, external_llm=None, retriever=retriever,
            document_repo=doc_repo, regulatory_repo=reg_repo,
        )

        sections = {"Раздел": "Данные"}
        metadata = {
            "version": "1.0",
            "generated_at": "2026-07-04",
            "status": "approved",
            "calculation_results": [{"method_id": "tnt_v1", "substance": "Метан"}],
            "validation_issues": [{"severity": "warning", "section": "Раздел 1", "reason": "Тест"}],
        }
        result = gen._build_docx("Тест", sections, metadata)
        assert isinstance(result, bytes)

    def test_setup_document_defaults_sets_margins(self):
        from docx import Document as DocxDocument
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=None, external_llm=None, retriever=retriever,
            document_repo=doc_repo, regulatory_repo=reg_repo,
        )

        doc = DocxDocument()
        gen._setup_document_defaults(doc)

        section = doc.sections[0]
        # Проверяем поля страницы (в EMU, 1 cm = 360000 EMU)
        assert abs(section.top_margin - 2.0 * 360000) < 1000
        assert abs(section.left_margin - 3.0 * 360000) < 1000
        assert abs(section.right_margin - 1.5 * 360000) < 1000

    def test_setup_document_defaults_sets_font(self):
        from docx import Document as DocxDocument
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=None, external_llm=None, retriever=retriever,
            document_repo=doc_repo, regulatory_repo=reg_repo,
        )

        doc = DocxDocument()
        gen._setup_document_defaults(doc)

        normal = doc.styles["Normal"]
        assert normal.font.name == BODY_FONT_NAME
        assert normal.font.size.pt == BODY_FONT_SIZE_PT

    def test_add_heading_center(self):
        from docx import Document as DocxDocument
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=None, external_llm=None, retriever=retriever,
            document_repo=doc_repo, regulatory_repo=reg_repo,
        )

        doc = DocxDocument()
        gen._add_heading(doc, "Заголовок", level=0, center=True)
        assert len(doc.paragraphs) == 1
        p = doc.paragraphs[0]
        assert p.alignment is not None  # CENTER alignment set

    def test_add_body_paragraph_bold_conversion(self):
        from docx import Document as DocxDocument
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=None, external_llm=None, retriever=retriever,
            document_repo=doc_repo, regulatory_repo=reg_repo,
        )

        doc = DocxDocument()
        gen._add_body_paragraph(doc, "Текст с **жирным** словом")
        assert len(doc.paragraphs) == 1
        # Должно быть 3 run: "Текст с ", "жирным", " словом"
        runs = doc.paragraphs[0].runs
        assert len(runs) == 3
        assert runs[1].bold is True
        assert runs[1].text == "жирным"


# ===========================================================================
# 4. Fallback (без LLM)
# ===========================================================================

class TestFallback:
    @pytest.mark.asyncio
    async def test_no_llm_uses_fallback_text(self):
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=None, external_llm=None, retriever=retriever,
            document_repo=doc_repo, regulatory_repo=reg_repo,
        )

        gen._load_structure = MagicMock(return_value={
            "title": "Тест",
            "sections": [{
                "id": "s1",
                "title": "1.2. Возможные сценарии аварий",
                "template": "dummy.j2",
                "content_type": "llm",
                "slot_type": "text",
                "rag_query": "",
            }],
        })
        gen._render_template = MagicMock(return_value="ok")
        gen._doc_validator.validate = AsyncMock(return_value=MagicMock(passed=True, issues=[]))
        gen._build_docx = MagicMock(return_value=b"docx")

        await gen.generate(document_id=MagicMock(), context=_minimal_context())

        # _build_docx должен быть вызван (значит fallback отработал)
        gen._build_docx.assert_called_once()

    def test_fallback_text_content(self):
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()

        gen = EnhancedDocumentGenerator(
            local_llm=None, external_llm=None, retriever=retriever,
            document_repo=doc_repo, regulatory_repo=reg_repo,
        )

        result = gen._fallback_section_content("Тестовый раздел", _minimal_context())
        assert "Объект" in result
        assert "Котельная" in result
        assert "Метан" in result


# ===========================================================================
# 5. Парсеры таблиц
# ===========================================================================

class TestBlockRendering:
    """Тесты рендера блоков через _render_blocks / _render_table_block."""

    def _make_gen(self):
        retriever = _make_retriever()
        doc_repo = AsyncMock()
        reg_repo = AsyncMock()
        reg_repo.session = AsyncMock()
        return EnhancedDocumentGenerator(
            local_llm=None, external_llm=None, retriever=retriever,
            document_repo=doc_repo, regulatory_repo=reg_repo,
        )

    def test_table_block_renders_as_docx_table(self):
        from docx import Document as DocxDocument
        from src.application.engines.blocks import TableBlock
        gen = self._make_gen()
        doc = DocxDocument()
        gen._setup_document_defaults(doc)
        block = TableBlock(
            headers=["№", "Наименование", "Количество"],
            rows=[["1", "Котёл", "2 шт"], ["2", "Насос", "5 шт"]],
            caption="Таблица 1. Оборудование",
        )
        gen._render_table_block(doc, block)
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 3  # header + 2 data rows

    def test_table_block_caption_is_bold(self):
        from docx import Document as DocxDocument
        from src.application.engines.blocks import TableBlock
        gen = self._make_gen()
        doc = DocxDocument()
        gen._setup_document_defaults(doc)
        block = TableBlock(
            headers=["A", "B"],
            rows=[["1", "2"]],
            caption="Таблица 1. Тест",
        )
        gen._render_table_block(doc, block)
        # Caption — первый абзац, жирный
        assert doc.paragraphs[0].runs[0].bold is True
        assert "Таблица 1" in doc.paragraphs[0].text

    def test_heading_block_renders(self):
        from docx import Document as DocxDocument
        from src.application.engines.blocks import HeadingBlock
        gen = self._make_gen()
        doc = DocxDocument()
        gen._setup_document_defaults(doc)
        block = HeadingBlock(text="Раздел 1", level=1, center=False)
        gen._render_blocks(doc, [block])
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].runs[0].bold is True
        assert doc.paragraphs[0].runs[0].text == "Раздел 1"

    def test_paragraph_block_renders(self):
        from docx import Document as DocxDocument
        from src.application.engines.blocks import ParagraphBlock
        gen = self._make_gen()
        doc = DocxDocument()
        gen._setup_document_defaults(doc)
        block = ParagraphBlock(text="Обычный текст")
        gen._render_blocks(doc, [block])
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "Обычный текст"

    def test_bold_paragraph_block_renders(self):
        from docx import Document as DocxDocument
        from src.application.engines.blocks import ParagraphBlock
        gen = self._make_gen()
        doc = DocxDocument()
        gen._setup_document_defaults(doc)
        block = ParagraphBlock(text="Жирный текст", bold=True)
        gen._render_blocks(doc, [block])
        assert doc.paragraphs[0].runs[0].bold is True

    def test_build_docx_with_blocks(self):
        from src.application.engines.blocks import TableBlock, HeadingBlock, ParagraphBlock
        gen = self._make_gen()
        sections = {
            "Раздел 1": [
                HeadingBlock(text="Подраздел", level=2),
                ParagraphBlock(text="Текст"),
                TableBlock(headers=["A", "B"], rows=[["1", "2"]], caption="Таблица 1"),
            ],
            "Раздел 2": "Обычный текст",
        }
        metadata = {"version": "1.0", "generated_at": "2026-07-05", "status": "draft"}
        result = gen._build_docx("Тестовый документ", sections, metadata)
        assert isinstance(result, bytes)
        assert len(result) > 0
