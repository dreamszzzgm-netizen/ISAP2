"""File parsers for Smart Import Center."""
from __future__ import annotations

import csv
import io
from dataclasses import dataclass
from pathlib import Path
from typing import Any


@dataclass(frozen=True)
class ParsedTable:
    headers: list[str]
    rows: list[dict[str, Any]]


class SmartImportParser:
    """Parse Excel/CSV files into a flat table."""

    def parse_bytes(self, filename: str, content: bytes) -> ParsedTable:
        suffix = Path(filename).suffix.lower()
        if suffix in {".xlsx", ".xlsm"}:
            return self._parse_xlsx(content)
        if suffix in {".csv", ".txt"}:
            return self._parse_csv(content)
        if suffix == ".docx":
            return self._parse_docx(content)
        raise ValueError("Поддерживаются только .xlsx, .xlsm, .csv, .txt, .docx")

    def _parse_csv(self, content: bytes) -> ParsedTable:
        text = content.decode("utf-8-sig")
        sample = text[:4096]
        dialect = csv.Sniffer().sniff(sample, delimiters=";,\t,") if sample.strip() else csv.excel
        reader = csv.DictReader(io.StringIO(text), dialect=dialect)
        headers = [str(header or "").strip() for header in (reader.fieldnames or [])]
        rows = [self._clean_row(row) for row in reader]
        return ParsedTable(headers=headers, rows=rows)

    def _parse_xlsx(self, content: bytes) -> ParsedTable:
        try:
            from openpyxl import load_workbook
        except ModuleNotFoundError as exc:
            raise RuntimeError("Для импорта Excel требуется openpyxl") from exc

        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sheet = workbook.active
        rows_iter = sheet.iter_rows(values_only=True)
        header_values = next(rows_iter, None)
        if not header_values:
            return ParsedTable(headers=[], rows=[])
        headers = [str(value or "").strip() for value in header_values]
        rows: list[dict[str, Any]] = []
        for values in rows_iter:
            raw = {headers[i]: values[i] if i < len(values) else None for i in range(len(headers))}
            if any(value not in (None, "") for value in raw.values()):
                rows.append(self._clean_row(raw))
        return ParsedTable(headers=headers, rows=rows)


    def _parse_docx(self, content: bytes) -> ParsedTable:
        """Extract questionnaire-like key/value data from DOCX.

        DOCX import is intentionally converted to a single preview row. The user
        must confirm it before the data becomes a PMLA questionnaire.
        """
        try:
            from docx import Document as DocxDocument
        except ModuleNotFoundError as exc:
            raise RuntimeError("Для импорта DOCX требуется python-docx") from exc

        import re

        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        table_pairs: dict[str, Any] = {}
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2 and cells[0] and cells[1]:
                    table_pairs[cells[0]] = cells[1]

        full_text = "\n".join([*paragraphs, *[f"{k}: {v}" for k, v in table_pairs.items()]])

        def find_value(patterns: list[str]) -> str:
            for pattern in patterns:
                match = re.search(pattern, full_text, flags=re.IGNORECASE | re.MULTILINE)
                if match:
                    value = match.group(1).strip(" :—-\t\r\n")
                    value = re.split(r"\n|\r", value)[0].strip()
                    return value[:1000]
            return ""

        def find_list(anchor_words: list[str]) -> str:
            lines = [line.strip(" •-\t") for line in full_text.splitlines()]
            capture = False
            result: list[str] = []
            for line in lines:
                low = line.lower().replace("ё", "е")
                if any(anchor in low for anchor in anchor_words):
                    capture = True
                    continue
                if capture:
                    if not line:
                        continue
                    if re.match(r"^\d+(\.\d+)*\s+", line) and result:
                        break
                    if len(result) >= 8:
                        break
                    if len(line) > 3:
                        result.append(line)
            return "; ".join(result)

        lower_text = full_text.lower().replace("ё", "е")
        has_incidents = ""
        incident_description = ""
        if "аварий" in lower_text or "инцидент" in lower_text:
            if "не зарегистр" in lower_text or "не было" in lower_text or "отсутств" in lower_text:
                has_incidents = "нет"
                incident_description = "За период эксплуатации аварии и инциденты не зарегистрированы."
            else:
                has_incidents = "требует проверки"
                incident_description = find_list(["авар", "инцидент"]) or "Найдены упоминания аварий/инцидентов, требуется проверка пользователем."

        row = {
            "Организация": find_value([r"(?:организация|заказчик|эксплуатирующая организация)\s*[:—-]\s*(.+)"]),
            "ОПО": find_value([r"(?:опасный производственный объект|наименование опо|объект)\s*[:—-]\s*(.+)"]),
            "Регистрационный номер ОПО": find_value([r"(?:регистрационный номер(?:\s+опо)?|рег\.\s*номер)\s*[:—-]\s*([А-ЯA-Z0-9\-]+)"]),
            "Были аварии/инциденты": has_incidents,
            "Описание аварий/инцидентов": incident_description,
            "Режим работы": find_value([r"(?:режим работы|сменность)\s*[:—-]\s*(.+)"]),
            "Персонал в смену": find_value([r"(?:персонал в смену|численность персонала|людей в смену)\s*[:—-]\s*(.+)"]),
            "Сценарии": find_list(["сценарии авар", "возможные сценарии", "наиболее вероятные аварии"]),
            "Другое / пользовательские сценарии": "",
            "Силы и средства": find_list(["силы и средства", "оснащение", "аварийный запас"]),
            "ПАСФ": find_value([r"(?:пасф|асф|аварийно[-\s]спасательное формирование)\s*[:—-]\s*(.+)"]),
            "Финансовый резерв": find_value([r"(?:финансовый резерв|резерв финансовых средств)\s*[:—-]\s*(.+)"]),
            "Тренировки": find_value([r"(?:тренировки|учения|учебно[-\s]тренировочные занятия)\s*[:—-]\s*(.+)"]),
            "Приложения": find_list(["приложения", "перечень приложений"]),
            "_docx_extracted_text_preview": full_text[:4000],
        }
        headers = list(row.keys())
        return ParsedTable(headers=headers, rows=[row])

    @staticmethod
    def _clean_row(row: dict[str, Any]) -> dict[str, Any]:
        cleaned: dict[str, Any] = {}
        for key, value in row.items():
            if key is None:
                continue
            header = str(key).strip()
            if not header:
                continue
            cleaned[header] = value
        return cleaned
