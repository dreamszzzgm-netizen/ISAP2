"""Парсер данных ОПО из Word-документов.

Поддерживаемые форматы:
1. Паспорт ОПО / Сведения, характеризующие ОПО (табличная форма)
2. Произвольный DOCX с ключевыми полями
"""
import io
import re
from datetime import datetime

from docx import Document


class WordImportService:
    """Извлечение полей ОПО из .docx файла."""

    FIELD_PATTERNS = {
        'f1_1': [
            r'Наименование\s+ОПО[:\s]+(.+)',
            r'Полное\s+наименование[:\s]+(.+)',
            r'^(\S.+(?:станция|компрессорная|котельная|газопровод|насосная|хранилище).+)$',
        ],
        'f1_4': [
            r'Адрес[:\s]+(.+)',
            r'Местоположение[:\s]+(.+)',
            r'Место\s+нахождения[:\s]+(.+)',
        ],
        'danger_class': [
            r'Класс\s+опасности[:\s]+([IIV]+)',
            r'Класс[:\s]+([IIV]+)',
        ],
        'f1_5': [
            r'ОКТМО[:\s]*(\d+)',
        ],
        'f1_7_1': [
            r'(?:Наименование|Собственник)[:\s]+(.+)',
        ],
        'f1_7_2': [
            r'ИНН\s+собственника[:\s]*(\d+)',
            r'ИНН[:\s]*(\d{10,12})',
        ],
        'f1_2': [
            r'Типовое\s+наименование[:\s]+(.+)',
            r'Тип\s+объекта[:\s]+(.+)',
        ],
        'f1_3': [
            r'Цифровое\s+обозначение[:\s]+(.+)',
            r'Код\s+отрасли[:\s]+(.+)',
        ],
        'f1_6': [
            r'Дата\s+ввода\s+в\s+эксплуатацию[:\s]*(\d{2}\.\d{2}\.\d{4})',
        ],
    }

    # Поля в табличной форме "Сведения, характеризующие ОПО"
    TABLE_FIELD_MAP = {
        '1.1': 'f1_1',
        '1.2': 'f1_2',
        '1.3': 'f1_3',
        '1.4': 'f1_4',
        '1.5': 'f1_5',
        '3.1': 'danger_class_I',
        '3.2': 'danger_class_II',
        '3.3': 'danger_class_III',
        '3.4': 'danger_class_IV',
    }

    def import_from_word(self, file_content: bytes) -> dict:
        """Парсит .docx и возвращает dict с полями формы."""
        doc = Document(io.BytesIO(file_content))

        # === Этап 1: парсим параграфы (plain text) ===
        lines = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)

        full_text = '\n'.join(lines)

        extracted = {}
        for field, patterns in self.FIELD_PATTERNS.items():
            for pat in patterns:
                m = re.search(pat, full_text, re.IGNORECASE | re.MULTILINE)
                if m:
                    val = m.group(1).strip()
                    if field == 'f1_6':
                        try:
                            val = datetime.strptime(val, '%d.%m.%Y').strftime('%Y-%m-%d')
                        except ValueError:
                            continue
                    if field == 'danger_class':
                        cmap = {'1': 'I', '2': 'II', '3': 'III', '4': 'IV'}
                        val = cmap.get(val, val.upper())
                    extracted[field] = val
                    break

        # === Этап 2: парсим таблицы (форма "Сведения, характеризующие ОПО") ===
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2:
                    continue

                key_cell = cells[0].strip()
                # Ищем ключ вида "1.1", "1.2" в начале ячейки
                key_match = re.match(r'^(\d+\.\d+)', key_cell)
                if key_match:
                    table_key = key_match.group(1)
                    target = self.TABLE_FIELD_MAP.get(table_key)
                    if target and len(cells) >= 2:
                        val = cells[1].strip()
                        if val and target not in extracted:
                            # Спецобработка для класса опасности
                            if target == 'danger_class_I' and val.upper() in ('V', 'X', '✓', '✔', '+', 'ДА'):
                                extracted['danger_class'] = 'I'
                            elif target == 'danger_class_II' and val.upper() in ('V', 'X', '✓', '✔', '+', 'ДА'):
                                extracted['danger_class'] = 'II'
                            elif target == 'danger_class_III' and val.upper() in ('V', 'X', '✓', '✔', '+', 'ДА'):
                                extracted['danger_class'] = 'III'
                            elif target == 'danger_class_IV' and val.upper() in ('V', 'X', '✓', '✔', '+', 'ДА'):
                                extracted['danger_class'] = 'IV'
                            elif target.startswith('f1_'):
                                extracted[target] = val

                # Также парсим ИНН из таблиц
                if 'ИНН' in key_cell:
                    inn_match = re.search(r'(\d{10,12})', cells[1] if len(cells) > 1 else '')
                    if inn_match and 'f1_7_2' not in extracted:
                        extracted['f1_7_2'] = inn_match.group(1)

                # ОГРНИП / ОГРН
                if 'ОГРН' in key_cell:
                    ogrn_match = re.search(r'(\d{13,15})', cells[1] if len(cells) > 1 else '')
                    if ogrn_match:
                        extracted['ogrn'] = ogrn_match.group(1)

                # ФИО из таблицы заявителя
                if 'Фамилия' in key_cell and 'имя' in key_cell and len(cells) > 1:
                    extracted['owner_name'] = cells[1].strip()

                # Адрес из таблицы заявителя
                if 'Адрес места жительства' in key_cell and len(cells) > 1:
                    if 'f1_4' not in extracted:
                        extracted['f1_4'] = cells[1].strip()

        # === Этап 3: доизвлечение из таблиц (простые пары ключ-значение) ===
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    key = cells[0]
                    val = cells[1]
                    if 'Класс' in key and 'опасности' in key:
                        cmap = {'1': 'I', '2': 'II', '3': 'III', '4': 'IV'}
                        v = val.upper().strip()
                        if v in ('I', 'II', 'III', 'IV'):
                            extracted['danger_class'] = v
                        elif v in cmap:
                            extracted['danger_class'] = cmap[v]
                    elif 'ОКТМО' in key:
                        extracted['f1_5'] = val
                    elif 'ИНН' in key:
                        inn_match = re.search(r'(\d{10,12})', val)
                        if inn_match and 'f1_7_2' not in extracted:
                            extracted['f1_7_2'] = inn_match.group(1)

        return {
            'success': bool(extracted),
            'data': extracted,
            'warnings': [] if extracted.get('f1_1') else ['Не удалось извлечь наименование ОПО'],
        }
