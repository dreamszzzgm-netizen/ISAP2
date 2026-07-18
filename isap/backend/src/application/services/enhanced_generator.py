"""Улучшенный генератор документов ПМЛА."""
import io
import logging
import json
import re
from datetime import UTC, datetime
from pathlib import Path
from uuid import UUID

from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.application.engines.blocks import (
    Block,
    HeadingBlock,
    ImageBlock,
    ParagraphBlock,
    TableBlock,
    deserialize_blocks,
    serialize_blocks,
)
from src.application.services.calculations import CalculationRegistry
from src.application.services.calculations.validation import CalculationValidator
from src.application.services.prompts import (
    CURRENT_PROMPT_VERSION,
    SYSTEM_PROMPT,
    build_section_prompt,
)
from src.application.services.types import GeneratedDocument
from src.application.services.validation import DocumentValidator
from src.infrastructure.export.docx_helpers import (
    add_appendices_manifest,
    add_approval_sheet,
    add_body_paragraph as helper_add_body_paragraph,
    add_correction_journal,
    add_data_table,
    add_heading as helper_add_heading,
    add_kv_table,
    add_toc_placeholder,
    configure_heading_styles,
    create_title_page,
    safe_text,
    sanitize_cyrillic_text,
    set_default_font,
    set_document_margins,
    strip_html,
)
from src.infrastructure.llm.providers import LLMMessage, LLMProvider
from src.infrastructure.rag.pipeline import Retriever
from src.infrastructure.repositories.document_repo import DocumentRepository
from src.infrastructure.repositories.regulatory_repo import RegulatoryRepository
from src.infrastructure.repositories.scenario_matrix_repo import (
    ScenarioMatrixRepository,
)

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent.parent.parent.parent / "templates"
CURRENT_TEMPLATE_VERSION = "1.0.0"

# Значения подобраны по эталонному документу ПМЛА (см. анализ styles.xml/sectPr)
BODY_FONT_NAME = "Times New Roman"
BODY_FONT_SIZE_PT = 12
HEADING_FONT_SIZE_PT = 14
PAGE_MARGINS_CM = {"top": 2.0, "bottom": 2.0, "left": 3.0, "right": 1.5}
FIRST_LINE_INDENT_CM = 1.25

# Поля из ResponsiblePersonModel / OrganizationModel, которые считаются
# персональными данными и не должны попадать в промпт внешнего LLM.
# Подтверждено по backend/src/infrastructure/database/models.py:
#   ResponsiblePersonModel: full_name, position, role, phone, email
#   OrganizationModel: phone, email (контакты организации — тоже чистим
#   по умолчанию, при желании можно вывести из списка отдельно)
PII_FIELD_NAMES = {"full_name", "phone", "email", "inn", "snils", "address", "birth_date"}


class PiiRoutingError(RuntimeError):
    """Секция помечена pii=true, но локальный LLM-провайдер не настроен."""


def _looks_like_pii_key(key: str) -> bool:
    return key.strip().lower() in PII_FIELD_NAMES


def strip_pii(value):
    """
    Рекурсивно возвращает копию структуры данных, где значения по
    персональным ключам (full_name/phone/email) заменены плейсхолдером.
    Используется перед сборкой промпта для внешнего (облачного) LLM.
    """
    if isinstance(value, dict):
        return {
            k: ("[скрыто]" if _looks_like_pii_key(str(k)) else strip_pii(v))
            for k, v in value.items()
        }
    if isinstance(value, list):
        return [strip_pii(v) for v in value]
    if hasattr(value, "__dict__") and not isinstance(value, (str, int, float, bool, type(None))):
        return {
            k: ("[скрыто]" if _looks_like_pii_key(k) else strip_pii(v))
            for k, v in vars(value).items()
            if not k.startswith("_")
        }
    return value


def _serialize_sections(sections: dict[str, str | list[Block]]) -> dict:
    """Сериализует секции для сохранения в JSONB."""
    result = {}
    for title, value in sections.items():
        if isinstance(value, list):
            result[title] = {"__blocks__": True, "data": serialize_blocks(value)}
        else:
            result[title] = {"__blocks__": False, "data": value}
    return result


def _deserialize_sections(data: dict) -> dict[str, str | list[Block]]:
    """Десериализует секции из JSONB."""
    result = {}
    for title, value in data.items():
        if isinstance(value, dict) and "__blocks__" in value:
            if value["__blocks__"]:
                result[title] = deserialize_blocks(value["data"])
            else:
                result[title] = value["data"]
        else:
            # Обратная совместимость со старым форматом (просто строка)
            result[title] = value
    return result


class EnhancedDocumentGenerator:
    """
    Улучшенный генератор ПМЛА с:
    - Расчётным блоком (method_id)
    - Защитой числовых полей
    - Автоматической валидацией
    - Версионированием
    - Раздельной маршрутизацией LLM: локальный провайдер для секций с
      персональными данными (section["pii"] = true) и для AI-ревью готового
      документа; внешний провайдер — для остального текста, с обязательной
      очисткой контекста через strip_pii() перед отправкой.
    """

    def __init__(
        self,
        local_llm: LLMProvider | None,
        external_llm: LLMProvider | None,
        retriever: Retriever | None,
        document_repo: DocumentRepository,
        regulatory_repo: RegulatoryRepository,
        scenario_matrix_repo: ScenarioMatrixRepository | None = None,
        sample_repo=None,
    ):
        self._local_llm = local_llm
        self._external_llm = external_llm
        self._retriever = retriever
        self._document_repo = document_repo
        self._regulatory_repo = regulatory_repo
        self._scenario_matrix_repo = scenario_matrix_repo
        self._sample_repo = sample_repo
        self._calc_validator = CalculationValidator()
        self._doc_validator = DocumentValidator(regulatory_repo)

    async def _select_scenarios(self, context: dict) -> list[dict]:
        """
        Детерминированный выбор сценариев из матрицы.
        Выбирает по facility_type × hazard_class (регистронезависимо).
        """
        if self._scenario_matrix_repo is None:
            return []

        facility = context.get("facility", {})
        facility_type = facility.get("facility_type", "")
        hazard_class = str(facility.get("hazard_class", ""))

        if not facility_type or not hazard_class:
            return []

        scenarios = await self._scenario_matrix_repo.get_by_type_and_class(
            facility_type=facility_type,
            hazard_class=hazard_class,
        )

        return [
            {
                "id": s.scenario_code,
                "name": s.scenario_name,
                "factor_type": s.factor_type or "",
                "calculation_method": s.calculation_method or "",
                "probability": s.probability or "средняя",
            }
            for s in scenarios
        ]

    def _enrich_context(self, context: dict, scenarios: list[dict], calculations: list[dict]) -> dict:
        """Обогащение контекста дополнительными переменными для шаблонов."""
        from datetime import datetime

        facility = context.get("facility", {})
        persons = context.get("responsible_persons", [])

        # Approver (для титульного листа)
        approver = {"name": "—", "position": "—"}
        if persons:
            p = persons[0] if isinstance(persons[0], dict) else {
                "full_name": getattr(persons[0], "full_name", "—"),
                "position": getattr(persons[0], "position", "—"),
            }
            approver = {
                "name": p.get("full_name", "—") if isinstance(p, dict) else p.get("full_name", "—"),
                "position": p.get("position", "—") if isinstance(p, dict) else p.get("position", "—"),
            }

        # Personnel (для раздела 11)
        personnel = []
        for p in persons:
            if isinstance(p, dict):
                personnel.append(p)
            else:
                personnel.append({
                    "full_name": getattr(p, "full_name", "—"),
                    "position": getattr(p, "position", "—"),
                    "role": getattr(p, "role", "—"),
                    "phone": getattr(p, "phone", "—"),
                })

        enriched = {
            **context,
            "year": datetime.now().year,
            "document_date": context.get("document_date") or datetime.now().strftime("%d.%m.%Y"),
            "approver": approver,
            "personnel": personnel,
            "scenarios": scenarios,
            "calculation_results": calculations,
            "facility_coords": {
                "latitude": facility.get("latitude"),
                "longitude": facility.get("longitude"),
            },
            "material_reserve": {
                "sip_amount": "—",
                "sip_source": "—",
                "fire_amount": "—",
                "fire_source": "—",
                "maintenance_amount": "—",
                "maintenance_source": "—",
                "training_amount": "—",
                "training_source": "—",
            },
            "material_reserves": [],
        }
        enriched["material_reserve"].update(context.get("material_reserve") or {})

        # RAG context for generated blocks — pre-fetch for all generated sections
        from src.application.services.pmla_rag_adapter import PmlaRagAdapter
        try:
            rag_adapter = PmlaRagAdapter()
            from src.application.services.pmla_assembly_blocks import get_generated_sections
            rag_contexts: dict[str, dict] = {}
            for sid in get_generated_sections():
                rag_ctx = rag_adapter.get_context(enriched, sid)
                if not rag_ctx.is_empty:
                    rag_contexts[sid] = {
                        "chunks": [
                            {"source_id": c.source_id, "source_title": c.source_title, "text": c.text}
                            for c in rag_ctx.chunks
                        ],
                        "summary": rag_ctx.summary,
                    }
            enriched["rag_contexts"] = rag_contexts
        except Exception as e:
            logger.warning("RAG adapter failed: %s", e)
            enriched["rag_contexts"] = {}

        # Манифест приложений: если он ещё не задан явно, синтезируем из
        # канонических записей реестра (5 приложений), статуса наличия из
        # attachments_checklist анкеты и выбранных документов ПАСФ.
        if not enriched.get("appendices_manifest"):
            enriched["appendices_manifest"] = _synthesize_appendices_manifest(
                enriched.get("attachments_checklist") or [],
                enriched.get("pasf_documents") or [],
            )

        return enriched

    async def _generate_sections_via_engines(
        self,
        context: dict,
        calculation_results: list[dict],
        selected_scenarios: list[dict],
        document_id: UUID,
        regenerate_sections: list[str] | None = None,
    ) -> dict[str, str | list[Block]]:
        """
        Генерация разделов через EngineRouter (6 движков).
        Заменяет монолитный цикл генерации на маршрутизацию по движкам.

        Возвращает {title: blocks_list | content_str} — реальные объекты для _build_docx.
        Для сохранения в БД используй serialize_sections().
        """
        from src.application.services.engine_integration import (
            build_document_context,
            create_engine_router,
        )

        # Создаём DocumentContext из raw dict
        doc_context = build_document_context(
            raw_context=context,
            calculation_results=calculation_results,
            scenarios=selected_scenarios,
        )

        # Создаём EngineRouter
        engine_router = create_engine_router(
            llm_provider=self._external_llm,
            retriever=self._retriever,
        )

        # Если перегенерация — загружаем существующие разделы из БД (десериализуем)
        all_sections: dict[str, str | list[Block]] = {}
        if regenerate_sections:
            doc = await self._document_repo.get(document_id)
            if doc and doc.rendered_sections:
                all_sections = _deserialize_sections(doc.rendered_sections)

        # Генерируем все разделы через движки
        structure = engine_router.load_structure("pmla")
        new_sections: dict[str, str | list[Block]] = {}

        for section_def in structure["sections"]:
            section_id = section_def["id"]
            title = section_def["title"]

            # Если задана перегенерация — пропускаем не указанные разделы
            if regenerate_sections and section_id not in regenerate_sections:
                continue

            result = await engine_router.generate_section(
                section_id=section_id,
                section_def=section_def,
                context=doc_context,
            )
            # Если движок вернул блоки — передаём их, иначе — текст
            if result.blocks:
                new_sections[title] = result.blocks
            else:
                new_sections[title] = result.content

        # Мерж: существующие + новые (новые перезаписывают старые)
        if regenerate_sections:
            all_sections.update(new_sections)
            return all_sections
        return new_sections

    async def generate(
        self,
        document_id: UUID,
        context: dict,
        regenerate_sections: list[str] | None = None,
    ) -> GeneratedDocument:
        """
        Полный пайплайн генерации ПМЛА:
        A. Обогащение данных (геокодирование, службы, СИЗ)
        B. Расчётный блок
        C. Генерация текста (LLM + RAG)
        D. Автоматическая валидация
        """
        structure = self._load_structure("pmla")

        # Этап B: Расчётный блок
        calculation_results = await self._run_calculations(context)

        # Логгер для метода
        import logging
        logger = logging.getLogger(__name__)

        # Этап B2: Выбор сценариев из матрицы (детерминированно)
        selected_scenarios = await self._select_scenarios(context)

        # Обогащение контекста для шаблонов
        context = self._enrich_context(context, selected_scenarios, calculation_results)

        # Этап C: Генерация разделов через EngineRouter (6 движков)
        rendered_sections = await self._generate_sections_via_engines(
            context=context,
            calculation_results=calculation_results,
            selected_scenarios=selected_scenarios,
            document_id=document_id,
            regenerate_sections=regenerate_sections,
        )

        # Этап D: Автоматическая валидация
        validation = await self._doc_validator.validate(
            rendered_sections=rendered_sections,
            context=context,
            calculation_results=[
                {"method_id": r["method_id"], "results": r["results"]}
                for r in calculation_results
            ],
        )

        # Этап D2: AI-ревью (если LLM доступен, валидация прошла, и включено в настройках)
        # ВАЖНО: ревью смотрит на уже ОТРЕНДЕРЕННЫЙ документ целиком, где в
        # data-секциях (Таблица 14, лист ознакомления, титульный лист) уже
        # подставлены реальные ФИО и телефоны. Поэтому ревью может идти
        # ТОЛЬКО через local_llm — здесь нет "безопасной части" документа,
        # которую можно было бы отдать внешнему провайдеру.
        ai_review_result = None
        from src.core.settings import settings
        if validation.passed and self._local_llm is not None and settings.ai_review_enabled:
            try:
                from src.application.services.ai_reviewer import AIReviewer
                reviewer = AIReviewer(self._local_llm)
                ai_review_result = await reviewer.review(rendered_sections, context)
            except Exception as e:
                logger.warning("AI review failed: %s", e)

        # Определяем статус
        if not validation.passed:
            status = "auto_validation_failed"
        elif ai_review_result and ai_review_result.decision == "auto_approve":
            status = "pending_review"  # всё равно требуем человеческое утверждение
        else:
            status = "pending_review"

        # Сборка DOCX
        metadata = {
            "version": "1.0",
            "generated_at": datetime.now(UTC).replace(tzinfo=None).isoformat(),
            "status": status,
            "prompt_version": CURRENT_PROMPT_VERSION,
            "template_version": CURRENT_TEMPLATE_VERSION,
            "calculation_results": calculation_results,
            "validation_issues": [
                {"section": i.section, "reason": i.reason, "severity": i.severity}
                for i in validation.issues
            ],
            "context": context,
        }
        docx_bytes = self._build_docx(structure["title"], rendered_sections, metadata)
        metadata["debug_report"] = {
            "context_keys": sorted(context.keys()),
            "section_count": len(rendered_sections),
            "section_titles": list(rendered_sections.keys()),
            "calculation_count": len(calculation_results),
            "validation_passed": validation.passed,
            "validation_issue_count": len(validation.issues),
            "ai_review_ran": ai_review_result is not None,
            "docx_size_bytes": len(docx_bytes),
            "docx_created": bool(docx_bytes),
        }

        # Сохранение в БД (сериализуем блоки в JSON)
        # Определяем начальный review_status на основе качества
        review_status = "needs_review"
        if validation.issues:
            has_critical = any(i.severity == "critical" for i in validation.issues)
            if has_critical:
                review_status = "needs_changes"

        await self._document_repo.update(
            document_id,
            {
                "content_docx": docx_bytes,
                "rendered_sections": _serialize_sections(rendered_sections),
                "status": status,
                "review_status": review_status,
                "generation_meta": metadata,
                "updated_at": datetime.now(UTC).replace(tzinfo=None),
            },
        )

        # Сохранение результатов расчётов
        for calc in calculation_results:
            from src.infrastructure.database.models import CalculationResultModel

            calc_model = CalculationResultModel(
                document_id=document_id,
                method_id=calc["method_id"],
                input_params=calc["input_params"],
                results=calc["results"],
                validation_status=calc["validation_status"],
            )
            await self._document_repo.add_calculation_result(calc_model)

        # Сохранение версии
        from src.infrastructure.database.models import DocumentVersionModel

        latest_version = await self._document_repo.get_latest_version(document_id)
        next_version = (latest_version.version_number + 1) if latest_version else 1

        # Снимок нормативов на момент генерации
        from src.application.services.regulatory_snapshot import (
            collect_regulatory_snapshot,
        )

        regulatory_snapshot = await collect_regulatory_snapshot(
            getattr(self._document_repo, "session", None)
        )

        version = DocumentVersionModel(
            document_id=document_id,
            version_number=next_version,
            input_data=context,
            prompt_version=CURRENT_PROMPT_VERSION,
            template_version=CURRENT_TEMPLATE_VERSION,
            calculation_results={"results": calculation_results},
            regulatory_snapshot=regulatory_snapshot,
            ai_review_confidence=ai_review_result.overall_confidence if ai_review_result else None,
            ai_review_decision=ai_review_result.decision if ai_review_result else None,
            ai_review_items=[
                {"id": i.check_id, "name": i.check_name, "passed": i.passed, "confidence": i.confidence, "details": i.details}
                for i in ai_review_result.items
            ] if ai_review_result else [],
            ai_review_summary=ai_review_result.summary if ai_review_result else None,
            content_docx=docx_bytes,
        )
        await self._document_repo.add_version(version)

        return GeneratedDocument(
            document_id=document_id,
            docx_bytes=docx_bytes,
            version_number=next_version,
            status=status,
        )

    async def _run_calculations(self, context: dict) -> list[dict]:
        """Запуск расчётного блока для всех веществ."""
        results = []
        substances = context.get("substances", [])

        for substance in substances:
            name = substance.get("name", "")
            quantity = substance.get("quantity_kg", 0)
            hazard_props = substance.get("hazard_properties", {})

            # Расчёт зоны взрыва (если есть энергия взрыва)
            if "explosion_energy_mj" in hazard_props:
                try:
                    from src.application.services.calculations.types import (
                        ExplosionParams,
                    )

                    params = ExplosionParams(
                        substance_name=name,
                        quantity_kg=quantity,
                        explosion_energy_mj=hazard_props["explosion_energy_mj"],
                        physical_state=hazard_props.get("physical_state", "газ"),
                    )
                    validation = self._calc_validator.validate_explosion(params)
                    if validation.is_valid:
                        calc_result = CalculationRegistry.calculate(
                            "tnt_equivalent_v1", params
                        )
                        results.append(
                            {
                                "method_id": "tnt_equivalent_v1",
                                "substance": name,
                                "input_params": calc_result.input_params,
                                "results": calc_result.results,
                                "validation_status": "valid",
                            }
                        )
                except Exception as e:  # noqa: BLE001 - расчёт не должен рвать генерацию
                    import logging

                    logging.getLogger(__name__).warning(
                        "Расчёт зоны взрыва для '%s' не выполнен: %s: %s",
                        name,
                        type(e).__name__,
                        str(e)[:200],
                    )
                    results.append(
                        {
                            "method_id": "tnt_equivalent_v1",
                            "substance": name,
                            "validation_status": "error",
                            "error": f"{type(e).__name__}: {e}",
                        }
                    )

            # Расчёт теплового излучения
            if "combustion_energy_mj_kg" in hazard_props:
                try:
                    from src.application.services.calculations.types import (
                        ThermalParams,
                    )

                    params = ThermalParams(
                        substance_name=name,
                        quantity_kg=quantity,
                        combustion_energy_mj_kg=hazard_props["combustion_energy_mj_kg"],
                    )
                    validation = self._calc_validator.validate_thermal(params)
                    if validation.is_valid:
                        calc_result = CalculationRegistry.calculate(
                            "thermal_radiation_v1", params
                        )
                        results.append(
                            {
                                "method_id": "thermal_radiation_v1",
                                "substance": name,
                                "input_params": calc_result.input_params,
                                "results": calc_result.results,
                                "validation_status": "valid",
                            }
                        )
                except Exception as e:  # noqa: BLE001 - расчёт не должен рвать генерацию
                    import logging

                    logging.getLogger(__name__).warning(
                        "Расчёт теплового излучения для '%s' не выполнен: %s: %s",
                        name,
                        type(e).__name__,
                        str(e)[:200],
                    )
                    results.append(
                        {
                            "method_id": "thermal_radiation_v1",
                            "substance": name,
                            "validation_status": "error",
                            "error": f"{type(e).__name__}: {e}",
                        }
                    )

            # Расчёт токсического поражения
            if "mac_mg_m3" in hazard_props:
                try:
                    from src.application.services.calculations.types import ToxicParams

                    params = ToxicParams(
                        substance_name=name,
                        quantity_kg=quantity,
                        mac_mg_m3=hazard_props["mac_mg_m3"],
                        lc50_mg_m3=hazard_props.get("lc50_mg_m3"),
                        physical_state=hazard_props.get("physical_state", "газ"),
                    )
                    validation = self._calc_validator.validate_toxic(params)
                    if validation.is_valid:
                        calc_result = CalculationRegistry.calculate(
                            "toxic_dispersion_v1", params
                        )
                        results.append(
                            {
                                "method_id": "toxic_dispersion_v1",
                                "substance": name,
                                "input_params": calc_result.input_params,
                                "results": calc_result.results,
                                "validation_status": "valid",
                            }
                        )
                except Exception as e:  # noqa: BLE001 - расчёт не должен рвать генерацию
                    import logging

                    logging.getLogger(__name__).warning(
                        "Расчёт токсического поражения для '%s' не выполнен: %s: %s",
                        name,
                        type(e).__name__,
                        str(e)[:200],
                    )
                    results.append(
                        {
                            "method_id": "toxic_dispersion_v1",
                            "substance": name,
                            "validation_status": "error",
                            "error": f"{type(e).__name__}: {e}",
                        }
                    )

        return results

    def _get_calc_placeholders(
        self, section_id: str, calculation_results: list[dict]
    ) -> dict:
        """Получение плейсхолдеров расчётных данных для раздела.

        Результаты с ``validation_status == "error"`` (расчёт не выполнен)
        пропускаются: для них нет данных, но они сохранены в логах и общем
        списке результатов для прозрачности.
        """
        placeholders = {}
        for calc in calculation_results:
            if calc.get("validation_status") == "error":
                continue
            method_id = calc["method_id"]
            results = calc.get("results", {})
            substance = calc.get("substance", "")

            for key, value in results.items():
                placeholder_key = f"{method_id}_{substance}_{key}"
                placeholders[f"CALCULATED:{placeholder_key}"] = value

        return placeholders

    async def _get_rag_context(self, section: dict, context: dict) -> str:
        """Получение RAG контекста с фильтрацией по актуальности."""
        if self._retriever is None:
            return ""
        try:
            rag_query = self._build_rag_query(section.get("rag_query", ""), context)
            rag_chunks = await self._retriever.retrieve(rag_query)

            # Фильтрация по актуальности (базовая проверка)
            valid_chunks = []
            for chunk in rag_chunks:
                # В будущем: проверка ссылок в чанке на документы из реестра
                valid_chunks.append(chunk)

            return "\n\n".join(c.content for c in valid_chunks)
        except Exception:
            return ""

    def _build_rag_query(self, query_template: str, context: dict) -> str:
        """Подстановка переменных в rag_query."""
        facility = context.get("facility", {})
        substances = context.get("substances", [])
        substance_names = " ".join(
            s.get("name", "") if isinstance(s, dict) else getattr(s, "name", "")
            for s in substances
        )
        try:
            return query_template.format(
                facility_type=facility.get("facility_type", ""),
                hazard_class=facility.get("hazard_class", ""),
                substance_names=substance_names,
            )
        except KeyError:
            return query_template

    async def _generate_section_llm(
        self,
        section_title: str,
        context: dict,
        rag_context: str,
        slot_type: str = "text",
        sample_context: dict | None = None,
        scenario_list: list[dict] | None = None,
        is_pii_section: bool = False,
    ) -> str:
        """Генерация текста раздела через LLM (v2: без расчётных данных в промпте)."""
        import logging
        logger = logging.getLogger(__name__)

        provider = self._local_llm if is_pii_section else self._external_llm

        if is_pii_section and provider is None:
            # Раздел явно помечен как содержащий персональные данные —
            # падаем, а не тихо уходим на внешний провайдер.
            raise PiiRoutingError(
                f"Раздел «{section_title}» помечен pii=true, но local_llm "
                f"не настроен. Генерация остановлена, чтобы не отправить "
                f"персональные данные во внешний ИИ."
            )

        if provider is None:
            from src.application.services.fallback_texts import get_fallback_text
            fallback = get_fallback_text(section_title)
            if not fallback.startswith("["):
                logger.info("Using fallback text for section '%s'", section_title)
                return fallback
            return self._fallback_section_content(section_title, context)

        facility = context.get("facility", {})
        substances = context.get("substances", [])
        equipment = context.get("equipment", [])
        persons = context.get("responsible_persons", [])

        # Для внешнего провайдера контекст ВСЕГДА чистим от персональных
        # полей (full_name/phone/email), даже если pii=false — это защита
        # от ошибки в structure.json, а не единственный барьер. Для
        # local_llm чистка не нужна: данные и так не покидают контур сети.
        if not is_pii_section:
            facility = strip_pii(facility)
            persons = strip_pii(persons)
            # organization в этой функции не используется напрямую, но если
            # он попадёт в build_section_prompt в будущем — тоже чистим здесь заранее.

        user_prompt = build_section_prompt(
            section_title=section_title,
            facility_data=facility,
            substances=substances,
            equipment=equipment,
            rag_context=rag_context,
            responsible_persons=persons,
            slot_type=slot_type,
            scenario_list=scenario_list,
            sample_context=sample_context,
        )

        try:
            response = await provider.complete(
                messages=[
                    LLMMessage(role="system", content=SYSTEM_PROMPT),
                    LLMMessage(role="user", content=user_prompt),
                ]
            )
            content = response.content.strip()
            if not content:
                logger.warning("LLM returned empty content for section '%s'", section_title)
                from src.application.services.fallback_texts import get_fallback_text
                return get_fallback_text(section_title)

            # Постобработка: удаление маркеров Markdown (**)
            content = content.replace("**", "")

            return content
        except Exception as e:
            logger.error("LLM error for section '%s': %s: %s", section_title, type(e).__name__, str(e)[:200])
            from src.application.services.fallback_texts import get_fallback_text
            fallback = get_fallback_text(section_title)
            if not fallback.startswith("["):
                return fallback
            return self._fallback_section_content(section_title, context, error=str(e))

    def _postprocess_placeholders(self, content: str, context: dict, calc_placeholders: dict) -> str:
        """Замена плейсхолдеров на реальные данные."""
        import re

        facility = context.get("facility", {})
        org = context.get("organization", {})
        persons = context.get("responsible_persons", [])

        # Данные для подстановки
        replacements = {
            "EMERGENCY_PHONE_NUMBER": "+7 (XXX) XXX-XX-XX",
            "FACILITY_NAME": facility.get("name", "объект"),
            "FACILITY_ADDRESS": facility.get("address", "адрес объекта"),
            "ORG_NAME": org.get("name", "организация"),
            "ORG_PHONE": org.get("phone", "+7 (XXX) XXX-XX-XX"),
        }

        # Подставляем телефоны ответственных лиц
        if persons:
            for p in persons:
                name = p.get("full_name", "") if isinstance(p, dict) else getattr(p, "full_name", "")
                phone = p.get("phone", "") if isinstance(p, dict) else getattr(p, "phone", "")
                role = p.get("role", "") if isinstance(p, dict) else getattr(p, "role", "")
                if "safety" in role or "начальник" in role.lower():
                    replacements["EMERGENCY_PHONE_NUMBER"] = phone
                    replacements["SAFETY_MANAGER_PHONE"] = phone
                    replacements["SAFETY_MANAGER_NAME"] = name
                if "director" in role or "директор" in role.lower():
                    replacements["DIRECTOR_PHONE"] = phone
                    replacements["DIRECTOR_NAME"] = name

        # Замена [CALCULATED:KEY] и [DATA:KEY]
        def replace_placeholder(match):
            key = match.group(1)
            return replacements.get(key, match.group(0))

        content = re.sub(r'\[(?:CALCULATED|DATA):([A-Za-z_0-9]+)\]', replace_placeholder, content)

        return content

    def _fallback_section_content(
        self,
        section_title: str,
        context: dict,
        error: str | None = None,
    ) -> str:
        """Fallback-контент при недоступности LLM."""
        facility = context.get("facility", {})
        substances = context.get("substances", [])
        equipment = context.get("equipment", [])

        lines = []
        if error:
            lines.append(f"[Автогенерация недоступна ({error[:100]}). Раздел сформирован по исходным данным.]")
            lines.append("")

        lines.append(f"Объект: {facility.get('name', '—')} ({facility.get('facility_type', '—')})")
        lines.append(f"Класс опасности: {facility.get('hazard_class', '—')}")
        lines.append("")

        if substances:
            lines.append("Опасные вещества:")
            for s in substances:
                name = s.get("name", "—") if isinstance(s, dict) else getattr(s, "name", "—")
                qty = s.get("quantity_kg", "—") if isinstance(s, dict) else getattr(s, "quantity_kg", "—")
                lines.append(f"- {name}: {qty} кг")
            lines.append("")

        if equipment:
            lines.append("Оборудование:")
            for eq in equipment:
                eq_name = eq.get("name", "—") if isinstance(eq, dict) else getattr(eq, "name", "—")
                lines.append(f"- {eq_name}")
            lines.append("")

        lines.append(f"[Раздел «{section_title}» требует доработки экспертом по ПБ.]")
        return "\n".join(lines)

    def _load_structure(self, document_type: str) -> dict:
        structure_path = TEMPLATES_DIR / document_type / "structure.json"
        if not structure_path.exists():
            raise ValueError(f"Шаблон не найден: {document_type}")
        return json.loads(structure_path.read_text(encoding="utf-8"))

    def _render_template(self, env, template_name: str, context: dict) -> str:
        template = env.get_template(template_name)
        return template.render(**context)

    def _setup_document_defaults(self, doc: DocxDocument) -> None:
        """Настраивает страницу и базовый стиль по образцу эталонного ПМЛА (ГОСТ)."""
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(PAGE_MARGINS_CM["top"])
        section.bottom_margin = Cm(PAGE_MARGINS_CM["bottom"])
        section.left_margin = Cm(PAGE_MARGINS_CM["left"])
        section.right_margin = Cm(PAGE_MARGINS_CM["right"])

        normal = doc.styles["Normal"]
        normal.font.name = BODY_FONT_NAME
        normal.font.size = Pt(BODY_FONT_SIZE_PT)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), BODY_FONT_NAME)

        pf = normal.paragraph_format
        pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.0
        pf.space_after = Pt(0)

        # Встроенные стили Heading 1/2 приводим к Times New Roman / чёрному,
        # чтобы Word TOC field собирал заголовки при сохранении визуала эталона.
        configure_heading_styles(doc)

    def _add_heading(self, doc: DocxDocument, text: str, *, level: int, center: bool = True) -> None:
        """Заголовок в стиле эталона: Times New Roman, жирный, без синего цвета Word-стиля.

        Для level >= 1 назначается встроенный стиль ``Heading {level}`` — это
        позволяет Word TOC field (``TOC \\o "1-2"``) собирать заголовки при
        обновлении поля. Визуальное оформление задаётся поверх стиля на уровне
        run, переопределяя стиль. level == 0 стиль не назначается.
        """
        paragraph = doc.add_paragraph()
        if level >= 1:
            try:
                paragraph.style = doc.styles[f"Heading {level}"]
            except KeyError:
                pass
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(sanitize_cyrillic_text(text))
        run.font.name = BODY_FONT_NAME
        run.font.bold = True
        run.font.size = Pt(HEADING_FONT_SIZE_PT if level == 0 else HEADING_FONT_SIZE_PT - 2)
        run.font.color.rgb = RGBColor(0, 0, 0)

    _BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

    def _add_body_paragraph(self, doc: DocxDocument, line: str) -> None:
        """
        Добавляет абзац содержимого, конвертируя markdown-разметку **жирный**
        в реальное жирное форматирование вместо буквальных звёздочек.
        HTML-теги удаляются, не-кириллические символы фильтруются.
        """
        line = strip_html(line)
        line = sanitize_cyrillic_text(line)
        paragraph = doc.add_paragraph()
        pos = 0
        for match in self._BOLD_RE.finditer(line):
            if match.start() > pos:
                paragraph.add_run(line[pos:match.start()])
            bold_run = paragraph.add_run(match.group(1))
            bold_run.font.bold = True
            pos = match.end()
        if pos < len(line):
            paragraph.add_run(line[pos:])

    def _build_docx(
        self, title: str, sections: dict[str, str | list[Block]], metadata: dict
    ) -> bytes:
        """Сборка DOCX с метаданными.

        sections: {section_title: content_str_or_blocks_list}
        Если значение — list[Block], рендерит блоки напрямую.
        Если значение — str, рендерит как раньше (split by newline).
        """
        doc = DocxDocument()
        self._setup_document_defaults(doc)

        context = metadata.get("context", {})

        # Front matter выделяется по section_id из ASSEMBLY_REGISTRY (а не по
        # хардкоженным русским строкам):title_page, approval_sheet,
        # correction_log, toc рендерятся специальными хелперами и не должны
        # попасть в общий цикл секций ниже.
        from src.application.services.pmla_assembly_blocks import (
            get_front_matter_section_ids,
            get_section_title,
        )
        for sid in get_front_matter_section_ids():
            title = get_section_title(sid)
            if title:
                sections.pop(title, None)

        # Титульный лист
        create_title_page(doc, context)

        # Лист согласования — служебный front matter, не review workflow.
        add_approval_sheet(doc, context)

        # Журнал корректировки — статичная DOCX-таблица
        add_correction_journal(doc, context.get("corrections"))

        # Содержание — Word TOC field (заголовок "Содержание" рендерится
        # без стиля Heading, чтобы оглавление не сошлалось само на себя).
        add_toc_placeholder(doc)

        # Основные разделы
        for section_title, content_or_blocks in sections.items():
            self._add_heading(doc, section_title, level=1, center=False)

            if isinstance(content_or_blocks, list) and content_or_blocks:
                # Новый путь: рендер блоков
                self._render_blocks(doc, content_or_blocks)
            elif isinstance(content_or_blocks, str):
                # Старый путь: текст, разбитый по строкам (strip HTML tags first)
                cleaned = strip_html(content_or_blocks)
                for line in cleaned.strip().split("\n"):
                    if line.strip():
                        self._add_body_paragraph(doc, line.strip())
            doc.add_paragraph()

        # Приложения — манифест таблица
        appendices_manifest = context.get("appendices_manifest") or []
        attachments = context.get("attachments_checklist") or []
        if appendices_manifest:
            add_appendices_manifest(doc, appendices_manifest)
        elif attachments:
            from src.infrastructure.export.docx_helpers import add_appendices_section
            add_appendices_section(doc, attachments)

        # Нормативные ссылки
        if metadata.get("calculation_results"):
            self._add_heading(doc, "Расчётные данные", level=1, center=False)
            for calc in metadata["calculation_results"]:
                self._add_body_paragraph(
                    doc, f"- {calc['method_id']}: {calc.get('substance', '—')}"
                )

        # Замечания валидации
        if metadata.get("validation_issues"):
            self._add_heading(doc, "Результаты автоматической валидации", level=1, center=False)
            for issue in metadata["validation_issues"]:
                self._add_body_paragraph(
                    doc, f"- [{issue['severity']}] {issue['section']}: {issue['reason']}"
                )

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _render_blocks(self, doc: DocxDocument, blocks: list[Block]) -> None:
        """Рендерит список блоков в DOCX-элементы."""
        for block in blocks:
            if isinstance(block, HeadingBlock):
                self._add_heading(doc, block.text, level=block.level, center=block.center)
            elif isinstance(block, ParagraphBlock):
                if block.bold:
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
                    run = p.add_run(sanitize_cyrillic_text(block.text))
                    run.font.name = BODY_FONT_NAME
                    run.font.bold = True
                    run.font.size = Pt(BODY_FONT_SIZE_PT)
                else:
                    self._add_body_paragraph(doc, block.text)
            elif isinstance(block, TableBlock):
                self._render_table_block(doc, block)
            elif isinstance(block, ImageBlock):
                self._render_image_block(doc, block)

    def _render_table_block(self, doc: DocxDocument, block: TableBlock) -> None:
        """Рендерит TableBlock как настоящую DOCX-таблицу с Table Grid."""
        # Подпись таблицы (сверху, как в эталоне)
        if block.caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.first_line_indent = Cm(0)
            run = cap.add_run(sanitize_cyrillic_text(block.caption))
            run.font.name = BODY_FONT_NAME
            run.font.bold = True
            run.font.size = Pt(BODY_FONT_SIZE_PT)

        num_cols = len(block.headers)
        num_rows = len(block.rows) + 1  # +1 для заголовка
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заголовки (жирные)
        for i, header in enumerate(block.headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(sanitize_cyrillic_text(header))
            run.font.name = BODY_FONT_NAME
            run.font.bold = True
            run.font.size = Pt(BODY_FONT_SIZE_PT)

        # Данные
        for row_idx, row_data in enumerate(block.rows, 1):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(sanitize_cyrillic_text(str(cell_text)))
                    run.font.name = BODY_FONT_NAME
                    run.font.size = Pt(BODY_FONT_SIZE_PT)

        doc.add_paragraph()  # отступ после таблицы

    def _render_image_block(self, doc: DocxDocument, block: ImageBlock) -> None:
        """Рендерит ImageBlock (заглушка — текстовая ссылка)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(f"[Изображение: {block.path}]")
        run.font.name = BODY_FONT_NAME
        run.font.italic = True
        run.font.size = Pt(BODY_FONT_SIZE_PT)

        if block.caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.first_line_indent = Cm(0)
            run = cap.add_run(block.caption)
            run.font.name = BODY_FONT_NAME
            run.font.size = Pt(BODY_FONT_SIZE_PT - 1)


# ---------------------------------------------------------------------------
# Манифест приложений: сопоставление реестра приложений и attachments_checklist
#
# Module-level helpers, используемые EnhancedDocumentGenerator._enrich_context.
# Ключевые слова для сопоставления канонического приложения (по section_id из
# ASSEMBLY_REGISTRY) с записями attachments_checklist анкеты. Сопоставление
# регистронезависимое, по подстроке в поле name записи checklist.
# ---------------------------------------------------------------------------

_APPENDIX_MATCH_KEYWORDS: dict[str, list[str]] = {
    "appendix_1": ["изучен", "порядок изучен", "обучен"],
    "appendix_2": ["оперативн", "сообщен об инцидент", "форма сообщен"],
    "appendix_3": ["пасф", "состав пасф", "состав формирован"],
    "appendix_4": ["оснащен", "оснащение пасф", "средств защит", "сиз"],
    "appendix_5": ["оповещ", "схема оповещ"],
}


def _checklist_matches(checklist: list, keywords: list[str]) -> bool:
    """Проверяет, есть ли в checklist запись с present=True и совпадением имени."""
    for item in checklist:
        if not isinstance(item, dict):
            continue
        if not item.get("present", False):
            continue
        name = str(item.get("name", "")).lower()
        if any(kw in name for kw in keywords):
            return True
    return False


def _synthesize_appendices_manifest(checklist: list, pasf_documents: list | None = None) -> list[dict]:
    """Синтезирует appendices_manifest из реестра приложений + статуса наличия.

    Возвращает записи вида::
        {"appendix_number": 1, "title": "Приложение 1. ...",
         "filename": "—", "present": bool, "source": "template"|"file"}
    source="template" означает, что приложение сгенерировано шаблоном и
    всегда присутствует в DOCX. source="file" — ожидается внешний файл.

    Нумерация:
    1. Канонические шаблонные приложения
    2. Выбранные документы ПАСФ
    """
    from src.application.services.pmla_assembly_blocks import (
        get_appendix_manifest_entries,
        get_block_def,
    )
    from src.application.services.pmla_assembly_blocks import BlockType

    entries = get_appendix_manifest_entries()
    manifest: list[dict] = []
    appendix_num = 1

    for entry in entries:
        sid = entry["section_id"]
        block_def = get_block_def(sid)

        # Template-generated appendices (Jinja2) are always "сформированы"
        # because they produce content directly in the DOCX.
        is_template = block_def and block_def.template is not None

        if is_template:
            present = True
            source = "template"
        else:
            # File-based appendices: check attachments_checklist
            keywords = _APPENDIX_MATCH_KEYWORDS.get(sid, [])
            present = _checklist_matches(checklist, keywords)
            source = "file"

        manifest.append({
            "appendix_number": appendix_num,
            "title": entry["title"],
            "filename": "—",
            "present": present,
            "source": source,
        })
        appendix_num += 1

    # Add PASF documents as file appendices
    for doc in (pasf_documents or []):
        if not isinstance(doc, dict):
            continue
        doc_type = doc.get("document_type", "other")
        title = doc.get("title") or doc.get("file_name") or f"Документ {doc_type}"
        file_name = doc.get("file_name") or "—"
        manifest.append({
            "appendix_number": appendix_num,
            "title": f"{title} ({doc_type})",
            "filename": file_name,
            "present": True,
            "source": "file",
            "document_type": doc_type,
            "document_number": doc.get("document_number", ""),
            "issued_at": doc.get("issued_at", ""),
            "valid_until": doc.get("valid_until", ""),
            "checksum": doc.get("checksum_sha256", ""),
        })
        appendix_num += 1

    return manifest
