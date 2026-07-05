"""Структурный RAG-пайплайн ПМЛА — парсер → детектор → чанкер → эмбеддинги → ChromaDB.

Заменяет простой чанкер на структурный, использующий границы разделов.
Каждый чанк связан с конкретным разделом документа.
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from uuid import UUID

from docx import Document as DocxDocument

from src.infrastructure.rag.parsers.section_detector import SectionDetector
from src.infrastructure.rag.parsers.chunker import StructuralChunker, StructuralChunk, ChunkingConfig
from src.infrastructure.rag.parsers.models import DetectionReport
from src.infrastructure.rag.pipeline import Embedder, VectorStore, Chunk

logger = logging.getLogger(__name__)

SAMPLES_COLLECTION = "isap_samples"


@dataclass
class PmlaParseResult:
    """Результат парсинга DOCX файла ПМЛА."""
    paragraphs: list[tuple[str, str | None, bool]] = field(default_factory=list)
    table_texts: list[str] = field(default_factory=list)
    report: DetectionReport | None = None
    chunks: list[StructuralChunk] = field(default_factory=list)


def parse_pmla_docx(file_path: str) -> PmlaParseResult:
    """Парсит DOCX файл ПМЛА: извлекает параграфы, таблицы, определяет разделы.

    Args:
        file_path: Путь к DOCX файлу.

    Returns:
        PmlaParseResult с параграфами, таблицами, отчётом детектора и чанками.
    """
    doc = DocxDocument(file_path)

    # Извлекаем текст из таблиц (для title_page и других разделов в таблицах)
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    table_texts.append(text)

    # Извлекаем параграфы (без сдвига индексов!)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else None
        is_bold = False
        if para.runs:
            is_bold = any(r.bold for r in para.runs if r.bold is not None)
        paragraphs.append((text, style, is_bold))

    # Детектируем разделы
    detector = SectionDetector()
    report = detector.detect_with_source(paragraphs, source_path=file_path)

    # Если title_page не найден в параграфах, проверяем таблицы
    if not any(s.section_id == "title_page" for s in report.detected_sections):
        for t_text in table_texts:
            if re.search(r"СОГЛАСОВАНО|УТВЕРЖДАЮ|ПЛАН\s+МЕРОПРИЯТИЙ", t_text, re.IGNORECASE):
                # Добавляем текст таблицы в начало и перезапускаем детекцию
                paragraphs.insert(0, (t_text, None, False))
                report = detector.detect_with_source(paragraphs, source_path=file_path)
                break

    # Чанкуем по структуре
    chunker = StructuralChunker()
    chunks = chunker.chunk(report, paragraphs, source_path=file_path)

    return PmlaParseResult(
        paragraphs=paragraphs,
        table_texts=table_texts,
        report=report,
        chunks=chunks,
    )


class StructuralSampleIndexer:
    """Индексирует образцы ПМЛА в ChromaDB с использованием структурного чанкера.

    Каждый чанк привязан к конкретному разделу документа и содержит
    метаданные о разделе для точного поиска.
    """

    def __init__(self):
        self._embedder = Embedder()

    async def index_sample(
        self,
        sample_id: UUID,
        file_path: str,
        file_type: str,
        facility_type: str | None,
        hazard_class: str | None,
    ) -> int:
        """Парсит DOCX, детектирует разделы, чанкует, эмбеддит, сохраняет в ChromaDB.

        Returns:
            Количество проиндексированных чанков.
        """
        if file_type != "docx":
            logger.warning("Only DOCX supported for structural indexing, got %s", file_type)
            return 0

        try:
            result = parse_pmla_docx(file_path)
        except Exception as e:
            logger.warning("Failed to parse sample %s: %s", sample_id, e)
            return 0

        if not result.chunks:
            logger.warning("No chunks generated for sample %s", sample_id)
            return 0

        # Добавляем метаданные к каждому чанку
        for sc in result.chunks:
            sc.metadata.update({
                "sample_id": str(sample_id),
                "facility_type": (facility_type or "").strip(),
                "hazard_class": (hazard_class or "").strip(),
                "doc_type": "sample",
            })

        # Конвертируем StructuralChunk → Chunk для эмбеддинга
        # Генерируем уникальные ID через source = sample_id:section_id
        pipeline_chunks = []
        for sc in result.chunks:
            # Уникальный source для каждого чанка образца
            unique_source = f"{sample_id}:{sc.section_id}"
            chunk = Chunk(
                content=sc.content,
                source=unique_source,
                chunk_index=sc.chunk_index,
                metadata=sc.metadata,
            )
            pipeline_chunks.append(chunk)

        try:
            embedded = await self._embedder.embed_chunks(pipeline_chunks)
            store = VectorStore(collection_name=SAMPLES_COLLECTION)
            store.add(embedded)
            logger.info(
                "Indexed %d structural chunks for sample %s (%d sections found)",
                len(embedded), sample_id,
                len(result.report.detected_sections) if result.report else 0,
            )
            return len(embedded)
        except Exception as e:
            logger.warning("Failed to index sample %s: %s", sample_id, e)
            return 0

    async def remove_sample(self, sample_id: UUID) -> None:
        """Удаляет все чанки образца из ChromaDB."""
        try:
            store = VectorStore(collection_name=SAMPLES_COLLECTION)
            store.delete(where={"sample_id": str(sample_id)})
            logger.info("Removed sample %s from ChromaDB", sample_id)
        except Exception as e:
            logger.warning("Failed to remove sample %s: %s", sample_id, e)


class StructuralSampleRetriever:
    """Ищет релевантные фрагменты по структурным чанкам с фильтрацией по разделам."""

    def __init__(self):
        self._embedder = Embedder()
        try:
            self._store = VectorStore(collection_name=SAMPLES_COLLECTION)
        except Exception:
            self._store = None

    async def retrieve_by_section(
        self,
        query: str,
        section_id: str | None = None,
        facility_type: str | None = None,
        hazard_class: str | None = None,
        top_k: int = 5,
    ) -> list[Chunk]:
        """RAG-поиск по образцам с опциональной фильтрацией по разделу.

        Args:
            query: Поисковый запрос.
            section_id: Фильтр по ID раздела (например, "section_2").
            facility_type: Фильтр по типу объекта.
            hazard_class: Фильтр по классу опасности.
            top_k: Количество результатов.

        Returns:
            Список релевантных чанков.
        """
        if self._store is None:
            return []

        try:
            query_embedding = await self._embedder.embed_query(query)

            # Строим фильтр
            conditions = []
            if section_id:
                conditions.append({"section_id": section_id})
            if facility_type:
                conditions.append({"facility_type": facility_type.strip()})
            if hazard_class:
                conditions.append({"hazard_class": hazard_class.strip()})

            where_filter = None
            if len(conditions) == 1:
                where_filter = conditions[0]
            elif len(conditions) > 1:
                where_filter = {"$and": conditions}

            results = self._store._collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                where=where_filter,
                include=["documents", "metadatas"],
            )

            chunks = []
            for doc, meta in zip(
                results["documents"][0] if results["documents"] else [],
                results["metadatas"][0] if results["metadatas"] else [],
            ):
                chunks.append(Chunk(
                    content=doc,
                    source=meta.get("source", ""),
                    chunk_index=meta.get("chunk_index", 0),
                    metadata=meta,
                ))
            return chunks
        except Exception as e:
            logger.warning("Structural retrieval failed: %s", e)
            return []

    async def retrieve_for_generation(
        self,
        section_id: str,
        query: str,
        facility_type: str | None = None,
        hazard_class: str | None = None,
        top_k: int = 3,
    ) -> list[Chunk]:
        """Retrieve specifically for PMLA generation — fetches chunks from
        the matching section in sample documents.

        Args:
            section_id: Target section (e.g., "section_2" for scenarios).
            query: Semantic query for RAG.
            facility_type: Filter by facility type.
            hazard_class: Filter by hazard class.
            top_k: Number of chunks to return.

        Returns:
            List of relevant chunks from the target section.
        """
        return await self.retrieve_by_section(
            query=query,
            section_id=section_id,
            facility_type=facility_type,
            hazard_class=hazard_class,
            top_k=top_k,
        )
