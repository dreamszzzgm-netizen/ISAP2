"""
RAG-компоненты (ADR-001).
Реализованы вручную — без фреймворков.
Каждый класс — единственная ответственность.
"""
from __future__ import annotations

import hashlib
from dataclasses import dataclass, field
from pathlib import Path

import chromadb

from src.core.settings import settings


# ---------------------------------------------------------------------------
# Базовые типы
# ---------------------------------------------------------------------------

@dataclass
class Document:
    """Исходный документ (PDF, TXT или текст)."""
    content: str
    source: str          # путь к файлу или URL
    metadata: dict = field(default_factory=dict)


@dataclass
class Chunk:
    """Фрагмент документа после разбивки."""
    content: str
    source: str
    chunk_index: int
    metadata: dict = field(default_factory=dict)

    @property
    def id(self) -> str:
        """Уникальный ID чанка на основе содержимого."""
        return hashlib.md5(f"{self.source}:{self.chunk_index}".encode()).hexdigest()


# ---------------------------------------------------------------------------
# DocumentLoader — загрузка PDF, TXT, DOCX
# ---------------------------------------------------------------------------

class DocumentLoader:
    """Загружает документы разных форматов и возвращает список Document."""

    def load(self, path: str | Path) -> list[Document]:
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._load_pdf(path)
        if suffix == ".txt":
            return self._load_text(path)
        if suffix == ".docx":
            return self._load_docx(path)
        raise ValueError(f"Неподдерживаемый формат: {suffix}")

    def _load_pdf(self, path: Path) -> list[Document]:
        import fitz  # PyMuPDF
        documents = []
        pdf = fitz.open(str(path))
        for page_num, page in enumerate(pdf):
            text = page.get_text().strip()
            if text:
                documents.append(Document(
                    content=text,
                    source=str(path),
                    metadata={"page": page_num + 1, "total_pages": len(pdf), "format": "pdf"},
                ))
        pdf.close()
        return documents

    def _load_text(self, path: Path) -> list[Document]:
        text = path.read_text(encoding="utf-8").strip()
        if not text:
            return []
        return [Document(content=text, source=str(path), metadata={"format": "txt"})]

    def _load_docx(self, path: Path) -> list[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        if not text:
            return []
        return [Document(content=text, source=str(path), metadata={"format": "docx"})]


# ---------------------------------------------------------------------------
# Chunker — разбивка на фрагменты
# ---------------------------------------------------------------------------

class Chunker:
    """
    Разбивает документы на фрагменты фиксированного размера
    с перекрытием (overlap).
    """

    def __init__(
        self,
        chunk_size: int = settings.chunk_size,
        overlap: int = settings.chunk_overlap,
    ):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, documents: list[Document]) -> list[Chunk]:
        chunks = []
        for doc in documents:
            doc_chunks = self._split_text(doc.content)
            for idx, text in enumerate(doc_chunks):
                chunks.append(Chunk(
                    content=text,
                    source=doc.source,
                    chunk_index=idx,
                    metadata=doc.metadata,
                ))
        return chunks

    def _split_text(self, text: str) -> list[str]:
        """Простая разбивка по словам с перекрытием."""
        words = text.split()
        result = []
        start = 0

        while start < len(words):
            end = min(start + self.chunk_size, len(words))
            chunk_text = " ".join(words[start:end])
            result.append(chunk_text)
            if end == len(words):
                break
            start += self.chunk_size - self.overlap

        return result


# ---------------------------------------------------------------------------
# Embedder — создание векторных представлений
# ---------------------------------------------------------------------------

class Embedder:
    """
    Создаёт эмбеддинги для чанков.
    Использует OpenAI или локальную модель через Ollama.
    """

    OPENAI_MODEL = "text-embedding-3-small"
    OLLAMA_EMBED_MODEL = "nomic-embed-text"

    def __init__(self):
        self._provider = settings.llm_provider

    async def embed_chunks(self, chunks: list[Chunk]) -> list[tuple[Chunk, list[float]]]:
        """Возвращает список (chunk, embedding)."""
        texts = [c.content for c in chunks]
        vectors = await self._get_embeddings(texts)
        return list(zip(chunks, vectors))

    async def embed_query(self, query: str) -> list[float]:
        """Эмбеддинг для поискового запроса."""
        vectors = await self._get_embeddings([query])
        return vectors[0]

    async def _get_embeddings(self, texts: list[str]) -> list[list[float]]:
        if self._provider == "openai":
            return await self._openai_embed(texts)
        return await self._ollama_embed(texts)

    async def _openai_embed(self, texts: list[str]) -> list[list[float]]:
        import asyncio
        client = AsyncOpenAI(api_key=settings.openai_api_key)
        # Батчи по 100 текстов
        all_vectors = []
        for i in range(0, len(texts), 100):
            batch = texts[i:i + 100]
            response = await client.embeddings.create(
                model=self.OPENAI_MODEL,
                input=batch,
            )
            all_vectors.extend([item.embedding for item in response.data])
        return all_vectors

    async def _ollama_embed(self, texts: list[str]) -> list[list[float]]:
        import httpx
        vectors = []
        async with httpx.AsyncClient(timeout=60.0) as client:
            for text in texts:
                response = await client.post(
                    f"{settings.ollama_base_url}/api/embeddings",
                    json={"model": self.OLLAMA_EMBED_MODEL, "prompt": text},
                )
                response.raise_for_status()
                vectors.append(response.json()["embedding"])
        return vectors


# ---------------------------------------------------------------------------
# VectorStore — хранение и поиск векторов
# ---------------------------------------------------------------------------

class VectorStore:
    """Обёртка над ChromaDB."""

    def __init__(self, collection_name: str | None = None):
        try:
            self._client = chromadb.HttpClient(
                host=settings.chroma_host,
                port=settings.chroma_port,
            )
            name = collection_name or settings.chroma_collection_name
            self._collection = self._client.get_or_create_collection(
                name=name,
                metadata={"hnsw:space": "cosine"},
            )
            self._available = True
        except Exception:
            self._client = None
            self._collection = None
            self._available = False

    def add(self, embedded_chunks: list[tuple[Chunk, list[float]]]) -> None:
        """Добавить чанки с эмбеддингами в коллекцию."""
        if not self._available:
            return
        ids = [chunk.id for chunk, _ in embedded_chunks]
        documents = [chunk.content for chunk, _ in embedded_chunks]
        embeddings = [vec for _, vec in embedded_chunks]
        metadatas = [
            {**chunk.metadata, "source": chunk.source, "chunk_index": chunk.chunk_index}
            for chunk, _ in embedded_chunks
        ]

        self._collection.upsert(
            ids=ids,
            documents=documents,
            embeddings=embeddings,
            metadatas=metadatas,
        )

    def search(self, query_embedding: list[float], top_k: int = settings.retrieval_top_k) -> list[Chunk]:
        """Поиск ближайших чанков по косинусному сходству."""
        if not self._available:
            return []
        results = self._collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
            include=["documents", "metadatas"],
        )

        chunks = []
        for doc, meta in zip(results["documents"][0], results["metadatas"][0]):
            chunks.append(Chunk(
                content=doc,
                source=meta.get("source", ""),
                chunk_index=meta.get("chunk_index", 0),
                metadata=meta,
            ))
        return chunks

    def delete(self, where: dict | None = None) -> None:
        """Удалить чанки из коллекции по фильтру."""
        if not self._available:
            return
        if where:
            self._collection.delete(where=where)


# ---------------------------------------------------------------------------
# Retriever — высокоуровневый интерфейс
# ---------------------------------------------------------------------------

class Retriever:
    """
    Высокоуровневый интерфейс для RAG-поиска.
    Принимает текстовый запрос, возвращает релевантные чанки.
    """

    def __init__(self, embedder: Embedder, vector_store: VectorStore):
        self._embedder = embedder
        self._vector_store = vector_store

    async def retrieve(self, query: str, top_k: int = settings.retrieval_top_k) -> list[Chunk]:
        query_embedding = await self._embedder.embed_query(query)
        return self._vector_store.search(query_embedding, top_k=top_k)
