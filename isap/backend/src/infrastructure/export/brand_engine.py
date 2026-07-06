"""Движок генерации DOCX для формы «Сведения об ОПО»."""
import io
import os
import subprocess
import tempfile
import uuid
from copy import deepcopy

from docx.oxml import OxmlElement
from docx.shared import Cm

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = BASE_DIR
NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

TABLE_WIDTHS = {
    0: [4990, 4989],
    1: [9384, 595],
    2: [9384, 596],
    3: [9384, 595],
    4: [9242, 737],
    5: [454, 1984, 2126, 1985, 2835, 594],
    6: [737, 5103, 4139],
    7: [5670, 4309],
    8: [4536, 170, 1701, 567, 397, 255, 1361, 369, 369, 392],
}


def _set_col_widths(table, widths_dxa):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_dxa):
            tc = cell._tc
            tcPr = tc.find(f'{NS}tcPr')
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcW = tcPr.find(f'{NS}tcW')
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(f'{NS}w', str(w))
            tcW.set(f'{NS}type', 'dxa')


def _set_cell(cell, val):
    tc = cell._tc if hasattr(cell, '_tc') else cell
    p = tc.find(f'.//{NS}p')
    if p is None:
        return
    for r in list(p.findall(f'{NS}r')):
        p.remove(r)
    r = p.makeelement(f'{NS}r', {})
    t = r.makeelement(f'{NS}t', {})
    t.text = val
    r.append(t)
    p.append(r)
    if val == 'V':
        pPr = p.find(f'{NS}pPr')
        if pPr is None:
            pPr = p.makeelement(f'{NS}pPr', {})
            p.insert(0, pPr)
        jc = pPr.find(f'{NS}jc')
        if jc is None:
            jc = p.makeelement(f'{NS}jc', {})
            pPr.append(jc)
        jc.set(f'{NS}val', 'center')


def _text(el):
    return ''.join(t.text or '' for t in el.findall(f'.//{NS}t'))


def generate_opo_docx(data: dict) -> bytes:
    """Генерация DOCX из данных формы. Возвращает байты."""
    from docxtpl import DocxTemplate

    is_legal = data.get('applicant_type', 'legal') == 'legal'

    applicant_lines = (
        [
            ('8.1.1', data.get('f8_1_1', '')),
            ('8.1.2', data.get('f8_1_2', '')),
            ('8.1.3', data.get('f8_1_3', '')),
            ('8.1.4', data.get('f8_1_4', '')),
            ('8.1.5', data.get('f8_1_5', '')),
            ('8.1.6', data.get('f8_1_6', '')),
            ('8.1.7', data.get('f8_1_7', '')),
            ('8.1.8', data.get('f8_1_8', '')),
            ('8.1.9', data.get('f8_1_9', '')),
            ('М.П.',  data.get('f8_1_10', '')),
        ] if is_legal else [
            ('8.2.1', data.get('f8_2_1', '')),
            ('8.2.2', data.get('f8_2_2', '')),
            ('8.2.3', data.get('f8_2_3', '')),
            ('8.2.4', data.get('f8_2_4', '')),
            ('8.2.5', data.get('f8_2_5', '')),
            ('8.2.6', data.get('f8_2_6', '')),
            ('', ''), ('', ''), ('', ''), ('', ''),
        ]
    )

    ctx = {
        'opo_name':            data.get('f1_1', ''),
        'object_type':         data.get('f1_2', ''),
        'industry_code':       data.get('f1_3', ''),
        'address':             data.get('f1_4', ''),
        'oktmo':               data.get('f1_5', ''),
        'commissioning_date':  data.get('f1_6', ''),
        'owner_name':          data.get('f1_7_1', ''),
        'owner_inn':           data.get('f1_7_2', ''),
        'processes_text':      data.get('processes_text', ''),
        'danger_class':        data.get('danger_class', ''),
        'classification_text': data.get('classification_text', ''),
        'licenses_text':       data.get('licenses_text', ''),
        'composition':         [],
        'total_amount':        data.get('totalAmount', '0'),
        'nearby_substances':   data.get('f7', ''),
        'applicant_label':     '',
        'applicant_lines':     applicant_lines,
        'reg_number':          data.get('f9_1', ''),
        'temp_number':         data.get('f9_2', ''),
        'reg_date':            data.get('f9_3', ''),
        'change_date':         data.get('f9_4', ''),
        'reg_org':             data.get('f9_5', ''),
        'auth_post':           data.get('f9_6', ''),
        'auth_fio':            data.get('f9_7', ''),
        'auth_sign':           data.get('f9_8', ''),
        'auth_sign_date':      data.get('f9_10', ''),
        'auth_mp':             data.get('f9_11', ''),
        'sign_dolj':           data.get('signDolj', ''),
        'sign_podp':           data.get('signPodp', ''),
        'sign_date':           data.get('signDate', ''),
        'sign_mp':             data.get('signMp', ''),
    }

    tpl_path = os.path.join(TEMPLATES_DIR, 'opo_template.docx')
    doc = DocxTemplate(tpl_path)
    doc.render(ctx)

    body = doc.element.body
    elements = list(body)

    for el in elements[:3]:
        text = _text(el)
        if 'Рекомендуемая' in text or ('форм' in text.lower() and 'сведени' in text.lower()):
            body.remove(el)
            break

    for el in body:
        for br in el.findall(f'.//{NS}br'):
            if br.get(f'{NS}type') == 'page':
                br.getparent().remove(br)

    elements = list(body)
    sec8_idx = None
    for i, el in enumerate(elements):
        if '8.' in _text(el) and 'Заявитель' in _text(el):
            sec8_idx = i
            break

    if sec8_idx is not None:
        empty_before = 0
        for i in range(sec8_idx - 1, -1, -1):
            el = elements[i]
            tag = el.tag.replace(NS, 'w:')
            if tag == 'w:p' and not _text(el).strip():
                empty_before += 1
            else:
                break

        target = 5
        sec8_el = elements[sec8_idx]
        if empty_before < target:
            for _ in range(target - empty_before):
                empty_p = OxmlElement('w:p')
                sec8_el.addprevious(empty_p)
        elif empty_before > target:
            count = 0
            for i in range(sec8_idx - 1, -1, -1):
                el = elements[i]
                tag = el.tag.replace(NS, 'w:')
                if tag == 'w:p' and not _text(el).strip():
                    count += 1
                    if count > target:
                        body.remove(el)
                else:
                    break

    for section in doc.sections:
        section.top_margin    = Cm(1.50)
        section.bottom_margin = Cm(1.00)
        section.left_margin   = Cm(2.00)
        section.right_margin  = Cm(1.50)

    for ti, table in enumerate(doc.tables):
        if ti in TABLE_WIDTHS:
            _set_col_widths(table, TABLE_WIDTHS[ti])

    t8 = doc.tables[8]
    for row in t8.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = trPr.makeelement(f'{NS}cantSplit', {})
        cantSplit.set(f'{NS}val', '1')
        trPr.append(cantSplit)

    sign_date = data.get('signDate', '')
    if sign_date:
        parts = sign_date.replace('«','').replace('»','').replace('г.','').strip().split()
        if len(parts) >= 3:
            _set_cell(t8.rows[0].cells[3], '«')
            _set_cell(t8.rows[0].cells[4], parts[0])
            _set_cell(t8.rows[0].cells[5], '»')
            _set_cell(t8.rows[0].cells[6], parts[1])
            _set_cell(t8.rows[0].cells[7], '20' + parts[2] if len(parts[2]) == 2 else parts[2])
            _set_cell(t8.rows[0].cells[9], 'г.')
        elif len(parts) == 2:
            _set_cell(t8.rows[0].cells[3], '«')
            _set_cell(t8.rows[0].cells[4], parts[0])
            _set_cell(t8.rows[0].cells[5], '»')
            _set_cell(t8.rows[0].cells[6], parts[1])
            _set_cell(t8.rows[0].cells[9], 'г.')
    _set_cell(t8.rows[1].cells[8], '')

    dc = data.get('danger_class', '')
    roman = {'I': 0, 'II': 1, 'III': 2, 'IV': 3}
    for name in ['IV', 'III', 'II', 'I']:
        if name in dc:
            _set_cell(doc.tables[2].rows[roman[name]].cells[1], 'V')
            break

    proc_ids = [s.strip() for s in data.get('processes_text', '').split(',') if s.strip()]
    t1 = doc.tables[1]
    proc_row_map = {
        '2.1': 0, '2.2а': 2, '2.2б': 3, '2.2в': 4,
        '2.3': 5, '2.4': 6, '2.5': 7, '2.6': 8
    }
    for pid in proc_ids:
        if pid in proc_row_map:
            _set_cell(t1.rows[proc_row_map[pid]].cells[1], 'V')

    classif_ids = [s.strip() for s in data.get('classification_text', '').split(';') if s.strip()]
    t3 = doc.tables[3]
    classif_row_map = {
        '4.1': 0, '4.2': 1, '4.3': 2, '4.4': 3, '4.5': 4,
        '4.6': 5, '4.7': 6, '4.8': 7, '4.9': 8, '4.10': 9, '4.11': 10,
        'на землях особо охраняемых природных территорий': 11,
        'на континентальном шельфе Российской Федерации': 12,
        'во внутренних морских водах, территориальном море или прилежащей зоне Российской Федерации': 13,
        'на искусственном земельном участке, созданном на водном объекте, находящемся в федеральной собственности': 14,
        '4.12. ОПО, аварии на котором могут иметь трансграничное воздействие': 15,
    }
    for cid in classif_ids:
        if cid in classif_row_map:
            _set_cell(t3.rows[classif_row_map[cid]].cells[1], 'V')
        else:
            num = cid.replace('4.', '').strip()
            try:
                row_idx = int(num) - 1
                if 0 <= row_idx < len(t3.rows):
                    _set_cell(t3.rows[row_idx].cells[1], 'V')
            except (ValueError, IndexError):
                pass

    license_ids = [s.strip() for s in data.get('licenses_text', '').split(';') if s.strip()]
    t4 = doc.tables[4]
    for lid in license_ids:
        num = lid.replace('5.', '').strip()
        try:
            row_idx = int(num) - 1
            if 0 <= row_idx < len(t4.rows):
                _set_cell(t4.rows[row_idx].cells[1], 'V')
        except (ValueError, IndexError):
            pass

    composition = data.get('composition', [])
    if composition:
        t5 = doc.tables[5]
        template_row = deepcopy(t5.rows[2]._tr)
        for i, item in enumerate(reversed(composition)):
            new_row = deepcopy(template_row)
            cells = new_row.findall(f'.//{NS}tc')
            values = [
                str(len(composition) - i),
                str(item.get('name', '')),
                str(item.get('danger', '')),
                str(item.get('substance', '')),
                str(item.get('characteristics', '')),
                str(item.get('processes', '')),
            ]
            for cell, val in zip(cells, values):
                _set_cell(cell, val)
            t5.rows[3]._tr.addprevious(new_row)
        t5._tbl.remove(t5.rows[2]._tr)

    t6 = doc.tables[6]
    if not is_legal:
        _set_cell(t6.rows[0].cells[0], '8.2')
        _set_cell(t6.rows[0].cells[1], 'Индивидуальный предприниматель')
        ip_labels = [
            'Фамилия, имя и отчество (при наличии) индивидуального предпринимателя',
            'Идентификационный номер налогоплательщика (ИНН)',
            'Основной государственный регистрационный номер индивидуального предпринимателя (ОГРНИП)',
            'Адрес места жительства индивидуального предпринимателя',
            'Подпись индивидуального предпринимателя',
            'Дата подписания',
        ]
        for ri in range(1, 7):
            _set_cell(t6.rows[ri].cells[0], f'8.2.{ri}')
            _set_cell(t6.rows[ri].cells[1], ip_labels[ri - 1])
        for ri in range(9, 6, -1):
            t6._tbl.remove(t6.rows[ri]._tr)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_opo_pdf(data: dict) -> bytes:
    """Генерация PDF из данных формы. Возвращает байты."""
    docx_bytes = generate_opo_docx(data)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, f"opo_{uuid.uuid4().hex[:8]}.docx")
        with open(docx_path, 'wb') as f:
            f.write(docx_bytes)

        soffice = _find_soffice()
        subprocess.run([
            soffice, '--headless', '--convert-to', 'pdf',
            '--outdir', tmpdir, docx_path
        ], check=True, timeout=120)

        pdf_name = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        pdf_path = os.path.join(tmpdir, pdf_name)
        with open(pdf_path, 'rb') as f:
            return f.read()


def _find_soffice():
    import platform
    if platform.system() == 'Windows':
        for candidate in [
            os.path.join(os.environ.get('PROGRAMFILES', ''), 'LibreOffice', 'program', 'soffice.exe'),
            os.path.join(os.environ.get('PROGRAMFILES(X86)', ''), 'LibreOffice', 'program', 'soffice.exe'),
        ]:
            if os.path.isfile(candidate):
                return candidate
    # In Docker container, soffice is in PATH
    import shutil
    if shutil.which('soffice'):
        return 'soffice'
    if shutil.which('libreoffice'):
        return 'libreoffice'
    return 'soffice'
