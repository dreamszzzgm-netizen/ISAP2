"""Helper functions for DOCX document generation.

Provides reusable formatting helpers for building professional DOCX documents
without mixing formatting logic into business logic.
"""
from __future__ import annotations

import re
from io import BytesIO

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

BODY_FONT_NAME = "Times New Roman"
BODY_FONT_SIZE_PT = 12
HEADING_FONT_SIZE_PT = 14
FIRST_LINE_INDENT_CM = 1.25

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
HTML_TAG_RE = re.compile(r"<[^>]+>")
# Characters allowed in Russian PMLA text: Cyrillic, Latin, digits, standard punctuation
ALLOWED_CHAR_RE = re.compile(
    r"[^\u0000-\u007F"   # ASCII (Latin, digits, punctuation)
    r"\u0400-\u04FF"     # Cyrillic
    r"\u0500-\u052F"     # Cyrillic Supplement
    r"\u2000-\u206F"     # General punctuation (em dash, etc.)
    r"\u2100-\u214F"     # Letterlike symbols (°, №, etc.)
    r"]+"
)


def strip_html(text: str) -> str:
    """Remove HTML tags from text, leaving only the inner content."""
    return HTML_TAG_RE.sub("", text)


def sanitize_cyrillic_text(text: str) -> str:
    """Remove non-Cyrillic/non-Latin characters (Chinese, Japanese, etc.)

    that may appear in LLM-generated text. Keeps Cyrillic, Latin, digits,
    and standard punctuation. Useful as post-processing for LLM output.
    """
    return ALLOWED_CHAR_RE.sub("", text)


def safe_text(value) -> str:
    """Convert value to safe string, replacing None/null/undefined with 'не указано'."""
    if value is None:
        return "не указано"
    if isinstance(value, bool):
        return "да" if value else "нет"
    if isinstance(value, (list, dict)):
        if not value:
            return "не указано"
        return str(value)
    s = str(value).strip()
    if not s or s in ("None", "null", "undefined", "{}", "[]"):
        return "не указано"
    return s


def set_document_margins(doc: DocxDocument, top: float = 2.0, bottom: float = 2.0, left: float = 3.0, right: float = 1.5) -> None:
    """Set page margins in centimeters."""
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_default_font(doc: DocxDocument, font_name: str = BODY_FONT_NAME, font_size: int = BODY_FONT_SIZE_PT) -> None:
    """Set the default document font."""
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)
    # Set East Asian font
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def configure_heading_styles(doc: DocxDocument) -> None:
    """Переопределяет встроенные стили Heading 1/2 под эталон ПМЛА.

    python-docx создаёт документы на дефолтном шаблоне, где ``Heading 1``/2
    имеют голубой цвет и шрифт Calibri. Чтобы Word-овский TOC field
    (``TOC \\o "1-2"``) собирал заголовки, параграфам назначаются встроенные
    стили — но при этом сами стили приводятся к Times New Roman / чёрному,
    сохраняя визуал эталона (ГОСТ).
    """
    from docx.enum.style import WD_STYLE_TYPE
    for level, size in ((1, HEADING_FONT_SIZE_PT), (2, HEADING_FONT_SIZE_PT - 1)):
        style_name = f"Heading {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:  # дефолтный шаблон всегда содержит Heading 1/2
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = BODY_FONT_NAME
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        # East Asian font binding (иначе Word подставляет свой CJK-шрифт)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), BODY_FONT_NAME)


def add_heading(doc: DocxDocument, text: str, level: int = 1, center: bool = False) -> None:
    """Add a heading with proper formatting.

    Для level >= 1 назначается встроенный стиль ``Heading {level}``, чтобы
    Word TOC field собирал заголовки разделов и подразделов. Визуальное
    оформление (bold, чёрный цвет, Times New Roman) задаётся поверх стиля
    на уровне run и переопределяет стиль. level == 0 (заголовок документа,
    служебные заголовки вне нумерации TOC) стиль не назначается.
    """
    paragraph = doc.add_paragraph()
    if level >= 1:
        try:
            paragraph.style = doc.styles[f"Heading {level}"]
        except KeyError:
            pass  # стиль не настроен — рендеримся как раньше
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(safe_text(text))
    run.font.name = BODY_FONT_NAME
    run.font.bold = True
    run.font.size = Pt(HEADING_FONT_SIZE_PT if level == 0 else max(HEADING_FONT_SIZE_PT - 2 * (level - 1), 12))
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_body_paragraph(doc: DocxDocument, text: str, bold: bool = False) -> None:
    """Add a body paragraph with optional bold formatting."""
    paragraph = doc.add_paragraph()
    if bold:
        paragraph.paragraph_format.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
        run = paragraph.add_run(safe_text(text))
        run.font.name = BODY_FONT_NAME
        run.font.bold = True
        run.font.size = Pt(BODY_FONT_SIZE_PT)
    else:
        # Handle **bold** markdown
        pos = 0
        for match in BOLD_RE.finditer(text):
            if match.start() > pos:
                paragraph.add_run(text[pos:match.start()])
            bold_run = paragraph.add_run(match.group(1))
            bold_run.font.bold = True
            pos = match.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])


def add_kv_table(doc: DocxDocument, rows: list[tuple[str, str]], caption: str | None = None) -> None:
    """Add a key-value table (2 columns: label, value)."""
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        run = cap.add_run(caption)
        run.font.name = BODY_FONT_NAME
        run.font.bold = True
        run.font.size = Pt(BODY_FONT_SIZE_PT)

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = 1  # WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(rows):
        # Label cell (bold)
        cell_label = table.rows[i].cells[0]
        cell_label.text = ""
        p = cell_label.paragraphs[0]
        run = p.add_run(safe_text(label))
        run.font.name = BODY_FONT_NAME
        run.font.bold = True
        run.font.size = Pt(BODY_FONT_SIZE_PT)

        # Value cell
        cell_value = table.rows[i].cells[1]
        cell_value.text = ""
        p = cell_value.paragraphs[0]
        run = p.add_run(safe_text(value))
        run.font.name = BODY_FONT_NAME
        run.font.size = Pt(BODY_FONT_SIZE_PT)

    doc.add_paragraph()


def add_data_table(doc: DocxDocument, headers: list[str], rows: list[list[str]], caption: str | None = None) -> None:
    """Add a data table with headers and rows."""
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        run = cap.add_run(caption)
        run.font.name = BODY_FONT_NAME
        run.font.bold = True
        run.font.size = Pt(BODY_FONT_SIZE_PT)

    num_cols = len(headers)
    num_rows = len(rows) + 1
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = 1  # WD_TABLE_ALIGNMENT.CENTER

    # Headers (bold)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(safe_text(header))
        run.font.name = BODY_FONT_NAME
        run.font.bold = True
        run.font.size = Pt(BODY_FONT_SIZE_PT)

    # Data rows
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(safe_text(cell_text))
                run.font.name = BODY_FONT_NAME
                run.font.size = Pt(BODY_FONT_SIZE_PT)

    doc.add_paragraph()


def _first_text(*values) -> str:
    """Return the first meaningful text value without leaking raw placeholders."""
    for value in values:
        if value is None:
            continue
        if isinstance(value, str):
            cleaned = value.strip()
            if cleaned and cleaned not in ("None", "null", "undefined", "{}", "[]", "—"):
                return cleaned
        elif not isinstance(value, (list, dict)):
            cleaned = str(value).strip()
            if cleaned:
                return cleaned
    return ""


def _person_name(person: dict | str | None) -> str:
    if isinstance(person, str):
        return _first_text(person)
    if not isinstance(person, dict):
        return ""
    return _first_text(
        person.get("full_name"),
        person.get("name"),
        person.get("fio"),
        person.get("display_name"),
    )


def _person_position(person: dict | str | None) -> str:
    if not isinstance(person, dict):
        return ""
    return _first_text(person.get("position"), person.get("title"), person.get("role_label"))


def _role_text(person: dict) -> str:
    role_parts = [
        person.get("role"),
        person.get("position"),
        person.get("title"),
        person.get("role_label"),
    ]
    return " ".join(str(part).lower() for part in role_parts if part)


def _find_person_by_role(persons: list[dict], keywords: tuple[str, ...]) -> dict | None:
    for person in persons:
        if not isinstance(person, dict):
            continue
        role_text = _role_text(person)
        if any(keyword in role_text for keyword in keywords):
            return person
    return None


def _person_from_any(value) -> dict:
    if isinstance(value, dict):
        return {
            "full_name": _person_name(value),
            "position": _person_position(value),
        }
    if isinstance(value, str) and value.strip():
        return {"full_name": value.strip(), "position": ""}
    return {}


def _approval_row(role: str, position: str, name: str) -> list[str]:
    return [
        role,
        position or "__________________",
        name or "__________________",
        "__________",
        "__________",
    ]


def build_approval_rows(context: dict) -> list[list[str]]:
    """Build approval-sheet rows from known PMLA context fields.

    The sheet is informational and does not change the review workflow. Missing
    people are rendered as blank signature fields.
    """
    persons = [
        person for person in context.get("responsible_persons", []) or []
        if isinstance(person, dict)
    ]
    questionnaire = context.get("questionnaire") or {}
    q_persons = questionnaire.get("responsible_persons") or []
    if isinstance(q_persons, list):
        persons.extend(person for person in q_persons if isinstance(person, dict))

    developer = (
        _find_person_by_role(persons, ("developer", "engineer", "industrial_safety", "разработ", "инженер"))
        or persons[0] if persons else {}
    )
    reviewer = (
        _find_person_by_role(persons, ("review", "check", "safety", "провер", "ответствен"))
        or {}
    )

    organization = context.get("organization") or {}
    facility = context.get("facility") or {}
    approver = _person_from_any(context.get("approver"))
    if not _person_name(approver):
        approver = _person_from_any(organization.get("director"))
    if not _person_name(approver):
        approver = _person_from_any(organization.get("manager"))
    if not _person_name(approver):
        approver = _person_from_any(facility.get("responsible_person"))

    return [
        _approval_row(
            "Разработал",
            _person_position(developer) or "Инженер",
            _person_name(developer),
        ),
        _approval_row(
            "Проверил",
            _person_position(reviewer) or "Ответственный специалист",
            _person_name(reviewer),
        ),
        _approval_row(
            "Утвердил",
            _person_position(approver) or "Руководитель организации",
            _person_name(approver),
        ),
    ]


def add_approval_sheet(doc: DocxDocument, context: dict | None = None) -> None:
    """Add the PMLA approval sheet after the title page."""
    add_heading(doc, "Лист согласования", level=1, center=False)
    add_data_table(
        doc,
        ["Роль", "Должность", "ФИО", "Подпись", "Дата"],
        build_approval_rows(context or {}),
    )
    doc.add_page_break()


def create_title_page(doc: DocxDocument, context: dict) -> None:
    """Create a professional title page for the PMLA document."""
    # Organization info
    org = context.get("organization", {})
    facility = context.get("facility", {})

    # Add some spacing at the top
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)

    # Main title
    title_text = "ПЛАН МЕРОПРИЯТИЙ\nпо локализации и ликвидации последствий аварий\nна опасном производственном объекте"
    add_heading(doc, title_text, level=0, center=True)

    # Spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # Info table
    info_rows = [
        ("Наименование организации", safe_text(org.get("name"))),
        ("Наименование ОПО", safe_text(facility.get("name"))),
        ("Регистрационный номер ОПО", safe_text(facility.get("reg_number"))),
        ("Класс опасности", safe_text(facility.get("hazard_class"))),
        ("Адрес ОПО", safe_text(facility.get("address"))),
    ]
    add_kv_table(doc, info_rows)

    # Spacing
    for _ in range(3):
        doc.add_paragraph()

    # Year
    from datetime import datetime
    year = str(datetime.now().year)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f"Год: {year}")
    run.font.name = BODY_FONT_NAME
    run.font.size = Pt(BODY_FONT_SIZE_PT)

    # Page break
    doc.add_page_break()


def _normalize_attachment(item) -> dict:
    """Normalize an attachment item to {name, present} dict.

    Accepts both string items ("схема расположения ОПО") and
    dict items ({"name": "...", "present": true}).
    """
    if isinstance(item, str):
        return {"name": item, "present": True}
    if isinstance(item, dict):
        return {
            "name": item.get("name", ""),
            "present": item.get("present", False),
        }
    return {"name": str(item), "present": True}


def add_appendices_section(doc: DocxDocument, attachments_checklist: list[dict | str] | None = None) -> None:
    """Add the Appendices section with checklist."""
    add_heading(doc, "Приложения", level=1, center=False)

    if not attachments_checklist:
        add_body_paragraph(doc, "Приложения не представлены.")
        return

    appendix_num = 1
    for raw_item in attachments_checklist:
        attachment = _normalize_attachment(raw_item)
        name = attachment.get("name") or f"Приложение {appendix_num}"
        present = attachment.get("present", False)
        status = "" if present else " — не представлено"
        add_body_paragraph(doc, f"Приложение {appendix_num}. {name}{status}")
        appendix_num += 1


# ---------------------------------------------------------------------------
# PMLA Assembly Layer helpers
# ---------------------------------------------------------------------------


def add_correction_journal(doc: DocxDocument, corrections: list[dict] | None = None) -> None:
    """Add the correction journal as a proper DOCX table.

    Columns: № п/п | Дата изменения | Раздел / страница | Содержание изменения | Основание | Подпись
    Default: header row + 1 empty data row.
    """
    add_heading(doc, "Журнал корректировки документа", level=1, center=False)

    headers = ["№ п/п", "Дата изменения", "Раздел / страница", "Содержание изменения", "Основание", "Подпись"]
    rows: list[list[str]] = []

    if corrections:
        for i, c in enumerate(corrections, 1):
            rows.append([
                str(i),
                str(c.get("date", "")),
                str(c.get("section_page", "")),
                str(c.get("description", "")),
                str(c.get("basis", "")),
                str(c.get("signature", "")),
            ])
    else:
        rows.append(["", "", "", "", "", ""])

    add_data_table(doc, headers, rows)
    doc.add_page_break()


def add_toc_placeholder(doc: DocxDocument) -> None:
    """Insert a Word TOC field placeholder.

    After opening in Word, user right-clicks → Update Field to populate page numbers.
    The TOC reads Heading 1 / Heading 2 styles from the document.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Заголовок «Содержание» — без стиля Heading (level=0), иначе оглавление
    # сошлётся само на себя.
    add_heading(doc, "Содержание", level=0, center=False)

    # Create TOC field: { TOC \o "1-2" \h \z \u }
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._element.append(instr_text)

    run3 = paragraph.add_run()
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld_char_separate)

    run4 = paragraph.add_run("[Обновите содержание: правый клик → Обновить поле]")
    run4.font.name = BODY_FONT_NAME
    run4.font.size = Pt(BODY_FONT_SIZE_PT)
    run4.font.italic = True

    run5 = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run5._element.append(fld_char_end)

    doc.add_page_break()


def add_appendices_manifest(doc: DocxDocument, appendices: list[dict] | None = None) -> None:
    """Add the Appendices manifest as a DOCX table.

    Columns: № приложения | Наименование | Файл | Наличие
    Each appendix entry gets a row.
    """
    add_heading(doc, "Приложения", level=1, center=False)

    if not appendices:
        add_body_paragraph(doc, "Приложения не представлены.")
        return

    headers = ["№ приложения", "Наименование", "Файл", "Наличие"]
    rows: list[list[str]] = []

    for entry in appendices:
        num = str(entry.get("appendix_number", ""))
        title = str(entry.get("title", ""))
        filename = str(entry.get("filename", "")) or "—"
        present = entry.get("present", False)
        source = entry.get("source", "file")
        if present and source == "template":
            status = "сформировано"
        elif present:
            status = "представлен"
        else:
            status = "не представлен"
        rows.append([num, title, filename, status])

    add_data_table(doc, headers, rows)
