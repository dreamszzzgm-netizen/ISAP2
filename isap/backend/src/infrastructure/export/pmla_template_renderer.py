"""PMLA Template Renderer — renders DOCX from template via docxtpl.

This is the core rendering engine for the v2 PMLA template. It takes a context
dict with all required fields and produces a rendered DOCX file.

Architecture:
- Uses docxtpl + Jinja2 for template rendering
- Post-processes to remove artifacts (yellow highlights, empty Jinja lines)
- Sets field-update-on-open for TOC/NUMPAGES
"""
from __future__ import annotations

import io
import logging
import os
import re
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Cm

logger = logging.getLogger(__name__)

# backend/src/infrastructure/export/ → 5 levels up = isap/ (project root)
TEMPLATES_DIR = Path(__file__).parent.parent.parent.parent.parent / "files"

# Jinja artifacts to remove from rendered document
JINJA_LINE_RE = re.compile(r"^\s*\{%.*?%\}\s*$", re.DOTALL)
JINJA_TAG_RE = re.compile(r"\{[%{].*?[%}]\}")


class PmlaTemplateRenderer:
    """Renders PMLA DOCX from a Jinja2 template using docxtpl."""

    def __init__(self, template_path: str | Path | None = None):
        self.template_path = Path(template_path) if template_path else TEMPLATES_DIR / "pmla_v2_template.docx"

    def render(self, context: dict) -> bytes:
        """Render the template with the given context. Returns DOCX bytes.
        
        Args:
            context: Dict with all template variables (organization, facility,
                     equipment_list, substance_params, etc.)
        
        Returns:
            Rendered DOCX as bytes.
        """
        from docxtpl import DocxTemplate

        logger.info("Rendering PMLA template: %s", self.template_path)
        doc = DocxTemplate(str(self.template_path))
        doc.render(context)

        # Post-process
        self._remove_jinja_artifacts(doc)
        self._fix_page_number_fields(doc)
        self._replace_toc_with_field(doc)
        self._remove_empty_headings(doc)
        self._remove_empty_tables(doc)
        self._fix_section_margins(doc)
        self._set_field_update_on_open(doc)

        # Save to bytes
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        # Second pass: remove any remaining Jinja lines via XML
        docx_bytes = self._clean_xml_jinja(docx_bytes)

        logger.info("Rendered PMLA: %d bytes", len(docx_bytes))
        return docx_bytes

    def render_to_file(self, context: dict, output_path: str | Path) -> Path:
        """Render and save to a file."""
        output_path = Path(output_path)
        docx_bytes = self.render(context)
        output_path.write_bytes(docx_bytes)
        return output_path

    def _remove_jinja_artifacts(self, doc: Document) -> None:
        """Remove yellow highlight from Jinja placeholders and empty Jinja lines."""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        # Remove paragraphs that are only Jinja control tags
                        if JINJA_LINE_RE.match(text):
                            # Clear the paragraph but keep it (don't remove XML element
                            # to avoid breaking table structure)
                            for run in para.runs:
                                run.text = ""

    def _set_field_update_on_open(self, doc: Document) -> None:
        """Set the document to update fields when opened in Word."""
        from lxml import etree
        # settings.element IS the lxml element in python-docx
        settings_el = doc.settings.element
        # Use lxml find on the element directly
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        update_fields = settings_el.find('.//w:updateFields', nsmap)
        if update_fields is None:
            update_fields = etree.SubElement(settings_el, qn("w:updateFields"))
        update_fields.set(qn("w:val"), "true")

    def _fix_page_number_fields(self, doc: Document) -> None:
        """Add dynamic PAGE fields to footers that lack them."""
        from docx.oxml import OxmlElement
        from lxml import etree

        def _add_page_field(parent_elem):
            """Add PAGE field XML to a paragraph element."""
            r1 = OxmlElement('w:r')
            f1 = OxmlElement('w:fldChar')
            f1.set(qn('w:fldCharType'), 'begin')
            r1.append(f1)
            parent_elem.append(r1)
            r2 = OxmlElement('w:r')
            instr = OxmlElement('w:instrText')
            instr.text = ' PAGE '
            r2.append(instr)
            parent_elem.append(r2)
            r3 = OxmlElement('w:r')
            f3 = OxmlElement('w:fldChar')
            f3.set(qn('w:fldCharType'), 'separate')
            r3.append(f3)
            parent_elem.append(r3)
            r4 = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = '1'
            r4.append(t)
            parent_elem.append(r4)
            r5 = OxmlElement('w:r')
            f5 = OxmlElement('w:fldChar')
            f5.set(qn('w:fldCharType'), 'end')
            r5.append(f5)
            parent_elem.append(r5)

        for i, section in enumerate(doc.sections):
            if section.footer.is_linked_to_previous:
                continue
            # Check if any paragraph in footer already has PAGE field
            has_page = False
            for p in section.footer.paragraphs:
                for run in p.runs:
                    for child in run._element:
                        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                        if tag == 'instrText' and child.text and 'PAGE' in child.text:
                            has_page = True
                            break
                    if has_page:
                        break
                if has_page:
                    break
            if has_page:
                continue
            # Find or create a paragraph in the footer
            footer_paras = list(section.footer.paragraphs)
            if footer_paras:
                para = footer_paras[0]
                # Clear existing runs
                for child in list(para._element):
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'w:r':
                        para._element.remove(child)
            else:
                # Create new paragraph in footer
                para_elem = etree.SubElement(section.footer._element, qn('w:p'))
                # Add paragraph properties with JC (center alignment)
                pPr = etree.SubElement(para_elem, qn('w:pPr'))
                jc = etree.SubElement(pPr, qn('w:jc'))
                jc.set(qn('w:val'), 'center')
                para = doc.add_paragraph()
                # We need to use the XML element, not the proxy object
                # The footer element already has our new p element directly
            # Add PAGE field to the paragraph element
            if footer_paras:
                p_element = footer_paras[0]._element
            else:
                p_element = section.footer._element[-1]
                # Verify it's a w:p element
                tag = p_element.tag.split('}')[-1] if '}' in p_element.tag else p_element.tag
                if tag != 'p':
                    p_element = None
            if p_element is not None:
                _add_page_field(p_element)

    def _replace_toc_with_field(self, doc: Document) -> None:
        """Replace the first empty toc-1 paragraph with a TOC Word field."""
        from docx.oxml import OxmlElement
        for para in doc.paragraphs:
            if not (para.style and getattr(para.style, 'style_id', '') == '13'):
                continue
            # Clear all w:r runs, keep pPr
            for child in list(para._element):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag == 'w:r':
                    para._element.remove(child)
            # Build TOC field
            r1 = OxmlElement('w:r')
            f1 = OxmlElement('w:fldChar')
            f1.set(qn('w:fldCharType'), 'begin')
            r1.append(f1)
            para._element.append(r1)
            r2 = OxmlElement('w:r')
            instr = OxmlElement('w:instrText')
            instr.text = r'TOC \o "1-4" \h \z \u'
            r2.append(instr)
            para._element.append(r2)
            r3 = OxmlElement('w:r')
            f3 = OxmlElement('w:fldChar')
            f3.set(qn('w:fldCharType'), 'separate')
            r3.append(f3)
            para._element.append(r3)
            r4 = OxmlElement('w:r')
            t_elem = OxmlElement('w:t')
            t_elem.text = '[Обновите оглавление]'
            r4.append(t_elem)
            para._element.append(r4)
            r5 = OxmlElement('w:r')
            f5 = OxmlElement('w:fldChar')
            f5.set(qn('w:fldCharType'), 'end')
            r5.append(f5)
            para._element.append(r5)
            break  # only replace first toc-1

    def _remove_empty_headings(self, doc: Document) -> None:
        """Remove Heading 1-4 paragraphs that have no text or only whitespace."""
        headings_to_remove = []
        for para in doc.paragraphs:
            if para.style and getattr(para.style, 'style_id', '') in ('1', '2', '3', '4'):
                if not para.text.strip():
                    headings_to_remove.append(para)
        for para in headings_to_remove:
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)
        if headings_to_remove:
            logger.info("Removed %d empty heading paragraphs", len(headings_to_remove))

    def _remove_empty_tables(self, doc: Document) -> None:
        """Remove tables with no data rows; deduplicate identical empty tables."""
        from lxml import etree
        single_row_tables = []
        empty_by_fingerprint = {}
        for table in doc.tables:
            rows = table.rows
            if len(rows) <= 1:
                single_row_tables.append(table)
                continue
            has_data = False
            for row in rows[1:]:
                if any(cell.text.strip() for cell in row.cells):
                    has_data = True
                    break
            if not has_data:
                fp = etree.tostring(table._element, method='c14n')
                empty_by_fingerprint.setdefault(fp, []).append(table)
        tables_to_remove = list(single_row_tables)
        for fp, table_list in empty_by_fingerprint.items():
            if len(table_list) > 1:
                tables_to_remove.extend(table_list[1:])  # keep first
            else:
                tables_to_remove.append(table_list[0])
        for table in tables_to_remove:
            parent = table._element.getparent()
            if parent is not None:
                parent.remove(table._element)
        if tables_to_remove:
            logger.info("Removed %d empty/duplicate tables", len(tables_to_remove))

    def _fix_section_margins(self, doc: Document) -> None:
        """Apply uniform margins to all portrait sections."""
        for section in doc.sections:
            if section.orientation == WD_ORIENT.PORTRAIT:
                section.top_margin = Cm(2.0)
                section.bottom_margin = Cm(2.0)
                section.left_margin = Cm(3.0)
                section.right_margin = Cm(1.5)

    def _clean_xml_jinja(self, docx_bytes: bytes) -> bytes:
        """Remove any remaining Jinja artifacts from the DOCX XML.
        
        Processes word/document.xml, headers, and footers to remove
        any stray Jinja tags that docxtpl may have left behind.
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            # Extract DOCX (it's a ZIP)
            docx_path = os.path.join(tmpdir, "input.docx")
            with open(docx_path, "wb") as f:
                f.write(docx_bytes)

            with zipfile.ZipFile(docx_path, "r") as zin:
                zin.extractall(tmpdir)

            # Process XML files — only document.xml has Jinja placeholders
            xml_files = [
                "word/document.xml",
            ]
            # Headers and footers contain Word PAGE/TOC fields in complex SDT
            # structures. They are NOT processed to avoid ElementTree corruption
            # of structured document tags (SDT).

            for xml_rel in xml_files:
                xml_path = os.path.join(tmpdir, xml_rel)
                if not os.path.exists(xml_path):
                    continue
                try:
                    tree = ET.parse(xml_path)
                    root = tree.getroot()
                    # Find all text elements and clean Jinja artifacts
                    for t_elem in root.iter(qn("w:t")):
                        if t_elem.text and JINJA_TAG_RE.search(t_elem.text):
                            # Remove Jinja tags from text
                            t_elem.text = JINJA_TAG_RE.sub("", t_elem.text)
                    tree.write(xml_path, xml_declaration=True, encoding="UTF-8")
                except ET.ParseError:
                    logger.warning("Could not parse XML: %s", xml_rel)

            # Repackage as DOCX
            output_buf = io.BytesIO()
            with zipfile.ZipFile(output_buf, "w", zipfile.ZIP_DEFLATED) as zout:
                for root_dir, dirs, files in os.walk(tmpdir):
                    for file in files:
                        file_path = os.path.join(root_dir, file)
                        arcname = os.path.relpath(file_path, tmpdir)
                        zout.write(file_path, arcname)

            return output_buf.getvalue()


def render_pmla_document(context: dict, template_path: str | Path | None = None) -> bytes:
    """Convenience function to render a PMLA document.
    
    Args:
        context: Template context dict
        template_path: Optional custom template path
    
    Returns:
        Rendered DOCX bytes
    """
    renderer = PmlaTemplateRenderer(template_path)
    return renderer.render(context)
