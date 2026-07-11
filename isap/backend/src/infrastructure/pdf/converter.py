"""Конвертация DOCX → PDF.

Поддерживает два бэкенда:
1. LibreOffice headless (Linux/Docker — быстрое, точное, сохраняет форматирование)
2. Fallback: python-docx → текст → fpdf2 (медленнее, без форматирования)
"""
from __future__ import annotations

import io
import logging
import os
import subprocess
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)


def docx_bytes_to_pdf(docx_bytes: bytes) -> bytes:
    """Конвертирует DOCX (bytes) → PDF (bytes)."""
    # Попытка 1: LibreOffice headless
    try:
        return _convert_via_libreoffice(docx_bytes)
    except Exception as e:
        logger.info("LibreOffice недоступен (%s), пробуем fpdf2 fallback", e)

    # Попытка 2: fpdf2 fallback
    return _convert_via_fpdf(docx_bytes)


def _find_soffice() -> str:
    """Find LibreOffice soffice binary on the system."""
    # Common installation paths
    candidates = [
        "soffice",  # PATH (Linux/macOS)
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
    ]
    for candidate in candidates:
        try:
            # Quick existence check first for full paths
            if os.sep in candidate or "/" in candidate:
                if Path(candidate).exists():
                    return candidate
                continue
            result = subprocess.run(
                [candidate, "--version"],
                capture_output=True,
                timeout=30,
            )
            if result.returncode == 0:
                return candidate
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    raise RuntimeError(
        "LibreOffice not found. Install LibreOffice and ensure 'soffice' "
        "is in PATH, or install via: winget install TheDocumentFoundation.LibreOffice"
    )


def _convert_via_libreoffice(docx_bytes: bytes) -> bytes:
    """Конвертация через LibreOffice headless — сохраняет полное форматирование."""
    soffice_path = _find_soffice()
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "input.docx"
        docx_path.write_bytes(docx_bytes)

        result = subprocess.run(
            [
                soffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", tmpdir,
                str(docx_path),
            ],
            capture_output=True,
            timeout=120,
        )

        if result.returncode != 0:
            stderr = result.stderr.decode(errors="replace")
            raise RuntimeError(f"soffice failed (rc={result.returncode}): {stderr}")

        pdf_path = docx_path.with_suffix(".pdf")
        if not pdf_path.exists():
            raise RuntimeError("soffice не создал PDF-файл")

        return pdf_path.read_bytes()


def _convert_via_fpdf(docx_bytes: bytes) -> bytes:
    """Fallback: извлекаем текст из DOCX и создаём PDF через fpdf2."""
    from docx import Document as DocxDocument
    from fpdf import FPDF

    doc = DocxDocument(io.BytesIO(docx_bytes))

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Пытаемся использовать шрифт с кириллицей
    font_path = _find_cyrillic_font()
    if font_path:
        pdf.add_font("DejaVu", "", font_path)
        pdf.add_font("DejaVu", "B", font_path)
        font_name = "DejaVu"
    else:
        font_name = "Helvetica"

    pdf.add_page()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            pdf.ln(4)
            continue

        style = para.style.name if para.style else ""

        if "Heading 1" in style or "title" in style.lower():
            pdf.set_font(font_name, "B", 16)
            pdf.multi_cell(0, 10, text)
            pdf.ln(2)
        elif "Heading 2" in style:
            pdf.set_font(font_name, "B", 13)
            pdf.multi_cell(0, 8, text)
            pdf.ln(1)
        elif "Heading 3" in style:
            pdf.set_font(font_name, "B", 11)
            pdf.multi_cell(0, 7, text)
            pdf.ln(1)
        else:
            pdf.set_font(font_name, "", 10)
            pdf.multi_cell(0, 6, text)
            pdf.ln(1)

    # Обработка таблиц
    for table in doc.tables:
        pdf.ln(4)
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(cells)
            pdf.set_font(font_name, "", 9)
            pdf.multi_cell(0, 5, line)
        pdf.ln(2)

    output = io.BytesIO()
    pdf.output(output)
    return output.getvalue()


def _find_cyrillic_font() -> str | None:
    """Ищет шрифт с поддержкой кириллицы (Windows + Linux)."""
    candidates = [
        # Windows
        r"C:\Windows\Fonts\DejaVuSans.ttf",
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\cour.ttf",
        # Linux
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    return None
