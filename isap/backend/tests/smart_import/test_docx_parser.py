import io

from docx import Document

from src.application.services.smart_import.parser import SmartImportParser
from src.application.services.smart_import.profiles import IMPORT_PROFILES


def _docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Организация: АО Хлебокомбинат")
    doc.add_paragraph("Опасный производственный объект: Сеть газопотребления Хлебозавода №2")
    doc.add_paragraph("Регистрационный номер ОПО: А01-0001-0005")
    doc.add_paragraph("За период эксплуатации аварии и инциденты не зарегистрированы.")
    doc.add_paragraph("Возможные сценарии аварий")
    doc.add_paragraph("- утечка природного газа")
    doc.add_paragraph("- воспламенение газовоздушной смеси")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_import_extracts_pmla_questionnaire_preview_row():
    table = SmartImportParser().parse_bytes("old_pmla.docx", _docx_bytes())
    assert len(table.rows) == 1
    row = table.rows[0]
    assert row["Организация"] == "АО Хлебокомбинат"
    assert "Сеть газопотребления" in row["ОПО"]
    assert row["Регистрационный номер ОПО"] == "А01-0001-0005"
    assert row["Были аварии/инциденты"] == "нет"

    profile = IMPORT_PROFILES["pmla_questionnaire"]
    mapping = profile.map_headers(table.headers)
    normalized = profile.normalize_row(row, mapping)
    errors, _warnings = profile.validate_row(normalized)
    assert errors == []
    assert normalized["organization_name"] == "АО Хлебокомбинат"
    assert normalized["has_incidents"] == "нет"
