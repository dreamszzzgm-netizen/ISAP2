"""E2E test: questionnaire data appears in generated DOCX."""
from __future__ import annotations

import json
import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import UUID, uuid4

import pytest
from docx import Document as DocxDocument
from docx.shared import Cm, Pt

from src.application.engines.base import DocumentContext
from src.application.services.pmla_generation_from_questionnaire_service import (
    PmlaGenerationFromQuestionnaireService,
)
from src.application.services.pmla_generation_context import PmlaGenerationContext


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def extract_docx_text(docx_bytes: bytes) -> str:
    """Extract all text from DOCX bytes (paragraphs + tables)."""
    import io
    doc = DocxDocument(io.BytesIO(docx_bytes))
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Test data
# ---------------------------------------------------------------------------

SAMPLE_QUESTIONNAIRE_DATA = {
    "questionnaire": {
        "incident_history": {
            "has_incidents": False,
            "period": "за период эксплуатации",
            "items": [],
        },
        "selected_scenarios": [
            "утечка опасного вещества",
            "загазованность помещения",
            "отказ автоматики",
        ],
        "custom_scenarios": [
            {
                "name": "Отказ запорной арматуры",
                "title": "Отказ запорной арматуры",
                "description": "Нарушение герметичности или невозможность полного перекрытия подачи газа.",
                "place": "газопровод / ГРУ",
                "source_equipment": "газопровод / ГРУ",
                "equipment": "запорная арматура",
                "hazardous_substance": "природный газ",
                "consequences": "загазованность, пожар, взрыв",
            }
        ],
        "organization_resources": {
            "actual_items": [
                {
                    "name": "Газоанализатор",
                    "type": "control",
                    "quantity": "1",
                    "storage_place": "операторная",
                    "responsible_person": "ответственный за производственный контроль",
                    "purpose": "контроль загазованности",
                },
                {
                    "name": "Огнетушитель порошковый",
                    "type": "firefighting",
                    "quantity": "4",
                    "storage_place": "котельная",
                    "responsible_person": "начальник участка",
                    "purpose": "ликвидация начального возгорания",
                },
            ]
        },
        "notification_scheme": {
            "first_receiver": "оператор котельной",
            "incident_commander": "ответственный руководитель работ",
            "pasf_caller": "дежурный диспетчер",
            "fire_caller": "оператор котельной",
            "medical_caller": "специалист по охране труда",
            "shutdown_responsible": "слесарь-ремонтник",
            "evacuation_responsible": "начальник смены",
            "service_meeting_responsible": "представитель администрации",
        },
        "financial_reserve": {
            "created": True,
            "order_number": "12-ПБ",
            "order_date": "2026-01-15",
            "amount": "500000",
            "responsible": "главный бухгалтер",
        },
        "insurance": {
            "has_contract": True,
            "company": "АО Страховая компания",
            "contract_number": "ГО-123456",
            "valid_until": "2027-01-15",
            "insured_amount": "10000000",
        },
    },
    "organization": {
        "id": str(uuid4()),
        "name": "ООО Газовый сервис",
        "inn": "7701234567",
        "address": "г. Тюмень, ул. Промышленная, 5",
        "phone": "+7(3452)12-34-56",
    },
    "facility": {
        "id": str(uuid4()),
        "name": "Котельная №1",
        "facility_type": "Сеть газопотребления",
        "hazard_class": "3",
        "address": "г. Тюмень, ул. Промышленная, 5",
        "reg_number": "77-1-2-0001-12345678",
        "latitude": 57.15,
        "longitude": 65.53,
    },
    "equipment": [
        {"name": "Котёл газовый", "equipment_type": "Теплогенератор", "serial_number": "KG-001", "manufacture_year": "2020"},
        {"name": "ГРУ", "equipment_type": "Газорегуляторный узел", "serial_number": "GRU-001", "manufacture_year": "2019"},
    ],
    "substances": [
        {"name": "природный газ", "cas_number": "74-82-8", "quantity_kg": 500},
    ],
    "responsible_persons": [
        {"full_name": "Иванов И.И.", "position": "Начальник участка", "phone": "+7(3452)12-34-57"},
    ],
    "emergency_services": [],
    "pasf": None,
    "material_reserve": {
        "fin_reserve_order": "12-ПБ от 2026-01-15",
        "fin_reserve_amount": "500000",
        "insurance_company": "АО Страховая компания",
        "insurance_contract": "ГО-123456",
    },
    "protective_equipment": [
        {"name": "Газоанализатор", "type": "control", "quantity": "1", "location": "операторная"},
        {"name": "Огнетушитель порошковый", "type": "firefighting", "quantity": "4", "location": "котельная"},
    ],
}


# ---------------------------------------------------------------------------
# Unit tests
# ---------------------------------------------------------------------------

class TestAdaptContextForGenerator:
    """Test adapt_context_for_generator maps questionnaire fields correctly."""

    def setup_method(self):
        self.service = PmlaGenerationFromQuestionnaireService(
            document_repo=MagicMock(),
            regulatory_repo=MagicMock(),
            scenario_matrix_repo=MagicMock(),
        )

    def test_incident_history_no_incidents(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        items = ctx.get("accidents_and_incidents", [])
        assert len(items) == 1
        desc = items[0].get("description", "").lower()
        assert "не зарегистрированы" in desc
        assert "опасного производственного объекта" in desc

    def test_custom_scenarios_appear(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        user_scenarios = ctx.get("user_scenarios", [])
        custom_names = [s.get("name", "") for s in user_scenarios if s.get("source") == "custom"]
        assert any("Отказ запорной арматуры" in name for name in custom_names)

    def test_protective_equipment_appears(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        pe = ctx.get("protective_equipment", [])
        names = [item.get("name", "") for item in pe]
        assert "Газоанализатор" in names
        assert any("Огнетушитель" in n for n in names)

    def test_notification_scheme_preserved(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        ns = ctx.get("notification_scheme", {})
        assert ns.get("first_receiver") == "оператор котельной"
        assert ns.get("pasf_caller") == "дежурный диспетчер"

    def test_material_reserve_financial(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        mr = ctx.get("material_reserve", {})
        assert "12-ПБ" in mr.get("fin_reserve_order", "")

    def test_insurance_in_material_reserve(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        mr = ctx.get("material_reserve", {})
        assert mr.get("insurance_company") == "АО Страховая компания"
        assert mr.get("insurance_contract") == "ГО-123456"


class TestDocumentContextFromDict:
    """Test DocumentContext.from_dict maps fields correctly."""

    def test_from_dict_maps_all_fields(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        doc_ctx = DocumentContext.from_dict(ctx)

        assert doc_ctx.facility.get("name") == "Котельная №1"
        assert doc_ctx.organization.get("name") == "ООО Газовый сервис"
        assert len(doc_ctx.equipment) == 2
        assert len(doc_ctx.substances) == 1
        assert len(doc_ctx.protective_equipment) == 2
        assert doc_ctx.notification_scheme.get("first_receiver") == "оператор котельной"
        assert "12-ПБ" in doc_ctx.material_reserve.get("fin_reserve_order", "")

    def test_notification_scheme_is_serialized_for_quality_review(self):
        """All UI notification roles survive context serialization."""
        source = PmlaGenerationContext()
        scheme = {
            "first_receiver": "Оператор",
            "responsible_manager": "Начальник смены",
            "calls_pasf": "Диспетчер ПАСФ",
            "calls_fire": "Диспетчер пожарной охраны",
            "calls_medical": "Диспетчер скорой помощи",
            "stops_equipment": "Машинист",
            "evacuation_responsible": "Начальник участка",
            "meets_services": "Дежурный",
        }
        source.notification_scheme = scheme

        assert source.to_dict()["notification_scheme"] == scheme

    def setup_method(self):
        self.service = PmlaGenerationFromQuestionnaireService(
            document_repo=MagicMock(),
            regulatory_repo=MagicMock(),
            scenario_matrix_repo=MagicMock(),
        )


class TestDocxOutputContainsQuestionnaireData:
    """Test that DOCX output contains questionnaire data via engines."""

    def setup_method(self):
        self.service = PmlaGenerationFromQuestionnaireService(
            document_repo=MagicMock(),
            regulatory_repo=MagicMock(),
            scenario_matrix_repo=MagicMock(),
        )

    def _make_doc_context(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        return DocumentContext.from_dict(ctx)

    @pytest.mark.asyncio
    async def test_scenario_engine_renders_custom_scenario(self):
        from src.application.engines.scenario_engine import ScenarioEngine

        engine = ScenarioEngine()
        doc_ctx = self._make_doc_context()
        # Set scenarios from user_scenarios
        doc_ctx.scenarios = doc_ctx.user_scenarios

        result = await engine.generate("section_2", {"title": "Сценарии"}, doc_ctx)
        text = result.content
        assert "Отказ запорной арматуры" in text, f"Custom scenario missing in section_2: {text[:500]}"

    @pytest.mark.asyncio
    async def test_data_engine_renders_protective_equipment(self):
        from src.application.engines.data_engine import DataEngine

        engine = DataEngine()
        doc_ctx = self._make_doc_context()

        result = await engine.generate("section_4", {"title": "Силы и средства"}, doc_ctx)
        text = result.content
        assert "Газоанализатор" in text, f"Protective equipment missing in section_4: {text[:500]}"
        assert "Огнетушитель" in text, f"Fire extinguisher missing in section_4: {text[:500]}"

    @pytest.mark.asyncio
    async def test_data_engine_renders_notification_scheme(self):
        from src.application.engines.data_engine import DataEngine

        engine = DataEngine()
        doc_ctx = self._make_doc_context()

        result = await engine.generate("section_8", {"title": "Управление, связь"}, doc_ctx)
        text = result.content
        assert "оператор котельной" in text, f"Notification role missing in section_8: {text[:500]}"
        assert "дежурный диспетчер" in text, f"PASF caller missing in section_8: {text[:500]}"

    @pytest.mark.asyncio
    async def test_data_engine_renders_financial_reserve(self):
        from src.application.engines.data_engine import DataEngine

        engine = DataEngine()
        doc_ctx = self._make_doc_context()

        result = await engine.generate("section_13", {"title": "Материально-техническое обеспечение"}, doc_ctx)
        text = result.content
        assert "12-ПБ" in text, f"Financial order missing in section_13: {text[:500]}"
        assert "500000" in text, f"Financial amount missing in section_13: {text[:500]}"

    @pytest.mark.asyncio
    async def test_data_engine_renders_insurance(self):
        from src.application.engines.data_engine import DataEngine

        engine = DataEngine()
        doc_ctx = self._make_doc_context()

        result = await engine.generate("section_13", {"title": "Материально-техническое обеспечение"}, doc_ctx)
        text = result.content
        assert "АО Страховая компания" in text, f"Insurance company missing in section_13: {text[:500]}"
        assert "ГО-123456" in text, f"Insurance contract missing in section_13: {text[:500]}"

    @pytest.mark.asyncio
    async def test_special_section_renders_custom_scenario(self):
        from src.application.engines.scenario_engine import ScenarioEngine

        engine = ScenarioEngine()
        doc_ctx = self._make_doc_context()
        doc_ctx.scenarios = doc_ctx.user_scenarios

        result = await engine.generate("special_section", {"title": "Специальный раздел"}, doc_ctx)
        text = result.content
        assert "Отказ запорной арматуры" in text, f"Custom scenario missing in special_section: {text[:500]}"


class TestValidateContext:
    """Test validate_questionnaire_context warnings."""

    def setup_method(self):
        self.service = PmlaGenerationFromQuestionnaireService(
            document_repo=MagicMock(),
            regulatory_repo=MagicMock(),
            scenario_matrix_repo=MagicMock(),
        )

    def test_full_context_passes(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        result = self.service.validate_questionnaire_context(ctx)
        assert result["passed"] is True
        assert result["summary"]["selected_scenarios"] == 3
        assert result["summary"]["custom_scenarios"] == 1

    def test_missing_facility_fails(self):
        data = dict(SAMPLE_QUESTIONNAIRE_DATA)
        data["facility"] = {}
        ctx = self.service.adapt_context_for_generator(data)
        result = self.service.validate_questionnaire_context(ctx)
        assert result["passed"] is False
        assert any("facility" in e for e in result["errors"])


class TestDocxExtraction:
    """Test DOCX text extraction utility."""

    def test_extract_docx_text(self):
        doc = DocxDocument()
        doc.add_paragraph("Тестовый абзац")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Ячейка 1"
        table.cell(0, 1).text = "Ячейка 2"
        import io
        buf = io.BytesIO()
        doc.save(buf)
        text = extract_docx_text(buf.getvalue())
        assert "Тестовый абзац" in text
        assert "Ячейка 1" in text
        assert "Ячейка 2" in text


class TestFullDocxBuildFromEngines:
    """Build a full DOCX from engine output and verify all questionnaire data appears."""

    def setup_method(self):
        self.service = PmlaGenerationFromQuestionnaireService(
            document_repo=MagicMock(),
            regulatory_repo=MagicMock(),
            scenario_matrix_repo=MagicMock(),
        )

    def _make_doc_context(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        return DocumentContext.from_dict(ctx)

    def _build_docx_text(self, all_blocks: dict) -> str:
        """Simulate DOCX building from engine block output and extract text."""
        import io
        from docx.shared import Pt, Cm
        from src.application.engines.blocks import HeadingBlock, ParagraphBlock, TableBlock

        doc = DocxDocument()
        section_layout = doc.sections[0]
        section_layout.top_margin = Cm(2.0)
        section_layout.bottom_margin = Cm(2.0)
        section_layout.left_margin = Cm(3.0)
        section_layout.right_margin = Cm(1.5)

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        for title, blocks_or_str in all_blocks.items():
            if isinstance(blocks_or_str, str):
                doc.add_paragraph(blocks_or_str)
            elif isinstance(blocks_or_str, list):
                for block in blocks_or_str:
                    if isinstance(block, HeadingBlock):
                        h = doc.add_heading(level=block.level)
                        h.text = block.text
                    elif isinstance(block, ParagraphBlock):
                        p = doc.add_paragraph()
                        run = p.add_run(block.text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                    elif isinstance(block, TableBlock):
                        num_rows = 1 + len(block.rows)
                        num_cols = len(block.headers)
                        t = doc.add_table(rows=num_rows, cols=num_cols)
                        for j, hdr in enumerate(block.headers):
                            t.cell(0, j).text = str(hdr)
                        for i, row in enumerate(block.rows):
                            for j, val in enumerate(row):
                                t.cell(i + 1, j).text = str(val)

        buf = io.BytesIO()
        doc.save(buf)
        return extract_docx_text(buf.getvalue())

    @pytest.mark.asyncio
    async def test_full_pipeline_text_contains_all_questionnaire_data(self):
        """Run DataEngine + ScenarioEngine for all relevant sections and build DOCX."""
        from src.application.engines.data_engine import DataEngine
        from src.application.engines.scenario_engine import ScenarioEngine

        doc_ctx = self._make_doc_context()
        doc_ctx.scenarios = doc_ctx.user_scenarios

        title_map = {
            "section_1": "Характеристика объекта",
            "section_2": "Сценарии аварий",
            "section_3": "Аварийность",
            "section_4": "Силы и средства",
            "section_6": "Состав и дислокация",
            "section_8": "Управление, связь, оповещение",
            "section_13": "Материально-техническое обеспечение",
            "special_section": "Специальный раздел",
        }

        sections = {}
        for engine in [DataEngine(), ScenarioEngine()]:
            ids = {
                DataEngine: ["section_1", "section_3", "section_4", "section_6", "section_8", "section_13"],
                ScenarioEngine: ["section_2", "special_section"],
            }
            for sid in ids[type(engine)]:
                result = await engine.generate(sid, {"title": title_map[sid]}, doc_ctx)
                sections[title_map[sid]] = result.blocks or result.content

        text = self._build_docx_text(sections)

        required = [
            "Отказ запорной арматуры",
            "Газоанализатор",
            "Огнетушитель",
            "оператор котельной",
            "дежурный диспетчер",
            "12-ПБ",
            "АО Страховая компания",
            "ГО-123456",
            "не зарегистрированы",
        ]

        found = []
        missing = []
        for item in required:
            if item in text:
                found.append(item)
            else:
                missing.append(item)

        assert not missing, f"Missing in DOCX: {missing}\n\nDOCX text preview:\n{text[:3000]}"
        assert len(found) == len(required), f"Only {len(found)}/{len(required)} found"


class TestDocxQualityPhrases:
    """Test that DOCX contains official-quality phrases, not raw data dumps."""

    def setup_method(self):
        self.service = PmlaGenerationFromQuestionnaireService(
            document_repo=MagicMock(),
            regulatory_repo=MagicMock(),
            scenario_matrix_repo=MagicMock(),
        )

    def _make_doc_context(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        return DocumentContext.from_dict(ctx)

    def _build_docx_text(self, all_blocks: dict) -> str:
        import io
        from docx.shared import Pt, Cm
        from src.application.engines.blocks import HeadingBlock, ParagraphBlock, TableBlock

        doc = DocxDocument()
        section_layout = doc.sections[0]
        section_layout.top_margin = Cm(2.0)
        section_layout.bottom_margin = Cm(2.0)
        section_layout.left_margin = Cm(3.0)
        section_layout.right_margin = Cm(1.5)
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        for title, blocks_or_str in all_blocks.items():
            if isinstance(blocks_or_str, str):
                doc.add_paragraph(blocks_or_str)
            elif isinstance(blocks_or_str, list):
                for block in blocks_or_str:
                    if isinstance(block, HeadingBlock):
                        h = doc.add_heading(level=block.level)
                        h.text = block.text
                    elif isinstance(block, ParagraphBlock):
                        p = doc.add_paragraph()
                        run = p.add_run(block.text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                    elif isinstance(block, TableBlock):
                        num_rows = 1 + len(block.rows)
                        num_cols = len(block.headers)
                        t = doc.add_table(rows=num_rows, cols=num_cols)
                        for j, hdr in enumerate(block.headers):
                            t.cell(0, j).text = str(hdr)
                        for i, row in enumerate(block.rows):
                            for j, val in enumerate(row):
                                t.cell(i + 1, j).text = str(val)
            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        return extract_docx_text(buf.getvalue())

    @pytest.mark.asyncio
    async def test_incidents_no_incidents_quality_phrase(self):
        """When no incidents — check for official phrase, not raw JSON."""
        from src.application.engines.data_engine import DataEngine

        engine = DataEngine()
        doc_ctx = self._make_doc_context()
        result = await engine.generate("section_3", {"title": "Аварийность"}, doc_ctx)
        text = result.content
        assert "не зарегистрированы" in text.lower(), f"Missing quality phrase in section_3: {text[:500]}"
        # Should NOT contain raw JSON/Python dict
        assert "None" not in text, f"Raw None found in section_3: {text[:500]}"
        assert "{'" not in text, f"Raw dict found in section_3: {text[:500]}"

    @pytest.mark.asyncio
    async def test_custom_scenario_quality_phrase(self):
        """Custom scenario should have structured narrative text."""
        from src.application.engines.scenario_engine import ScenarioEngine

        engine = ScenarioEngine()
        doc_ctx = self._make_doc_context()
        doc_ctx.scenarios = doc_ctx.user_scenarios
        result = await engine.generate("section_2", {"title": "Сценарии"}, doc_ctx)
        text = result.content
        # Quality phrases from the task spec
        assert "Место возможного возникновения аварии" in text, f"Missing scenario place in section_2: {text[:500]}"
        assert "газопровод / ГРУ" in text, f"Missing scenario place value in section_2: {text[:500]}"
        assert "Задействованное оборудование" in text, f"Missing scenario equipment in section_2: {text[:500]}"
        assert "Опасное вещество" in text, f"Missing scenario substance in section_2: {text[:500]}"

    @pytest.mark.asyncio
    async def test_resources_table_columns(self):
        """Resources table should have extended columns with human-readable type."""
        from src.application.engines.data_engine import DataEngine

        engine = DataEngine()
        doc_ctx = self._make_doc_context()
        result = await engine.generate("section_4", {"title": "Силы и средства"}, doc_ctx)
        text = result.content
        # Check for extended table columns
        assert "Тип" in text, f"Missing 'Тип' column in section_4 table: {text[:500]}"
        assert "Ответственное лицо" in text or "Назначение" in text, \
            f"Missing extended columns in section_4: {text[:500]}"
        # No raw Python representations
        assert "{'" not in text, f"Raw dict in section_4: {text[:500]}"

    @pytest.mark.asyncio
    async def test_notification_scheme_russian_phrases(self):
        """Notification scheme should have Russian narrative, not English labels."""
        from src.application.engines.data_engine import DataEngine

        engine = DataEngine()
        doc_ctx = self._make_doc_context()
        result = await engine.generate("section_8", {"title": "Оповещение"}, doc_ctx)
        text = result.content
        # Check for Russian action phrases
        assert "Первое сообщение об аварии принимает" in text, \
            f"Missing Russian notification phrase: {text[:500]}"
        assert "оператор котельной" in text, f"Missing notification role: {text[:500]}"
        # Should NOT contain English labels
        assert "first receiver" not in text.lower(), f"English label found in section_8: {text[:500]}"
        assert "PASF caller" not in text, f"English label found in section_8: {text[:500]}"

    @pytest.mark.asyncio
    async def test_financial_reserve_quality_phrase(self):
        """Financial reserve should have official Russian text."""
        from src.application.engines.data_engine import DataEngine

        engine = DataEngine()
        doc_ctx = self._make_doc_context()
        result = await engine.generate("section_13", {"title": "Матобеспечение"}, doc_ctx)
        text = result.content
        # Check for official phrases
        assert "приказа № 12-ПБ" in text or "приказа №" in text, \
            f"Missing financial reserve order phrase: {text[:500]}"
        assert "Размер финансового резерва" in text or "500000" in text, \
            f"Missing financial reserve amount: {text[:500]}"
        # Should NOT contain English labels
        assert "Financial reserve" not in text, f"English label in section_13: {text[:500]}"

    @pytest.mark.asyncio
    async def test_insurance_quality_phrase(self):
        """Insurance should have official Russian text."""
        from src.application.engines.data_engine import DataEngine

        engine = DataEngine()
        doc_ctx = self._make_doc_context()
        result = await engine.generate("section_13", {"title": "Матобеспечение"}, doc_ctx)
        text = result.content
        # Check for official phrases
        assert "Гражданская ответственность" in text, \
            f"Missing insurance phrase: {text[:500]}"
        assert "АО Страховая компания" in text, f"Missing insurance company: {text[:500]}"
        assert "ГО-123456" in text, f"Missing insurance contract: {text[:500]}"
        # Should NOT contain English labels
        assert "Insurance company" not in text, f"English label in section_13: {text[:500]}"

    @pytest.mark.asyncio
    async def test_no_raw_json_in_docx(self):
        """DOCX should not contain raw JSON/Python dict/list representations."""
        from src.application.engines.data_engine import DataEngine
        from src.application.engines.scenario_engine import ScenarioEngine

        doc_ctx = self._make_doc_context()
        doc_ctx.scenarios = doc_ctx.user_scenarios

        all_blocks = {}
        for engine in [DataEngine(), ScenarioEngine()]:
            for sid in ["section_3", "section_4", "section_8", "section_13"]:
                result = await engine.generate(sid, {"title": sid}, doc_ctx)
                if result.blocks:
                    all_blocks[sid] = result.blocks

        for sid, blocks in all_blocks.items():
            for block in blocks:
                if hasattr(block, "text"):
                    text = block.text
                    assert "{'" not in text, f"Raw dict in {sid}: {text[:200]}"
                    assert "None" not in text, f"Raw None in {sid}: {text[:200]}"
                elif hasattr(block, "rows"):
                    for row in block.rows:
                        for cell in row:
                            assert "{'" not in str(cell), f"Raw dict in {sid} table: {cell[:200]}"


class TestFullEngineRouterPipeline:
    """End-to-end: run ALL engines through EngineRouter, build real DOCX, verify data."""

    def setup_method(self):
        self.service = PmlaGenerationFromQuestionnaireService(
            document_repo=MagicMock(),
            regulatory_repo=MagicMock(),
            scenario_matrix_repo=MagicMock(),
        )

    def _make_doc_context(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        return DocumentContext.from_dict(ctx)

    def _build_real_docx(self, all_sections: dict) -> bytes:
        """Build DOCX using the actual EnhancedDocumentGenerator._build_docx logic."""
        from src.application.engines.blocks import HeadingBlock, ParagraphBlock, TableBlock

        doc = DocxDocument()
        # Setup defaults (same as EnhancedDocumentGenerator._setup_document_defaults)
        section_layout = doc.sections[0]
        section_layout.top_margin = Cm(2.0)
        section_layout.bottom_margin = Cm(2.0)
        section_layout.left_margin = Cm(3.0)
        section_layout.right_margin = Cm(1.5)
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        # Title
        h = doc.add_heading(level=0)
        h.text = "План мероприятий по локализации и ликвидации последствий аварий"

        for section_title, content_or_blocks in all_sections.items():
            h = doc.add_heading(level=1)
            h.text = section_title
            if isinstance(content_or_blocks, list):
                for block in content_or_blocks:
                    if isinstance(block, HeadingBlock):
                        h2 = doc.add_heading(level=block.level)
                        h2.text = block.text
                    elif isinstance(block, ParagraphBlock):
                        p = doc.add_paragraph()
                        run = p.add_run(block.text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                    elif isinstance(block, TableBlock):
                        num_rows = len(block.rows) + 1
                        num_cols = len(block.headers)
                        t = doc.add_table(rows=num_rows, cols=num_cols)
                        for j, hdr in enumerate(block.headers):
                            t.cell(0, j).text = str(hdr)
                        for i, row in enumerate(block.rows):
                            for j, val in enumerate(row):
                                if j < num_cols:
                                    t.cell(i + 1, j).text = str(val)
            elif isinstance(content_or_blocks, str):
                for line in content_or_blocks.strip().split("\n"):
                    if line.strip():
                        p = doc.add_paragraph()
                        run = p.add_run(line.strip())
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
            doc.add_paragraph()

        import io
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @pytest.mark.asyncio
    async def test_engine_router_all_sections_produce_questionnaire_data_in_docx(self):
        """Run EngineRouter.generate_all() and verify all questionnaire data in DOCX."""
        from src.application.engines.router import EngineRouter
        from src.application.services.engine_integration import create_engine_router

        doc_ctx = self._make_doc_context()
        doc_ctx.scenarios = doc_ctx.user_scenarios

        engine_router = create_engine_router()
        all_results = await engine_router.generate_all(doc_ctx)

        # Convert SectionContent to {title: blocks_or_str}
        sections = {}
        for section_id, result in all_results.items():
            if result.blocks:
                sections[result.title] = result.blocks
            else:
                sections[result.title] = result.content

        docx_bytes = self._build_real_docx(sections)
        text = extract_docx_text(docx_bytes)

        # All questionnaire data must survive the full pipeline
        required = [
            "Отказ запорной арматуры",
            "Газоанализатор",
            "Огнетушитель",
            "оператор котельной",
            "дежурный диспетчер",
            "12-ПБ",
            "АО Страховая компания",
            "ГО-123456",
            "не зарегистрированы",
        ]

        found = []
        missing = []
        for item in required:
            if item in text:
                found.append(item)
            else:
                missing.append(item)

        assert not missing, (
            f"Missing in DOCX after full EngineRouter pipeline: {missing}\n\n"
            f"DOCX text preview:\n{text[:3000]}"
        )
        assert len(found) == len(required), f"Only {len(found)}/{len(required)} found"

    @pytest.mark.asyncio
    async def test_engine_router_quality_phrases_in_docx(self):
        """Verify quality phrases appear in the full DOCX pipeline output."""
        from src.application.services.engine_integration import create_engine_router

        doc_ctx = self._make_doc_context()
        doc_ctx.scenarios = doc_ctx.user_scenarios

        engine_router = create_engine_router()
        all_results = await engine_router.generate_all(doc_ctx)

        sections = {}
        for section_id, result in all_results.items():
            if result.blocks:
                sections[result.title] = result.blocks
            else:
                sections[result.title] = result.content

        docx_bytes = self._build_real_docx(sections)
        text = extract_docx_text(docx_bytes)

        # Quality phrases from the task spec
        quality_phrases = [
            "За период эксплуатации опасного производственного объекта аварии и инциденты",
            "Сценарий: Отказ запорной арматуры",
            "Место возможного возникновения аварии",
            "Первое сообщение об аварии принимает оператор котельной",
            "Финансовый резерв",
            "приказа № 12-ПБ",
            "Гражданская ответственность",
            "АО Страховая компания",
            "ГО-123456",
        ]

        found = []
        missing = []
        for phrase in quality_phrases:
            if phrase in text:
                found.append(phrase)
            else:
                missing.append(phrase)

        assert not missing, (
            f"Missing quality phrases in full DOCX: {missing}\n\n"
            f"DOCX text preview:\n{text[:3000]}"
        )
        assert len(found) == len(quality_phrases), f"Only {len(found)}/{len(quality_phrases)} found"

    @pytest.mark.asyncio
    async def test_engine_router_sections_are_non_empty(self):
        """Every section generated by EngineRouter must have non-empty content."""
        from src.application.services.engine_integration import create_engine_router

        doc_ctx = self._make_doc_context()
        doc_ctx.scenarios = doc_ctx.user_scenarios

        engine_router = create_engine_router()
        all_results = await engine_router.generate_all(doc_ctx)

        empty_sections = []
        for section_id, result in all_results.items():
            if not result.blocks and not result.content.strip():
                empty_sections.append(section_id)

        assert not empty_sections, f"Empty sections after engine routing: {empty_sections}"


class TestDocxStructure:
    """Test DOCX document structure matches the required PMLA template."""

    def setup_method(self):
        self.service = PmlaGenerationFromQuestionnaireService(
            document_repo=MagicMock(),
            regulatory_repo=MagicMock(),
            scenario_matrix_repo=MagicMock(),
        )

    def _make_doc_context(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        return DocumentContext.from_dict(ctx)

    def _build_real_docx(self, all_sections: dict, context: dict = None) -> bytes:
        """Build DOCX using the actual EnhancedDocumentGenerator._build_docx logic."""
        from src.application.engines.blocks import HeadingBlock, ParagraphBlock, TableBlock
        from src.infrastructure.export.docx_helpers import create_title_page

        doc = DocxDocument()
        # Setup defaults (same as EnhancedDocumentGenerator._setup_document_defaults)
        section_layout = doc.sections[0]
        section_layout.top_margin = Cm(2.0)
        section_layout.bottom_margin = Cm(2.0)
        section_layout.left_margin = Cm(3.0)
        section_layout.right_margin = Cm(1.5)
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        # Title page
        if context:
            create_title_page(doc, context)

        for section_title, content_or_blocks in all_sections.items():
            h = doc.add_heading(level=1)
            h.text = section_title
            if isinstance(content_or_blocks, list):
                for block in content_or_blocks:
                    if isinstance(block, HeadingBlock):
                        h2 = doc.add_heading(level=block.level)
                        h2.text = block.text
                    elif isinstance(block, ParagraphBlock):
                        p = doc.add_paragraph()
                        run = p.add_run(block.text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                    elif isinstance(block, TableBlock):
                        num_rows = len(block.rows) + 1
                        num_cols = len(block.headers)
                        t = doc.add_table(rows=num_rows, cols=num_cols)
                        for j, hdr in enumerate(block.headers):
                            t.cell(0, j).text = str(hdr)
                        for i, row in enumerate(block.rows):
                            for j, val in enumerate(row):
                                if j < num_cols:
                                    t.cell(i + 1, j).text = str(val)
            elif isinstance(content_or_blocks, str):
                for line in content_or_blocks.strip().split("\n"):
                    if line.strip():
                        p = doc.add_paragraph()
                        run = p.add_run(line.strip())
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
            doc.add_paragraph()

        import io
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @pytest.mark.asyncio
    async def test_docx_contains_title_page(self):
        """DOCX should contain a proper title page with organization info."""
        from src.infrastructure.export.docx_helpers import create_title_page

        context = {
            "organization": {"name": "ООО Газовый сервис"},
            "facility": {"name": "Котельная №1", "reg_number": "77-1-2-0001-12345678"},
        }

        doc = DocxDocument()
        create_title_page(doc, context)

        import io
        buf = io.BytesIO()
        doc.save(buf)
        text = extract_docx_text(buf.getvalue())

        assert "ПЛАН МЕРОПРИЯТИЙ" in text, "Title page missing main title"
        assert "по локализации и ликвидации последствий аварий" in text, "Title page missing subtitle"
        assert "на опасном производственном объекте" in text, "Title page missing OPO reference"
        assert "ООО Газовый сервис" in text, "Title page missing organization name"
        assert "Котельная №1" in text, "Title page missing facility name"
        assert "77-1-2-0001-12345678" in text, "Title page missing reg number"

    @pytest.mark.asyncio
    async def test_docx_has_all_required_sections(self):
        """DOCX should contain all required PMLA sections."""
        from src.application.services.engine_integration import create_engine_router

        doc_ctx = self._make_doc_context()
        doc_ctx.scenarios = doc_ctx.user_scenarios

        engine_router = create_engine_router()
        all_results = await engine_router.generate_all(doc_ctx)

        sections = {}
        for section_id, result in all_results.items():
            if result.blocks:
                sections[result.title] = result.blocks
            else:
                sections[result.title] = result.content

        context = {
            "organization": SAMPLE_QUESTIONNAIRE_DATA.get("organization", {}),
            "facility": SAMPLE_QUESTIONNAIRE_DATA.get("facility", {}),
        }
        docx_bytes = self._build_real_docx(sections, context)
        text = extract_docx_text(docx_bytes)

        # Check for required sections (use partial matches)
        required_sections = [
            "ПЛАН МЕРОПРИЯТИЙ",
            "Характеристика",
            "Сценарии",
            "Характеристика аварийности",
            "сил и средств",
            "управления, связи",
            "материально-технического обеспечения",
        ]

        found = []
        missing = []
        for section in required_sections:
            if section in text:
                found.append(section)
            else:
                missing.append(section)

        assert not missing, f"Missing sections in DOCX: {missing}"
        assert len(found) == len(required_sections), f"Only {len(found)}/{len(required_sections)} sections found"

    @pytest.mark.asyncio
    async def test_docx_no_none_null_in_output(self):
        """DOCX should not contain None, null, undefined, or empty dict/list representations."""
        from src.application.services.engine_integration import create_engine_router

        doc_ctx = self._make_doc_context()
        doc_ctx.scenarios = doc_ctx.user_scenarios

        engine_router = create_engine_router()
        all_results = await engine_router.generate_all(doc_ctx)

        sections = {}
        for section_id, result in all_results.items():
            if result.blocks:
                sections[result.title] = result.blocks
            else:
                sections[result.title] = result.content

        context = {
            "organization": SAMPLE_QUESTIONNAIRE_DATA.get("organization", {}),
            "facility": SAMPLE_QUESTIONNAIRE_DATA.get("facility", {}),
        }
        docx_bytes = self._build_real_docx(sections, context)
        text = extract_docx_text(docx_bytes)

        # Check for raw data representations
        bad_patterns = ["None", "null", "undefined", "{'", "[{", "{}", "[]"]
        found_bad = []
        for pattern in bad_patterns:
            if pattern in text:
                # Check context to see if it's in a table cell or paragraph
                found_bad.append(pattern)

        # Allow "None" in specific contexts (like table headers)
        # But overall, the document should be clean
        assert not found_bad, f"Raw data representations found in DOCX: {found_bad}"

    @pytest.mark.asyncio
    async def test_docx_tables_have_proper_formatting(self):
        """Tables in DOCX should have headers, rows, and proper structure."""
        from src.application.engines.blocks import TableBlock
        from src.application.engines.data_engine import DataEngine

        engine = DataEngine()
        doc_ctx = self._make_doc_context()

        result = await engine.generate("section_4", {"title": "Силы и средства"}, doc_ctx)

        # Check that the engine produces table blocks
        assert result.blocks, "Section 4 should produce blocks"
        table_blocks = [b for b in result.blocks if isinstance(b, TableBlock)]
        assert table_blocks, "Section 4 should have at least one table"

        # Check table structure
        for table_block in table_blocks:
            assert table_block.headers, "Table should have headers"
            assert table_block.rows, "Table should have rows"
            assert len(table_block.headers) > 0, "Table should have at least one header"
            assert len(table_block.rows) > 0, "Table should have at least one data row"
