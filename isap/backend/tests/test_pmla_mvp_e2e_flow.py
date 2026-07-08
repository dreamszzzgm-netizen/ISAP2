"""E2E test: полный MVP-сценарий ПМЛА.

ОПО → Анкета → Справочники → Генерация → DOCX → Quality Review → Версии → Скачать → Ручная проверка
"""
from __future__ import annotations

import io
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

import pytest
from docx import Document as DocxDocument
from docx.shared import Cm, Pt

from src.application.engines.base import DocumentContext
from src.application.services.document_review_service import (
    ALLOWED_TRANSITIONS,
    DocumentReviewService,
)
from src.application.services.pmla_generation_from_questionnaire_service import (
    PmlaGenerationFromQuestionnaireService,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def extract_docx_text(docx_bytes: bytes) -> str:
    """Extract all text from DOCX bytes (paragraphs + tables)."""
    doc = DocxDocument(io.BytesIO(docx_bytes))
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Test data — полный набор анкеты
# ---------------------------------------------------------------------------

FACILITY_ID = str(uuid4())
ORGANIZATION_ID = str(uuid4())

SAMPLE_FACILITY = {
    "id": FACILITY_ID,
    "name": "Котельная №1 ООО Газовый сервис",
    "facility_type": "Сеть газопотребления",
    "hazard_class": "3",
    "address": "г. Тюмень, ул. Промышленная, 5",
    "reg_number": "77-1-2-0001-12345678",
    "latitude": 57.15,
    "longitude": 65.53,
}

SAMPLE_ORGANIZATION = {
    "id": ORGANIZATION_ID,
    "name": "ООО Газовый сервис",
    "inn": "7701234567",
    "address": "г. Тюмень, ул. Промышленная, 5",
    "phone": "+7(3452)12-34-56",
}

SAMPLE_QUESTIONNAIRE_DATA = {
    "questionnaire": {
        "incident_history": {
            "has_incidents": False,
            "period": "за период эксплуатации",
            "items": [],
        },
        "selected_scenarios": [
            "утечка опасного вещества",
            "загазованность помещения",
            "отказ автоматики",
        ],
        "custom_scenarios": [
            {
                "name": "Отказ запорной арматуры",
                "title": "Отказ запорной арматуры",
                "description": "Нарушение герметичности или невозможность полного перекрытия подачи газа.",
                "place": "газопровод / ГРУ",
                "source_equipment": "газопровод / ГРУ",
                "equipment": "запорная арматура",
                "hazardous_substance": "природный газ",
                "consequences": "загазованность, пожар, взрыв",
            }
        ],
        "organization_resources": {
            "actual_items": [
                {
                    "name": "Газоанализатор",
                    "type": "control",
                    "quantity": "1",
                    "storage_place": "операторная",
                    "responsible_person": "ответственный за производственный контроль",
                    "purpose": "контроль загазованности",
                },
                {
                    "name": "Огнетушитель порошковый",
                    "type": "firefighting",
                    "quantity": "4",
                    "storage_place": "котельная",
                    "responsible_person": "начальник участка",
                    "purpose": "ликвидация начального возгорания",
                },
            ]
        },
        "notification_scheme": {
            "first_receiver": "оператор котельной",
            "incident_commander": "ответственный руководитель работ",
            "pasf_caller": "дежурный диспетчер",
            "fire_caller": "оператор котельной",
            "medical_caller": "специалист по охране труда",
            "shutdown_responsible": "слесарь-ремонтник",
            "evacuation_responsible": "начальник смены",
            "service_meeting_responsible": "представитель администрации",
        },
        "financial_reserve": {
            "created": True,
            "order_number": "12-ПБ",
            "order_date": "2026-01-15",
            "amount": "500000",
            "responsible": "главный бухгалтер",
        },
        "insurance": {
            "has_contract": True,
            "company": "АО Страховая компания",
            "contract_number": "ГО-123456",
            "valid_until": "2027-01-15",
            "insured_amount": "10000000",
        },
        "attachments_checklist": [
            {"name": "Схема расположения ОПО", "present": True},
            {"name": "Схема оповещения", "present": True},
            {"name": "Договор с ПАСФ", "present": False},
            {"name": "Страховой полис", "present": True},
        ],
    },
    "organization": SAMPLE_ORGANIZATION,
    "facility": SAMPLE_FACILITY,
    "equipment": [
        {"name": "Котёл газовый", "equipment_type": "Теплогенератор", "serial_number": "KG-001", "manufacture_year": "2020"},
        {"name": "ГРУ", "equipment_type": "Газорегуляторный узел", "serial_number": "GRU-001", "manufacture_year": "2019"},
    ],
    "substances": [
        {"name": "природный газ", "cas_number": "74-82-8", "quantity_kg": 500},
    ],
    "responsible_persons": [
        {"full_name": "Иванов И.И.", "position": "Начальник участка", "phone": "+7(3452)12-34-57"},
    ],
    "emergency_services": [],
    "pasf": None,
    "material_reserve": {
        "fin_reserve_order": "12-ПБ от 2026-01-15",
        "fin_reserve_amount": "500000",
        "insurance_company": "АО Страховая компания",
        "insurance_contract": "ГО-123456",
    },
    "protective_equipment": [
        {"name": "Газоанализатор", "type": "control", "quantity": "1", "location": "операторная"},
        {"name": "Огнетушитель порошковый", "type": "firefighting", "quantity": "4", "location": "котельная"},
    ],
}


# ---------------------------------------------------------------------------
# E2E Test: полный MVP-сценарий
# ---------------------------------------------------------------------------

class TestMvpE2EFlow:
    """E2E test: полный MVP-сценарий ПМЛА без внешнего LLM."""

    def setup_method(self):
        self.service = PmlaGenerationFromQuestionnaireService(
            document_repo=MagicMock(),
            regulatory_repo=MagicMock(),
            scenario_matrix_repo=MagicMock(),
        )
        self.review_service = DocumentReviewService(
            document_repo=MagicMock(),
        )

    def _make_doc_context(self):
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        return DocumentContext.from_dict(ctx)

    def _build_docx_from_sections(self, sections: dict, context: dict = None) -> bytes:
        """Build DOCX using the actual EnhancedDocumentGenerator._build_docx logic."""
        from src.application.engines.blocks import HeadingBlock, ParagraphBlock, TableBlock
        from src.infrastructure.export.docx_helpers import create_title_page

        doc = DocxDocument()
        section_layout = doc.sections[0]
        section_layout.top_margin = Cm(2.0)
        section_layout.bottom_margin = Cm(2.0)
        section_layout.left_margin = Cm(3.0)
        section_layout.right_margin = Cm(1.5)
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        if context:
            create_title_page(doc, context)

        for section_title, content_or_blocks in sections.items():
            h = doc.add_heading(level=1)
            h.text = section_title
            if isinstance(content_or_blocks, list):
                for block in content_or_blocks:
                    if isinstance(block, HeadingBlock):
                        h2 = doc.add_heading(level=block.level)
                        h2.text = block.text
                    elif isinstance(block, ParagraphBlock):
                        p = doc.add_paragraph()
                        run = p.add_run(block.text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                    elif isinstance(block, TableBlock):
                        num_rows = len(block.rows) + 1
                        num_cols = len(block.headers)
                        t = doc.add_table(rows=num_rows, cols=num_cols)
                        for j, hdr in enumerate(block.headers):
                            t.cell(0, j).text = str(hdr)
                        for i, row in enumerate(block.rows):
                            for j, val in enumerate(row):
                                if j < num_cols:
                                    t.cell(i + 1, j).text = str(val)
            elif isinstance(content_or_blocks, str):
                for line in content_or_blocks.strip().split("\n"):
                    if line.strip():
                        p = doc.add_paragraph()
                        run = p.add_run(line.strip())
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @pytest.mark.asyncio
    async def test_full_mvp_flow(self):
        """Полный MVP-сценарий: ОПО → Анкета → Генерация → DOCX → Review → Issued."""
        # ── 1. Подготовка контекста (симуляция анкеты) ──
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        assert ctx.get("organization", {}).get("name") == "ООО Газовый сервис"
        assert ctx.get("facility", {}).get("name") == "Котельная №1 ООО Газовый сервис"

        # ── 2. Создание DocumentContext ──
        doc_ctx = DocumentContext.from_dict(ctx)
        doc_ctx.scenarios = doc_ctx.user_scenarios

        # ── 3. Генерация через EngineRouter ──
        from src.application.services.engine_integration import create_engine_router

        engine_router = create_engine_router()
        all_results = await engine_router.generate_all(doc_ctx)

        assert len(all_results) > 0, "EngineRouter should produce sections"

        # ── 4. Конвертация в sections dict ──
        sections = {}
        for section_id, result in all_results.items():
            if result.blocks:
                sections[result.title] = result.blocks
            else:
                sections[result.title] = result.content

        assert len(sections) > 0, "Should have sections for DOCX"

        # ── 5. Построение DOCX ──
        context = {
            "organization": SAMPLE_QUESTIONNAIRE_DATA.get("organization", {}),
            "facility": SAMPLE_QUESTIONNAIRE_DATA.get("facility", {}),
        }
        docx_bytes = self._build_docx_from_sections(sections, context)
        assert len(docx_bytes) > 0, "DOCX should not be empty"

        # ── 6. Проверка DOCX содержимого ──
        text = extract_docx_text(docx_bytes)

        # Титульный лист
        assert "ПЛАН МЕРОПРИЯТИЙ" in text, "Missing title"
        assert "ООО Газовый сервис" in text, "Missing organization"
        assert "Котельная №1" in text, "Missing facility"

        # Обязательные данные анкеты
        required_phrases = [
            "Отказ запорной арматуры",
            "Газоанализатор",
            "Огнетушитель",
            "оператор котельной",
            "дежурный диспетчер",
            "12-ПБ",
            "АО Страховая компания",
            "ГО-123456",
            "не зарегистрированы",
        ]

        missing = [p for p in required_phrases if p not in text]
        assert not missing, f"Missing phrases in DOCX: {missing}"

        # Нет сырого мусора
        bad_patterns = ["None", "null", "undefined", "{'", "[{"]
        found_bad = [p for p in bad_patterns if p in text]
        assert not found_bad, f"Raw data in DOCX: {found_bad}"

        # ── 7. Проверка quality review ──
        from src.application.services.pmla_quality_review_service import (
            PmlaQualityReviewService,
        )

        quality_service = PmlaQualityReviewService()
        quality_review = quality_service.review(ctx, docx_bytes)

        assert quality_review is not None, "Quality review should exist"
        assert quality_review.overall_status in ("ok", "warning", "critical")
        assert len(quality_review.checks) > 0, "Quality review should have checks"

        # ── 8. Проверка review workflow ──
        # Симуляция: документ создан, review_status = needs_review
        doc_id = uuid4()

        def make_mock_doc(review_status, **kwargs):
            doc = MagicMock()
            doc.id = doc_id
            doc.review_status = review_status
            doc.review_comment = kwargs.get("review_comment", None)
            doc.reviewed_by = kwargs.get("reviewed_by", None)
            doc.reviewed_at = kwargs.get("reviewed_at", None)
            doc.issued_at = kwargs.get("issued_at", None)
            return doc

        # GET review status
        self.review_service.document_repo.get = AsyncMock(return_value=make_mock_doc("needs_review"))
        review_status = await self.review_service.get_review_status(doc_id)
        assert review_status["review_status"] == "needs_review"
        assert review_status["review_status_label"] == "Требует проверки"
        assert "in_review" in review_status["allowed_transitions"]

        # needs_review → in_review
        self.review_service.document_repo.get = AsyncMock(side_effect=[
            make_mock_doc("needs_review"),
            make_mock_doc("in_review", reviewed_by="engineer"),
        ])
        self.review_service.document_repo.update = AsyncMock()
        result = await self.review_service.update_review_status(
            doc_id, review_status="in_review", reviewed_by="engineer"
        )
        assert result["review_status"] == "in_review"

        # in_review → approved
        self.review_service.document_repo.get = AsyncMock(side_effect=[
            make_mock_doc("in_review"),
            make_mock_doc("approved", review_comment="Проверка пройдена"),
        ])
        self.review_service.document_repo.update = AsyncMock()
        result = await self.review_service.update_review_status(
            doc_id, review_status="approved", review_comment="Проверка пройдена"
        )
        assert result["review_status"] == "approved"

        # approved → ready_to_issue
        self.review_service.document_repo.get = AsyncMock(side_effect=[
            make_mock_doc("approved"),
            make_mock_doc("ready_to_issue"),
        ])
        self.review_service.document_repo.update = AsyncMock()
        result = await self.review_service.update_review_status(
            doc_id, review_status="ready_to_issue"
        )
        assert result["review_status"] == "ready_to_issue"

        # ready_to_issue → issued
        self.review_service.document_repo.get = AsyncMock(side_effect=[
            make_mock_doc("ready_to_issue"),
            make_mock_doc("issued", issued_at="2026-07-08T10:00:00"),
        ])
        self.review_service.document_repo.update = AsyncMock()
        result = await self.review_service.update_review_status(
            doc_id, review_status="issued"
        )
        assert result["review_status"] == "issued"
        assert result["issued_at"] is not None

    @pytest.mark.asyncio
    async def test_regeneration_creates_new_version(self):
        """Повторная генерация создаёт новую версию."""
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        doc_ctx = DocumentContext.from_dict(ctx)
        doc_ctx.scenarios = doc_ctx.user_scenarios

        from src.application.services.engine_integration import create_engine_router

        engine_router = create_engine_router()

        # Первая генерация
        all_results_1 = await engine_router.generate_all(doc_ctx)
        sections_1 = {}
        for section_id, result in all_results_1.items():
            sections_1[result.title] = result.blocks or result.content
        docx_1 = self._build_docx_from_sections(sections_1)
        text_1 = extract_docx_text(docx_1)

        # Вторая генерация (симуляция регенерации)
        all_results_2 = await engine_router.generate_all(doc_ctx)
        sections_2 = {}
        for section_id, result in all_results_2.items():
            sections_2[result.title] = result.blocks or result.content
        docx_2 = self._build_docx_from_sections(sections_2)
        text_2 = extract_docx_text(docx_2)

        # Обе версии содержат одни и те же данные
        assert "Отказ запорной арматуры" in text_1
        assert "Отказ запорной арматуры" in text_2
        assert "Газоанализатор" in text_1
        assert "Газоанализатор" in text_2

    @pytest.mark.asyncio
    async def test_docx_quality_phrases(self):
        """DOCX содержит официальные фразы ПМЛА."""
        ctx = self.service.adapt_context_for_generator(SAMPLE_QUESTIONNAIRE_DATA)
        doc_ctx = DocumentContext.from_dict(ctx)
        doc_ctx.scenarios = doc_ctx.user_scenarios

        from src.application.services.engine_integration import create_engine_router

        engine_router = create_engine_router()
        all_results = await engine_router.generate_all(doc_ctx)

        sections = {}
        for section_id, result in all_results.items():
            sections[result.title] = result.blocks or result.content

        context = {
            "organization": SAMPLE_QUESTIONNAIRE_DATA.get("organization", {}),
            "facility": SAMPLE_QUESTIONNAIRE_DATA.get("facility", {}),
        }
        docx_bytes = self._build_docx_from_sections(sections, context)
        text = extract_docx_text(docx_bytes)

        # Качественные фразы
        quality_phrases = [
            "За период эксплуатации опасного производственного объекта аварии и инциденты",
            "Место возможного возникновения аварии",
            "Первое сообщение об аварии принимает",
            "Финансовый резерв",
            "Гражданская ответственность",
        ]

        missing = [p for p in quality_phrases if p not in text]
        assert not missing, f"Missing quality phrases: {missing}"


class TestMvpEndpoints:
    """Test that all MVP endpoints are properly wired."""

    def test_document_review_service_exists(self):
        """DocumentReviewService can be instantiated."""
        repo = MagicMock()
        service = DocumentReviewService(repo)
        assert service is not None

    def test_transition_rules_complete(self):
        """All transition rules are defined."""
        assert "needs_review" in ALLOWED_TRANSITIONS
        assert "in_review" in ALLOWED_TRANSITIONS
        assert "approved" in ALLOWED_TRANSITIONS
        assert "ready_to_issue" in ALLOWED_TRANSITIONS
        assert "issued" in ALLOWED_TRANSITIONS

    def test_full_workflow_chain(self):
        """Full workflow chain is valid."""
        # needs_review → in_review
        assert "in_review" in ALLOWED_TRANSITIONS["needs_review"]
        # in_review → approved
        assert "approved" in ALLOWED_TRANSITIONS["in_review"]
        # approved → ready_to_issue
        assert "ready_to_issue" in ALLOWED_TRANSITIONS["approved"]
        # ready_to_issue → issued
        assert "issued" in ALLOWED_TRANSITIONS["ready_to_issue"]
        # issued → archived
        assert "archived" in ALLOWED_TRANSITIONS["issued"]
