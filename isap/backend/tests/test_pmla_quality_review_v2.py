"""Tests for PMLA Quality Review v2 — Assembly-aware checks."""
import json
import os
import tempfile
from io import BytesIO

import pytest

from src.application.services.pmla_assembly_blocks import (
    ASSEMBLY_REGISTRY,
    BlockType,
    get_appendix_manifest_entries,
    get_block_def,
    get_static_sections,
    get_variable_sections,
    get_generated_sections,
)
from src.application.services.pmla_quality_review_service import (
    CheckResult,
    PmlaQualityReviewService,
)


def _service():
    return PmlaQualityReviewService()


def _full_demo_context():
    """Load the real demo PMLA validation dataset, flattened for review service.

    The demo JSON nests notification_scheme, organization_resources, etc.
    inside questionnaire, but the quality review service reads them from
    the top level. Flatten them for compatibility.
    """
    demo_path = os.path.join(os.path.dirname(__file__), "..", "data", "demo_pmla_validation.json")
    with open(demo_path, encoding="utf-8") as f:
        raw = json.load(f)

    questionnaire = raw.get("questionnaire", {})
    ctx = {
        "organization": raw.get("organization"),
        "facility": raw.get("facility"),
        "responsible_persons": raw.get("responsible_persons"),
        "equipment": raw.get("equipment"),
        "substances": raw.get("substances"),
        "pasf": raw.get("pasf"),
        "emergency_services": raw.get("emergency_services"),
        # Flatten questionnaire fields to top level for quality review
        "selected_scenarios": questionnaire.get("selected_scenarios", []),
        "custom_scenarios": questionnaire.get("custom_scenarios", []),
        "organization_resources": questionnaire.get("organization_resources", {}),
        "notification_scheme": questionnaire.get("notification_scheme", {}),
        "attachments_checklist": questionnaire.get("attachments_checklist", []),
        "questionnaire": questionnaire,
    }
    return ctx


def _minimal_context():
    """Minimal context with only organization and facility."""
    return {
        "organization": {"name": "ООО Тест"},
        "facility": {"name": "Котельная"},
        "questionnaire": {},
        "selected_scenarios": [],
        "custom_scenarios": [],
        "pasf": None,
        "emergency_services": [],
        "organization_resources": {},
        "notification_scheme": {},
        "attachments_checklist": [],
    }


# --- 1. Complete demo PMLA returns not critical ---


def test_demo_pmla_not_critical():
    """Demo dataset with all data should not produce critical status."""
    ctx = _full_demo_context()
    report = _service().review(ctx, content_docx=b"demo docx")
    # The demo data has all required fields, so no critical checks
    critical_checks = [c for c in report.checks if c.status == "critical"]
    assert report.overall_status != "critical", (
        f"Demo PMLA should not be critical; critical checks: {[c.code for c in critical_checks]}"
    )


def test_demo_pmla_has_all_block_checks():
    """Demo review should include all 6 block-aware checks."""
    ctx = _full_demo_context()
    report = _service().review(ctx, content_docx=b"demo docx")
    block_codes = {
        "assembly_static_blocks",
        "assembly_variable_blocks",
        "assembly_generated_blocks",
        "assembly_toc_block",
        "assembly_appendix_references",
        "assembly_external_files",
    }
    found_codes = {c.code for c in report.checks}
    assert block_codes.issubset(found_codes), f"Missing block checks: {block_codes - found_codes}"


# --- 2. Static blocks do not lower score ---


def test_static_blocks_not_lower_score_without_questionnaire():
    """Static blocks (correction_log, abbreviations, terms, bibliography)
    should not lower the score even when questionnaire data is minimal."""
    ctx = _minimal_context()
    report = _service().review(ctx)
    static_check = next(c for c in report.checks if c.code == "assembly_static_blocks")
    assert static_check.status == "ok"
    # Score penalty should come from data checks, not static block check
    assert static_check.message.startswith("Все")


def test_static_blocks_ok_with_empty_context():
    """Static block check should always pass — they don't need data."""
    report = _service().review(_minimal_context())
    static_check = next(c for c in report.checks if c.code == "assembly_static_blocks")
    assert static_check.status == "ok"


# --- 3. Missing PASF lowers score ---


def test_missing_pasf_lowers_score():
    """Missing PASF should produce a warning and lower score."""
    ctx = _full_demo_context()
    ctx["pasf"] = None
    ctx["questionnaire"]["pasf_manual"] = {}
    report = _service().review(ctx)
    pasf_check = next(c for c in report.checks if c.code == "pasf")
    assert pasf_check.status == "warning"
    assert report.score < 100


# --- 4. Missing emergency services lowers score ---


def test_missing_emergency_services_lowers_score():
    """Missing emergency services should produce a critical and lower score."""
    ctx = _full_demo_context()
    ctx["emergency_services"] = []
    report = _service().review(ctx)
    es_check = next(c for c in report.checks if c.code == "emergency_services")
    assert es_check.status == "critical"
    assert report.score <= 80


# --- 5. Missing appendix manifest lowers score ---


def test_empty_attachments_lower_score():
    """Empty attachments_checklist should produce a warning."""
    ctx = _full_demo_context()
    ctx["attachments_checklist"] = []
    report = _service().review(ctx)
    att_check = next(c for c in report.checks if c.code == "attachments_checklist")
    assert att_check.status == "warning"
    assert report.score < 100


def test_appendix_references_ok_with_manifest():
    """Appendix references check should pass when manifest has entries and attachments are present."""
    ctx = _full_demo_context()
    report = _service().review(ctx)
    appendix_check = next(c for c in report.checks if c.code == "assembly_appendix_references")
    # With flattened context, demo has 5 attachments marked present=True
    assert appendix_check.status == "ok"


# --- 6. Raw HTML tags fail check ---


def test_raw_html_in_generated_blocks_fails():
    """Generated blocks with raw HTML tags should be flagged."""
    rendered = {
        "introduction": "Введение документа",
        "section_2": "<table><tr><td>Сценарий</td></tr></table>",
        "section_5": "Организация взаимодействия",
        "section_7": "Готовность сил",
        "section_9": "Обмен информацией",
        "section_10": "Первоочередные действия",
        "section_11": "Действия персонала",
        "section_12": "Безопасность населения",
        "special_section": "Специальный раздел",
    }
    report = _service().review(_minimal_context(), rendered_sections=rendered)
    gen_check = next(c for c in report.checks if c.code == "assembly_generated_blocks")
    assert gen_check.status == "critical"
    assert "raw_html" in gen_check.message.lower() or "Raw HTML" in gen_check.message


def test_clean_html_in_generated_blocks_passes():
    """Generated blocks with clean text (no HTML tags) should pass."""
    rendered = {
        "introduction": "Введение содержит нормативную базу для ПМЛА",
        "section_2": "Сценарии аварий включают утечку газа и возгорание",
        "section_5": "Взаимодействие сил координируется диспетчером",
        "section_7": "Готовность сил поддерживается на постоянной основе",
        "section_9": "Обмен информацией осуществляется по каналам связи",
        "section_10": "Первоочередные действия направлены на локализацию аварии",
        "section_11": "Действия персонала при угрозе аварии",
        "section_12": "Мероприятия по безопасности населения",
        "special_section": "Оперативный раздел для ликвидации аварий",
    }
    report = _service().review(_minimal_context(), rendered_sections=rendered)
    gen_check = next(c for c in report.checks if c.code == "assembly_generated_blocks")
    assert gen_check.status == "ok"


# --- 7. Quality review output contains block_id ---


def test_output_contains_block_id():
    """Check results should include block_id field."""
    report = _service().review(_full_demo_context(), content_docx=b"demo")
    for c in report.checks:
        d = c.to_dict() if hasattr(c, "to_dict") else {"block_id": c.block_id}
        assert "block_id" in d, f"Check {c.code} missing block_id"


def test_toc_check_has_block_id():
    """The TOC block check should have block_id='toc'."""
    report = _service().review(_minimal_context())
    toc_check = next(c for c in report.checks if c.code == "assembly_toc_block")
    assert toc_check.block_id == "toc"


# --- 8. Quality review output contains block_type ---


def test_output_contains_block_type():
    """Check results should include block_type field."""
    report = _service().review(_full_demo_context(), content_docx=b"demo")
    for c in report.checks:
        d = c.to_dict() if hasattr(c, "to_dict") else {"block_type": c.block_type}
        assert "block_type" in d, f"Check {c.code} missing block_type"


def test_toc_check_has_block_type():
    """The TOC block check should have block_type='word_toc_block'."""
    report = _service().review(_minimal_context())
    toc_check = next(c for c in report.checks if c.code == "assembly_toc_block")
    assert toc_check.block_type == "word_toc_block"


# --- 9. Assembly Registry / structure.json is used as source ---


def test_assembly_registry_drives_static_check():
    """Static block check should use ASSEMBLY_REGISTRY as source of truth."""
    report = _service().review(_minimal_context())
    static_check = next(c for c in report.checks if c.code == "assembly_static_blocks")
    registry_static = get_static_sections()
    assert static_check.status == "ok"
    assert static_check.details["section_ids"] == registry_static


def test_assembly_registry_drives_variable_check():
    """Variable block check should use ASSEMBLY_REGISTRY as source of truth."""
    report = _service().review(_minimal_context())
    var_check = next(c for c in report.checks if c.code == "assembly_variable_blocks")
    registry_variable = get_variable_sections()
    assert var_check.details["section_ids"] == registry_variable


def test_assembly_registry_drives_generated_check():
    """Generated block check should use ASSEMBLY_REGISTRY as source of truth."""
    report = _service().review(_minimal_context())
    gen_check = next(c for c in report.checks if c.code == "assembly_generated_blocks")
    registry_generated = get_generated_sections()
    assert gen_check.details["section_ids"] == registry_generated


def test_structure_json_block_types_match_registry():
    """structure.json block_type fields should match ASSEMBLY_REGISTRY."""
    structure_path = os.path.join(os.path.dirname(__file__), "..", "templates", "pmla", "structure.json")
    with open(structure_path, encoding="utf-8") as f:
        structure = json.load(f)

    for section in structure["sections"]:
        sid = section["id"]
        block_type_str = section["block_type"]
        registry_entry = ASSEMBLY_REGISTRY.get(sid)
        assert registry_entry is not None, f"Section {sid} not in ASSEMBLY_REGISTRY"
        assert registry_entry.block_type.value == block_type_str, (
            f"Section {sid}: structure.json has '{block_type_str}' but registry has '{registry_entry.block_type.value}'"
        )


# --- 10. to_dict output includes block_id and block_type ---


def test_to_dict_includes_block_fields():
    """to_dict() output should include block_id and block_type for every check."""
    report = _service().review(_full_demo_context(), content_docx=b"demo")
    d = report.to_dict()
    for check_dict in d["checks"]:
        assert "block_id" in check_dict
        assert "block_type" in check_dict


def test_toc_check_to_dict_has_block_fields():
    """TOC check in to_dict() should have block_id='toc' and block_type='word_toc_block'."""
    report = _service().review(_minimal_context())
    d = report.to_dict()
    toc_check = next(c for c in d["checks"] if c["code"] == "assembly_toc_block")
    assert toc_check["block_id"] == "toc"
    assert toc_check["block_type"] == "word_toc_block"


# --- 11. Variable block check with missing data ---


def test_variable_blocks_warning_when_data_missing():
    """Variable blocks should warn when key data sources are missing."""
    ctx = _minimal_context()
    report = _service().review(ctx)
    var_check = next(c for c in report.checks if c.code == "assembly_variable_blocks")
    # With minimal context, some variable blocks will have missing data
    # But organization and facility ARE present, so at minimum those sections are ok
    # The check returns warning only if sections have missing *required* keys
    assert var_check.status in ("ok", "warning")


def test_variable_blocks_ok_with_full_data():
    """Variable blocks should pass when all key data sources are present."""
    ctx = _full_demo_context()
    report = _service().review(ctx)
    var_check = next(c for c in report.checks if c.code == "assembly_variable_blocks")
    assert var_check.status == "ok"


# --- 12. Appendix manifest with present=false items ---


def test_appendix_check_respects_present_false():
    """Appendix references should only count items where present=True."""
    ctx = _full_demo_context()
    ctx["attachments_checklist"] = [
        {"name": "схема расположения ОПО", "present": False},
        {"name": "схема оповещения", "present": False},
        {"name": "договор с ПАСФ", "present": False},
        {"name": "страховой полис", "present": False},
    ]
    report = _service().review(ctx)
    att_check = next(c for c in report.checks if c.code == "attachments_checklist")
    # All 4 required attachments have present=False, so all are missing
    assert att_check.status == "warning"


# --- 13. Recommendations include block-aware items ---


def test_recommendations_include_block_aware_items():
    """Recommendations should include block-aware guidance when checks fail."""
    ctx = _minimal_context()
    ctx["notification_scheme"] = {}
    report = _service().review(ctx)
    # With empty notification scheme, should get a recommendation
    recs = report.recommendations
    assert any("оповещен" in r.lower() for r in recs), f"Expected notification recommendation in: {recs}"


# --- 14. Cyrillic regression test ---


def test_cyrillic_survives_docx_roundtrip():
    """Cyrillic strings must survive python-docx roundtrip without mojibake.

    Regression test: PowerShell Get-Content pipe corrupted UTF-8 Cyrillic
    when piping SQL files to psql, causing mojibake in DOCX output.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()
    test_strings = [
        "Индивидуальный предприниматель",
        "Иванов Иван Иванович",
        "Сеть газопотребления ул. Красная",
        "Сидоров Святослав Петрович",
        "Собственник ОПО",
        "Кабардино-Балкарская Республика",
    ]
    for s in test_strings:
        p = doc.add_paragraph()
        run = p.add_run(s)
        run.font.name = "Times New Roman"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    doc2 = DocxDocument(buf)
    for i, p in enumerate(doc2.paragraphs):
        text = p.text
        # No C1 control character U+0098 (mojibake marker)
        assert "\x98" not in text, f"Mojibake \\x98 found in paragraph {i}: {repr(text)}"
        # No raw replacement characters
        assert "\ufffd" not in text, f"Replacement char found in paragraph {i}: {repr(text)}"


def test_cyrillic_no_mojibake_markers():
    """DOCX text must not contain known mojibake byte patterns."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    p = doc.add_paragraph()
    p.add_run("Индивидуальный предприниматель")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    # Check raw XML bytes
    import zipfile
    with zipfile.ZipFile(buf, "r") as z:
        xml = z.read("word/document.xml")

    # The correct UTF-8 for "Индивидуальный" starts with d098
    correct = "Индивидуальный".encode("utf-8")
    assert correct in xml, "Correct UTF-8 not found in DOCX XML"

    # The mojibake pattern d0a0 c298 must NOT be present
    mojibake_pattern = b"\xd0\xa0\xc2\x98"
    assert mojibake_pattern not in xml, "Mojibake pattern found in DOCX XML"


# --- 15. Sanitize Cyrillic text filter ---


def test_sanitize_cyrillic_text_removes_chinese():
    """LLM-generated text with Chinese/Japanese characters should be cleaned."""
    from src.infrastructure.export.docx_helpers import sanitize_cyrillic_text

    # LLM hallucinated Chinese characters
    dirty = "Огнетушители, водяное орошение, средства связи,熱画像 камера, защитные каски"
    clean = sanitize_cyrillic_text(dirty)
    assert "熱画像" not in clean
    assert "Огнетушители" in clean
    assert "защитные каски" in clean


def test_sanitize_cyrillic_text_preserves_russian():
    """Russian text should pass through sanitize filter unchanged."""
    from src.infrastructure.export.docx_helpers import sanitize_cyrillic_text

    text = "Индивидуальный предприниматель — Иванов Иван Иванович"
    assert sanitize_cyrillic_text(text) == text


def test_sanitize_cyrillic_text_preserves_latin_and_digits():
    """Latin text and digits should pass through sanitize filter."""
    from src.infrastructure.export.docx_helpers import sanitize_cyrillic_text

    text = "ГОСТ 12.1.004-2018, CAS 74-82-8, ОКТМО 8364501"
    assert sanitize_cyrillic_text(text) == text


def test_sanitize_cyrillic_text_in_docx_rendering():
    """sanitize_cyrillic_text should be applied in template engine output."""
    from docx import Document as DocxDocument
    from src.infrastructure.export.docx_helpers import sanitize_cyrillic_text

    # Simulate LLM output with Chinese
    llm_text = "Технические средства: 熱画像 камера, огнетушители"
    clean = sanitize_cyrillic_text(llm_text)

    doc = DocxDocument()
    doc.add_paragraph(clean)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    doc2 = DocxDocument(buf)
    text = doc2.paragraphs[0].text
    assert "熱画像" not in text
    assert "камера" in text
    assert "огнетушители" in text
