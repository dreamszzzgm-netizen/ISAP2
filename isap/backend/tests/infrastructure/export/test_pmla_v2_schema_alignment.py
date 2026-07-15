"""PMLA v2 schema–template–renderer alignment tests.

Validates that:
- All Jinja template variables exist in schema or approved computed list
- All template lists are described in schema
- All renderer keys exist in schema
- All required schema fields are available in renderer
- Array item fields match between template and schema
- No name conflicts (scenarios/accident_scenarios, material_resources/material_reserve, etc.)
- After render, no Jinja tags remain
- Empty context doesn't break document
- DOCX is valid ZIP
- All XML files parse without errors

All tests are marked ``audit_only`` because the schema is not yet finalized.
"""
from __future__ import annotations

import io
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
_PROJECT_ROOT = Path(__file__).resolve().parents[4]  # backend/tests/infrastructure/export → isap/isap
_FILES_DIR = _PROJECT_ROOT / "files"
# Fallback for Docker container: /files/ is mounted from host ./files/
_CONTAINER_FILES_DIR = Path("/files")
_SCHEMA_PATH = _FILES_DIR / "pmla_v2.schema.json"
if not _SCHEMA_PATH.exists():
    _SCHEMA_PATH = _CONTAINER_FILES_DIR / "pmla_v2.schema.json"
_TEMPLATE_KEYS_PATH = _FILES_DIR / "pmla_v2_template_keys.json"
if not _TEMPLATE_KEYS_PATH.exists():
    _TEMPLATE_KEYS_PATH = _CONTAINER_FILES_DIR / "pmla_v2_template_keys.json"
_CONTEXT_KEYS_PATH = _FILES_DIR / "pmla_v2_context_keys.json"
if not _CONTEXT_KEYS_PATH.exists():
    _CONTEXT_KEYS_PATH = _CONTAINER_FILES_DIR / "pmla_v2_context_keys.json"
_TEMPLATE_PATH = _FILES_DIR / "pmla_v2_template.docx"
if not _TEMPLATE_PATH.exists():
    _TEMPLATE_PATH = _CONTAINER_FILES_DIR / "pmla_v2_template.docx"

# ---------------------------------------------------------------------------
# Load audit artefacts (shared across all tests)
# ---------------------------------------------------------------------------

def _load_json(path: Path) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


@pytest.fixture(scope="module")
def schema() -> dict:
    return _load_json(_SCHEMA_PATH)


@pytest.fixture(scope="module")
def template_keys() -> dict:
    return _load_json(_TEMPLATE_KEYS_PATH)


@pytest.fixture(scope="module")
def context_keys() -> dict:
    return _load_json(_CONTEXT_KEYS_PATH)


# ---------------------------------------------------------------------------
# Approved computed keys — scalar variables present in the template but NOT
# in schema.required_fields because they are computed or derived at runtime.
# This list must be kept in sync with the renderer / export service.
# ---------------------------------------------------------------------------
APPROVED_COMPUTED_SCALARS: set[str] = {
    # Additional scalar variables found in the template by the audit
    # that are NOT in schema.required_fields but ARE provided by the
    # export service at render time.
    "settlement_name",
    "gas_supplier_branch",
    "dislocation_address",
    "settlement_district",
    "edds_name",
    "edds_district",
    "electric_company",
    "local_admin",
}

# ---------------------------------------------------------------------------
# Known hardcoded template values — keys that exist in test contexts but
# are hardcoded in the template itself (not Jinja-templated).
# These are expected to be "unused" by the renderer.
# ---------------------------------------------------------------------------
KNOWN_HARDCODED_IN_TEMPLATE: set[str] = {
    "notification_chairman_phone",
    "notification_edds_phone",
    "notification_pasf_phone",
    "notification_gas_phone",
    "notification_rostechnadzor_phone",
}

# Full set of scalar keys the schema declares
def _schema_scalar_keys(schema: dict) -> set[str]:
    """Extract scalar (string/number) property names from Draft-2020-12 schema.
    Excludes arrays and object-$ref properties."""
    result: set[str] = set()
    for name, prop in schema.get("properties", {}).items():
        if prop.get("type") == "array":
            continue
        # Resolve $ref to check if it's a complex object, not a scalar
        if "$ref" in prop:
            def_name = prop["$ref"].split("/")[-1]
            def_obj = schema.get("$defs", {}).get(def_name, {})
            if def_obj.get("type") == "object":
                continue  # e.g. financial_reserve → FinancialReserve
        result.add(name)
    return result


# Full set of loop list names the schema declares
def _schema_loop_list_names(schema: dict) -> set[str]:
    """Extract array property names from Draft-2020-12 schema."""
    result: set[str] = set()
    for name, prop in schema.get("properties", {}).items():
        if prop.get("type") == "array":
            result.add(name)
    return result


# Resolve $ref to get item field names for a loop list
def _schema_array_item_fields(schema: dict, list_name: str) -> list[str]:
    """Resolve $ref for an array's items to get field names from $defs."""
    items = schema.get("properties", {}).get(list_name, {}).get("items", {})
    ref: str = items.get("$ref", "")
    def_name = ref.split("/")[-1]  # e.g. "EquipmentItem"
    def_obj = schema.get("$defs", {}).get(def_name, {})
    return list(def_obj.get("properties", {}).keys())


# ---------------------------------------------------------------------------
# 1. All Jinja template variables exist in schema or approved computed list
# ---------------------------------------------------------------------------
@pytest.mark.audit_only
class TestTemplateVarsInSchema:
    """Every scalar variable found in the template must appear in the schema
    (required_fields + notification_phones) or in APPROVED_COMPUTED_SCALARS."""

    def test_all_scalar_template_vars_in_schema_or_approved(
        self, template_keys: dict, schema: dict,
    ):
        schema_scalars = _schema_scalar_keys(schema)
        approved = APPROVED_COMPUTED_SCALARS

        for var in template_keys.get("variables", []):
            if var["type"] != "scalar":
                continue
            # Strip Jinja filters like | default('', true)
            raw_name = var["name"].split("|")[0].strip()
            # Skip loop built-ins and complex expressions
            if raw_name.startswith("loop.") or "if " in raw_name or raw_name.startswith("'"):
                continue
            assert raw_name in schema_scalars or raw_name in approved, (
                f"Scalar variable '{raw_name}' found in template but NOT in "
                f"schema required_fields/notification_phones nor in "
                f"APPROVED_COMPUTED_SCALARS"
            )


# ---------------------------------------------------------------------------
# 2. All template lists are described in schema
# ---------------------------------------------------------------------------
@pytest.mark.audit_only
class TestTemplateListsInSchema:
    """Every loop list found in the template must have a corresponding
    entry in schema.required_tables."""

    def test_all_loop_lists_described_in_schema(
        self, template_keys: dict, schema: dict,
    ):
        schema_lists = _schema_loop_list_names(schema)

        for loop in template_keys.get("loop_lists", []):
            list_name = loop["list_name"]
            assert list_name in schema_lists, (
                f"Loop list '{list_name}' found in template but NOT in "
                f"schema.required_tables"
            )


# ---------------------------------------------------------------------------
# 3. All renderer keys exist in schema
# ---------------------------------------------------------------------------
@pytest.mark.audit_only
class TestRendererKeysInSchema:
    """The renderer (PmlaTemplateRenderer) accepts a context dict. Every
    key the renderer passes to docxtpl must exist in the schema."""

    def test_renderer_scalar_keys_in_schema(self, schema: dict, context_keys: dict):
        schema_scalars = _schema_scalar_keys(schema)
        # The "approved computed" keys from context_keys that are in the
        # full_context_keys but not in schema must be in APPROVED_COMPUTED_SCALARS
        full_ctx = set(
            context_keys.get("full_context_keys", {}).get(
                "test_pmla_v2_render.FULL_CONTEXT", []
            )
        )
        for key in full_ctx:
            if key in schema_scalars:
                continue
            if key in APPROVED_COMPUTED_SCALARS:
                continue
            if key in KNOWN_HARDCODED_IN_TEMPLATE:
                continue
            # Check if it's a loop list
            if key in _schema_loop_list_names(schema):
                continue
            pytest.fail(
                f"Renderer context key '{key}' not in schema required_fields/"
                f"notification_phones, not in APPROVED_COMPUTED_SCALARS, "
                f"and not a known loop list"
            )


# ---------------------------------------------------------------------------
# 4. All required schema fields are available in renderer
# ---------------------------------------------------------------------------
@pytest.mark.audit_only
class TestSchemaFieldsAvailableInRenderer:
    """Every field in schema.required_fields and schema.notification_phones
    must appear in the renderer's context (either in full_context_keys
    or be a known loop list)."""

    def test_required_fields_in_renderer_context(
        self, schema: dict, context_keys: dict,
    ):
        full_ctx = set(
            context_keys.get("full_context_keys", {}).get(
                "test_pmla_v2_render.FFULL_CONTEXT",
                context_keys.get("full_context_keys", {}).get(
                    "test_pmla_v2_render.FULL_CONTEXT", []
                ),
            )
        )
        all_schema_scalars = _schema_scalar_keys(schema)
        # Also include approved computed (these are in the template, so the
        # renderer DOES provide them)
        all_renderer_scalars = full_ctx | APPROVED_COMPUTED_SCALARS | KNOWN_HARDCODED_IN_TEMPLATE

        missing = all_schema_scalars - all_renderer_scalars
        assert not missing, (
            f"Schema fields NOT available in renderer context: {missing}"
        )

    def test_schema_loop_lists_in_renderer_context(
        self, schema: dict, context_keys: dict,
    ):
        full_ctx = set(
            context_keys.get("full_context_keys", {}).get(
                "test_pmla_v2_render.FULL_CONTEXT", []
            )
        )
        schema_lists = _schema_loop_list_names(schema)
        missing = schema_lists - full_ctx
        assert not missing, (
            f"Schema loop lists NOT in renderer context: {missing}"
        )


# ---------------------------------------------------------------------------
# 5. Array item fields match between template and schema
# ---------------------------------------------------------------------------
@pytest.mark.audit_only
class TestArrayItemFieldsMatch:
    """For each loop list, the fields accessed in the template (via
    loop_variable.field) must match the item_fields declared in the
    schema.required_tables."""

    def _check_array_fields(
        self, list_name: str, prefix: str, schema: dict, template_keys: dict,
    ):
        schema_fields = set(_schema_array_item_fields(schema, list_name))
        template_fields: set[str] = set()
        for v in template_keys.get("variables", []):
            if v["type"] != "nested":
                continue
            # Use regex to extract all <prefix>.<field> occurrences
            for m in re.finditer(rf"{re.escape(prefix)}\.(\w+)", v["name"]):
                template_fields.add(m.group(1))
        assert template_fields == schema_fields, (
            f"{list_name} fields mismatch: "
            f"in template only: {template_fields - schema_fields}, "
            f"in schema only: {schema_fields - template_fields}"
        )

    def test_equipment_list_fields(self, template_keys: dict, schema: dict):
        self._check_array_fields("equipment_list", "eq", schema, template_keys)

    def test_substance_params_fields(self, template_keys: dict, schema: dict):
        self._check_array_fields("substance_params", "param", schema, template_keys)

    def test_equipment_scenario_links_fields(self, template_keys: dict, schema: dict):
        self._check_array_fields("equipment_scenario_links", "link", schema, template_keys)

    def test_accident_scenarios_fields(self, template_keys: dict, schema: dict):
        self._check_array_fields("accident_scenarios", "scenario", schema, template_keys)

    def test_injury_history_fields(self, template_keys: dict, schema: dict):
        self._check_array_fields("injury_history", "injury", schema, template_keys)

    def test_accident_history_fields(self, template_keys: dict, schema: dict):
        self._check_array_fields("accident_history", "accident", schema, template_keys)

    def test_material_reserve_fields(self, template_keys: dict, schema: dict):
        self._check_array_fields("material_reserve", "item", schema, template_keys)

    def test_countermeasures_fields(self, template_keys: dict, schema: dict):
        self._check_array_fields("countermeasures", "cm", schema, template_keys)


# ---------------------------------------------------------------------------
# 6. No name conflicts
# ---------------------------------------------------------------------------
@pytest.mark.audit_only
class TestNoNameConflicts:
    """Ensure there are no confusingly similar names between different
    loop lists or between scalar and list keys."""

    KNOWN_CONFLICTS = [
        # (schema_list_key, template_key) pairs that MUST be distinct
        ("accident_scenarios", "scenario"),  # list vs loop var
        ("material_reserve", "item"),  # list vs loop var
    ]

    def test_loop_var_names_dont_shadow_list_names(self, schema: dict, template_keys: dict):
        """Loop variable names (eq, param, link, scenario, ...) must not
        collide with any top-level list key in the schema."""
        list_names = _schema_loop_list_names(schema)
        loop_vars = {
            loop["loop_variable"]
            for loop in template_keys.get("loop_lists", [])
        }
        conflicts = list_names & loop_vars
        assert not conflicts, (
            f"Loop variable names collide with list key names: {conflicts}"
        )

    def test_no_duplicate_top_level_keys(self, schema: dict):
        """All property names must be unique (no duplicates in Draft-2020-12)."""
        # In the flat Draft-2020-12 schema, properties are already unique by construction
        prop_names = set(schema.get("properties", {}).keys())
        assert len(prop_names) == len(schema.get("properties", {})), (
            "Duplicate property names found in schema"
        )

    def test_list_keys_not_shadowed_by_scalars(self, schema: dict):
        """No loop list name should also appear as a scalar in required_fields."""
        list_names = _schema_loop_list_names(schema)
        scalar_names = _schema_scalar_keys(schema)
        shadowed = list_names & scalar_names
        assert not shadowed, (
            f"Loop list names also declared as scalars: {shadowed}"
        )


# ---------------------------------------------------------------------------
# Helper: full context for render tests
# ---------------------------------------------------------------------------
def _build_full_context() -> dict:
    """Full render context matching test_pmla_v2_render.FULL_CONTEXT."""
    return {
        "organization_full_name": "ООО «ТестПром»",
        "organization_short_name": "ООО «ТестПром»",
        "legal_address": "123456, г. Москва, ул. Тестовая, д. 1",
        "inn": "7701234567",
        "ogrn": "1027700000123",
        "phone": "+7 (495) 123-45-67",
        "email": "info@testprom.ru",
        "director_position_fullname": "Генеральный директор Иванов И.И.",
        "director_initials_surname": "И.И. Иванов",
        "director_initials_surname_full": "Иванов Иван Иванович",
        "deputy_chairman_fullname": "Петров Пётр Петрович",
        "main_activity_description": "01.11 Выращивание зерновых культур",
        "facility_name": "Сеть газопотребления",
        "facility_reg_number": "А34-99999-0001",
        "facility_location": "Московская область, г. Тест",
        "hazard_class": "III",
        "hazardous_substances_info": "Природный газ (метан)",
        "hazard_characteristics_116fz": "Использование горючих газов",
        "contractor_organization_name": "ООО «Спасатель»",
        "contractor_organization_short_name": "«Спасатель»",
        "contractor_agreement_date": "01.01.2026",
        "gas_supplier_name": "ПАО «Газпром»",
        "total_hazardous_substance_quantity": "0.5",
        "settlement_name": "г. Тест",
        "gas_supplier_branch": "Абонентский отдел",
        "dislocation_address": "г. Тест, ул. Аварийная, д. 1",
        "settlement_district": "Тестовый район",
        "edds_name": "ЕДДС г. Тест",
        "edds_district": "Тестовый",
        "electric_company": "ПАО «Россети»",
        "local_admin": "Администрация г. Тест",
        "equipment_list": [
            {
                "location": "Площадка ГРП",
                "hazard_characteristic": "Газы",
                "device_name": "ГРПШ-1",
                "specifications": "Р=0.6 МПа",
                "process_codes": "2.1",
            },
        ],
        "substance_params": [
            {"parameter": "Класс опасности", "value": "4"},
        ],
        "equipment_scenario_links": [
            {
                "equipment_name": "ГРПШ-1",
                "scenario_codes": "С-1, С-2",
                "description": "Утечка газа",
                "damaging_factors": "Взрыв",
            },
        ],
        "equipment_defects": [
            {
                "equipment_name": "ГРПШ-1",
                "defect": "Разгерметизация",
                "cause": "Износ",
                "source": "Фланцы",
                "scenario": "С-1",
            },
        ],
        "accident_scenarios": [
            {
                "code": "С-1",
                "name": "Выброс газа",
                "source": "ГРПШ",
                "preconditions": "Разгерметизация",
                "signs": "Запах одоранта",
                "damaging_factors": "Взрыв",
            },
        ],
        "injury_history": [],
        "accident_history": [],
        "material_reserve": [
            {"is_group_header": True, "group_name": "СИЗ"},
            {"name": "Противогаз", "quantity": "4 шт.", "location": "Шкаф"},
        ],
        "countermeasures": [
            {
                "scenario_label": "С-1 Выброс газа",
                "signs": "Запах одоранта",
                "protection": "Отключить оборудование",
                "technical_means": "Газоанализатор",
                "executors": "Оператор",
            },
        ],
        "person": {
            "position": "Генеральный директор",
            "phone": "+7 900 000-00-00",
        },
        "notification_chairman_phone": "+7 928 709-95-15",
        "notification_deputy_phone": "+7 906 881-07-07",
        "notification_edds_phone": "112",
        "notification_pasf_phone": "+7 (903) 495-75-57",
        "notification_fire_phone": "+7 (8663) 04-14-91",
        "notification_ambulance_phone": "112/03/103",
        "notification_gas_phone": "+7 (86630) 4-18-68",
        "notification_electric_phone": "+7 (86630) 4-27-70",
        "notification_mchs_phone": "+7 (8662) 39-99-99",
        "notification_rostechnadzor_phone": "+7 (928) 307-04-62",
        "notification_admin_phone": "+7 (86630) 7-63-99",
    }


def _build_empty_context() -> dict:
    """All list fields empty, scalars kept."""
    ctx = _build_full_context()
    for key in [
        "equipment_list", "substance_params", "equipment_scenario_links",
        "equipment_defects", "accident_scenarios", "injury_history",
        "accident_history", "material_reserve", "countermeasures",
    ]:
        ctx[key] = []
    return ctx


# ---------------------------------------------------------------------------
# 7. After render, no Jinja tags remain
# ---------------------------------------------------------------------------
JINJA_TAG_RE = re.compile(r"\{[%{].*?[%}]\}")


@pytest.mark.audit_only
class TestNoJinjaArtifactsAfterRender:
    """Rendered DOCX must not contain any Jinja template tags."""

    def test_full_render_no_jinja_tags(self):
        from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer

        renderer = PmlaTemplateRenderer()
        docx_bytes = renderer.render(_build_full_context())

        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        artifacts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if JINJA_TAG_RE.search(cell.text):
                        artifacts.append(cell.text[:120])
        assert not artifacts, (
            f"Jinja artifacts found in rendered DOCX: {artifacts[:5]}"
        )

    def test_empty_render_no_jinja_tags(self):
        from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer

        renderer = PmlaTemplateRenderer()
        docx_bytes = renderer.render(_build_empty_context())

        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        artifacts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if JINJA_TAG_RE.search(cell.text):
                        artifacts.append(cell.text[:120])
        assert not artifacts, (
            f"Jinja artifacts found in empty-rendered DOCX: {artifacts[:5]}"
        )


# ---------------------------------------------------------------------------
# 8. Empty context doesn't break document
# ---------------------------------------------------------------------------
@pytest.mark.audit_only
class TestEmptyContextRenders:
    """Rendering with all lists empty must produce valid DOCX bytes."""

    def test_empty_context_produces_valid_docx(self):
        from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer

        renderer = PmlaTemplateRenderer()
        docx_bytes = renderer.render(_build_empty_context())

        assert len(docx_bytes) > 0, "Empty-context render produced 0 bytes"

        # Must be openable by python-docx
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        assert len(doc.tables) > 0, "Rendered DOCX has no tables"


# ---------------------------------------------------------------------------
# 9. DOCX is valid ZIP
# ---------------------------------------------------------------------------
@pytest.mark.audit_only
class TestDocxIsValidZip:
    """DOCX output must be a valid ZIP archive with the expected structure."""

    def test_full_render_is_valid_zip(self):
        from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer

        renderer = PmlaTemplateRenderer()
        docx_bytes = renderer.render(_build_full_context())

        assert zipfile.is_zipfile(io.BytesIO(docx_bytes)), (
            "Rendered DOCX is not a valid ZIP file"
        )

    def test_empty_render_is_valid_zip(self):
        from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer

        renderer = PmlaTemplateRenderer()
        docx_bytes = renderer.render(_build_empty_context())

        assert zipfile.is_zipfile(io.BytesIO(docx_bytes)), (
            "Empty-context rendered DOCX is not a valid ZIP file"
        )

    def test_docx_contains_required_parts(self):
        from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer

        renderer = PmlaTemplateRenderer()
        docx_bytes = renderer.render(_build_full_context())

        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
            names = zf.namelist()
            assert "[Content_Types].xml" in names, "Missing [Content_Types].xml"
            assert "word/document.xml" in names, "Missing word/document.xml"


# ---------------------------------------------------------------------------
# 10. All XML files parse without errors
# ---------------------------------------------------------------------------
@pytest.mark.audit_only
class TestAllXmlFilesParse:
    """Every XML file inside the rendered DOCX must parse without errors."""

    def test_full_render_xml_parses(self):
        from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer

        renderer = PmlaTemplateRenderer()
        docx_bytes = renderer.render(_build_full_context())

        parse_errors: list[str] = []
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
            for name in zf.namelist():
                if not name.endswith(".xml"):
                    continue
                try:
                    data = zf.read(name)
                    ET.fromstring(data)
                except ET.ParseError as e:
                    parse_errors.append(f"{name}: {e}")

        assert not parse_errors, (
            f"XML parse errors in rendered DOCX:\n"
            + "\n".join(parse_errors[:10])
        )

    def test_empty_render_xml_parses(self):
        from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer

        renderer = PmlaTemplateRenderer()
        docx_bytes = renderer.render(_build_empty_context())

        parse_errors: list[str] = []
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
            for name in zf.namelist():
                if not name.endswith(".xml"):
                    continue
                try:
                    data = zf.read(name)
                    ET.fromstring(data)
                except ET.ParseError as e:
                    parse_errors.append(f"{name}: {e}")

        assert not parse_errors, (
            f"XML parse errors in empty-rendered DOCX:\n"
            + "\n".join(parse_errors[:10])
        )

    def test_template_xml_parses(self):
        """The template itself must have parseable XML (sanity check)."""
        if not _TEMPLATE_PATH.exists():
            pytest.skip("Template file not found")

        parse_errors: list[str] = []
        with zipfile.ZipFile(_TEMPLATE_PATH, "r") as zf:
            for name in zf.namelist():
                if not name.endswith(".xml"):
                    continue
                try:
                    data = zf.read(name)
                    ET.fromstring(data)
                except ET.ParseError as e:
                    parse_errors.append(f"{name}: {e}")

        assert not parse_errors, (
            f"XML parse errors in template:\n" + "\n".join(parse_errors[:10])
        )


# ---------------------------------------------------------------------------
# Bonus: cross-check context_keys consistency with schema
# ---------------------------------------------------------------------------
@pytest.mark.audit_only
class TestContextKeysConsistency:
    """The audit context_keys JSON must be consistent with the schema."""

    def test_empty_context_lists_match_schema_lists(
        self, schema: dict, context_keys: dict,
    ):
        """All lists emptied in the empty context must be schema loop lists."""
        emptied = set(
            context_keys.get("empty_context_keys", {})
            .get("test_pmla_v2_render.EMPTY_CONTEXT", {})
            .get("emptied_lists", [])
        )
        schema_lists = _schema_loop_list_names(schema)
        unexpected = emptied - schema_lists
        assert not unexpected, (
            f"Empty context empties lists not in schema: {unexpected}"
        )

    def test_array_item_keys_match_schema(
        self, context_keys: dict, schema: dict,
    ):
        """The array_item_keys in context_keys must match schema item fields."""
        ck_items = context_keys.get("array_item_keys", {})
        for list_name, fields in ck_items.items():
            if list_name not in _schema_loop_list_names(schema):
                continue  # known: equipment_defects may not be in all schemas
            schema_fields = set(_schema_array_item_fields(schema, list_name))
            ck_fields = set(fields)
            missing_in_schema = ck_fields - schema_fields
            missing_in_ck = schema_fields - ck_fields
            assert not missing_in_schema and not missing_in_ck, (
                f"array_item_keys['{list_name}'] mismatch with schema: "
                f"in context_keys only: {missing_in_schema}, "
                f"in schema only: {missing_in_ck}"
            )
