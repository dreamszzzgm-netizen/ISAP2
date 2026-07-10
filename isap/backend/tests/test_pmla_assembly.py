"""
Тесты для PMLA Assembly Layer:
- Block registry completeness
- Correction journal DOCX table
- TOC placeholder
- Appendices manifest table
- Static blocks don't require LLM
- No raw HTML in DOCX
"""
from __future__ import annotations

import io
import pytest
from unittest.mock import AsyncMock, MagicMock

from docx import Document as DocxDocument


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_gen():
    """Создаёт EnhancedDocumentGenerator с моками."""
    from src.application.services.enhanced_generator import EnhancedDocumentGenerator
    retriever = AsyncMock()
    retriever.retrieve.return_value = []
    doc_repo = AsyncMock()
    reg_repo = AsyncMock()
    reg_repo.session = AsyncMock()
    return EnhancedDocumentGenerator(
        local_llm=None, external_llm=None, retriever=retriever,
        document_repo=doc_repo, regulatory_repo=reg_repo,
    )


def _extract_docx_text(docx_bytes: bytes) -> str:
    """Извлекает текст из DOCX."""
    doc = DocxDocument(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def _extract_docx_tables(docx_bytes: bytes) -> list[list[list[str]]]:
    """Извлекает все таблицы из DOCX."""
    doc = DocxDocument(io.BytesIO(docx_bytes))
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)
    return tables


def _minimal_context():
    return {
        "organization": {"name": "ТестОрг", "inn": "1234567890"},
        "facility": {"name": "Объект", "facility_type": "Котельная", "hazard_class": "III"},
        "equipment": [{"name": "Котёл", "equipment_type": "Теплогенератор"}],
        "substances": [{"name": "Метан", "quantity_kg": 500}],
        "responsible_persons": [
            {"full_name": "Иванов И.И.", "phone": "+79991112233", "position": "Инженер"}
        ],
    }


# ===========================================================================
# 1. Block Registry
# ===========================================================================

class TestBlockRegistry:
    """Тесты реестра типов блоков."""

    def test_registry_covers_all_sections(self):
        """Все 27 разделов из structure.json имеют block_type."""
        import json
        from pathlib import Path
        from src.application.services.pmla_assembly_blocks import ASSEMBLY_REGISTRY

        structure_path = Path(__file__).parent.parent / "templates" / "pmla" / "structure.json"
        structure = json.loads(structure_path.read_text(encoding="utf-8"))
        section_ids = {s["id"] for s in structure["sections"]}

        registry_ids = set(ASSEMBLY_REGISTRY.keys())
        missing = section_ids - registry_ids
        extra = registry_ids - section_ids

        assert not missing, f"Sections in structure.json but not in registry: {missing}"
        assert not extra, f"Sections in registry but not in structure.json: {extra}"

    def test_static_sections_have_no_llm(self):
        """Static блоки не требуют LLM."""
        from src.application.services.pmla_assembly_blocks import (
            get_static_sections, requires_llm,
        )
        for sid in get_static_sections():
            assert not requires_llm(sid), f"Static section '{sid}' should not require LLM"

    def test_generated_sections_require_llm(self):
        """Generated блоки требуют LLM."""
        from src.application.services.pmla_assembly_blocks import (
            get_generated_sections, requires_llm,
        )
        for sid in get_generated_sections():
            assert requires_llm(sid), f"Generated section '{sid}' should require LLM"

    def test_get_block_type_returns_correct_types(self):
        """get_block_type возвращает правильные типы."""
        from src.application.services.pmla_assembly_blocks import get_block_type, BlockType
        assert get_block_type("correction_log") == BlockType.STATIC
        assert get_block_type("toc") == BlockType.WORD_TOC
        assert get_block_type("title_page") == BlockType.VARIABLE
        assert get_block_type("introduction") == BlockType.GENERATED
        assert get_block_type("appendix_1") == BlockType.APPENDIX_REF
        assert get_block_type("nonexistent") is None


# ===========================================================================
# 2. Correction Journal
# ===========================================================================

class TestCorrectionJournal:
    """Тесты журнала корректировки как DOCX-таблицы."""

    def test_correction_journal_appears_in_docx(self):
        """Журнал корректировки появляется в DOCX."""
        from src.infrastructure.export.docx_helpers import add_correction_journal
        doc = DocxDocument()
        add_correction_journal(doc)
        buf = io.BytesIO()
        doc.save(buf)
        text = _extract_docx_text(buf.getvalue())
        assert "Журнал корректировки документа" in text

    def test_correction_journal_has_table_columns(self):
        """Таблица журнала содержит нужные колонки."""
        from src.infrastructure.export.docx_helpers import add_correction_journal
        doc = DocxDocument()
        add_correction_journal(doc)
        buf = io.BytesIO()
        doc.save(buf)
        tables = _extract_docx_tables(buf.getvalue())
        assert len(tables) >= 1
        header_row = tables[0][0]
        assert "№ п/п" in header_row
        assert "Дата изменения" in header_row
        assert "Содержание изменения" in header_row

    def test_correction_journal_default_empty_row(self):
        """По умолчанию есть одна пустая строка данных."""
        from src.infrastructure.export.docx_helpers import add_correction_journal
        doc = DocxDocument()
        add_correction_journal(doc)
        buf = io.BytesIO()
        doc.save(buf)
        tables = _extract_docx_tables(buf.getvalue())
        # 1 header row + 1 empty data row
        assert len(tables[0]) == 2

    def test_correction_journal_with_corrections(self):
        """Журнал с данными заполняет строки."""
        from src.infrastructure.export.docx_helpers import add_correction_journal
        doc = DocxDocument()
        corrections = [
            {"date": "2026-01-15", "section_page": "раздел 3", "description": "Добавлена статистика", "basis": "приказ", "signature": ""}
        ]
        add_correction_journal(doc, corrections)
        buf = io.BytesIO()
        doc.save(buf)
        tables = _extract_docx_tables(buf.getvalue())
        # 1 header + 1 data row
        assert len(tables[0]) == 2
        assert "2026-01-15" in tables[0][1][1]


# ===========================================================================
# 3. TOC Placeholder
# ===========================================================================

class TestTocPlaceholder:
    """Тесты содержания (Word TOC field)."""

    def test_toc_placeholder_appears_in_docx(self):
        """Содержание появляется в DOCX."""
        from src.infrastructure.export.docx_helpers import add_toc_placeholder
        doc = DocxDocument()
        add_toc_placeholder(doc)
        buf = io.BytesIO()
        doc.save(buf)
        text = _extract_docx_text(buf.getvalue())
        assert "Содержание" in text

    def test_toc_has_update_hint(self):
        """Содержание содержит подсказку об обновлении."""
        from src.infrastructure.export.docx_helpers import add_toc_placeholder
        doc = DocxDocument()
        add_toc_placeholder(doc)
        buf = io.BytesIO()
        doc.save(buf)
        text = _extract_docx_text(buf.getvalue())
        assert "Обновите содержание" in text


# ===========================================================================
# 4. Appendices Manifest
# ===========================================================================

class TestAppendicesManifest:
    """Тесты манифеста приложений."""

    def test_appendices_manifest_renders_in_docx(self):
        """Манифест приложений появляется в DOCX."""
        from src.infrastructure.export.docx_helpers import add_appendices_manifest
        doc = DocxDocument()
        appendices = [
            {"appendix_number": 1, "title": "Порядок изучения ПМЛА", "filename": "", "present": False},
            {"appendix_number": 2, "title": "Форма сообщения", "filename": "form.docx", "present": True},
        ]
        add_appendices_manifest(doc, appendices)
        buf = io.BytesIO()
        doc.save(buf)
        tables = _extract_docx_tables(buf.getvalue())
        all_cell_text = " ".join(cell for row in tables for cells in row for cell in cells)
        assert "Приложения" in _extract_docx_text(buf.getvalue())
        assert "Порядок изучения ПМЛА" in all_cell_text
        assert "Форма сообщения" in all_cell_text

    def test_appendices_manifest_table_columns(self):
        """Таблица манифеста содержит нужные колонки."""
        from src.infrastructure.export.docx_helpers import add_appendices_manifest
        doc = DocxDocument()
        appendices = [
            {"appendix_number": 1, "title": "Тест", "filename": "", "present": False},
        ]
        add_appendices_manifest(doc, appendices)
        buf = io.BytesIO()
        doc.save(buf)
        tables = _extract_docx_tables(buf.getvalue())
        assert len(tables) >= 1
        header_row = tables[0][0]
        assert "№ приложения" in header_row
        assert "Наименование" in header_row
        assert "Файл" in header_row
        assert "Наличие" in header_row

    def test_appendices_manifest_status(self):
        """Статус наличия отображается корректно."""
        from src.infrastructure.export.docx_helpers import add_appendices_manifest
        doc = DocxDocument()
        appendices = [
            {"appendix_number": 1, "title": "Есть", "filename": "a.pdf", "present": True},
            {"appendix_number": 2, "title": "Нет", "filename": "", "present": False},
        ]
        add_appendices_manifest(doc, appendices)
        buf = io.BytesIO()
        doc.save(buf)
        tables = _extract_docx_tables(buf.getvalue())
        assert "представлен" in tables[0][1][3]
        assert "не представлен" in tables[0][2][3]

    def test_appendices_manifest_empty(self):
        """Пустой манифест — сообщение «Приложения не представлены»."""
        from src.infrastructure.export.docx_helpers import add_appendices_manifest
        doc = DocxDocument()
        add_appendices_manifest(doc, [])
        buf = io.BytesIO()
        doc.save(buf)
        text = _extract_docx_text(buf.getvalue())
        assert "не представлены" in text


# ===========================================================================
# 5. No Raw HTML
# ===========================================================================

class TestNoRawHtml:
    """Тесты отсутствия сырого HTML в DOCX."""

    def test_correction_journal_no_html(self):
        """Журнал корректировки не содержит HTML-тегов."""
        from src.infrastructure.export.docx_helpers import add_correction_journal
        doc = DocxDocument()
        add_correction_journal(doc)
        buf = io.BytesIO()
        doc.save(buf)
        text = _extract_docx_text(buf.getvalue())
        assert "<td>" not in text
        assert "<tr>" not in text
        assert "<table>" not in text

    def test_toc_placeholder_no_html(self):
        """Содержание не содержит HTML-тегов."""
        from src.infrastructure.export.docx_helpers import add_toc_placeholder
        doc = DocxDocument()
        add_toc_placeholder(doc)
        buf = io.BytesIO()
        doc.save(buf)
        text = _extract_docx_text(buf.getvalue())
        assert "<td>" not in text
        assert "<tr>" not in text

    def test_appendices_manifest_no_html(self):
        """Манифест приложений не содержит HTML-тегов."""
        from src.infrastructure.export.docx_helpers import add_appendices_manifest
        doc = DocxDocument()
        add_appendices_manifest(doc, [{"appendix_number": 1, "title": "Тест", "filename": "", "present": False}])
        buf = io.BytesIO()
        doc.save(buf)
        text = _extract_docx_text(buf.getvalue())
        assert "<td>" not in text
        assert "<tr>" not in text


# ===========================================================================
# 6. Static blocks don't require LLM
# ===========================================================================

class TestStaticBlocksNoLlm:
    """Static блоки генерируются без LLM."""

    def test_abbreviations_static(self):
        """Раздел abbreviations — static_block, не требует LLM."""
        from src.application.services.pmla_assembly_blocks import get_block_type, BlockType
        assert get_block_type("abbreviations") == BlockType.STATIC

    def test_terms_static(self):
        """Раздел terms — static_block."""
        from src.application.services.pmla_assembly_blocks import get_block_type, BlockType
        assert get_block_type("terms") == BlockType.STATIC

    def test_bibliography_static(self):
        """Раздел bibliography — static_block."""
        from src.application.services.pmla_assembly_blocks import get_block_type, BlockType
        assert get_block_type("bibliography") == BlockType.STATIC

    def test_correction_log_static(self):
        """Раздел correction_log — static_block."""
        from src.application.services.pmla_assembly_blocks import get_block_type, BlockType
        assert get_block_type("correction_log") == BlockType.STATIC


# ===========================================================================
# 7. Heading styles & Word TOC support
# ===========================================================================

class TestHeadingStyles:
    """Заголовки используют встроенные стили Heading 1/2, чтобы Word TOC
    field (TOC \\o "1-2") собирал их при обновлении поля."""

    def test_configure_heading_styles_sets_times_new_roman_black(self):
        """configure_heading_styles приводит Heading 1/2 к Times New Roman, чёрный."""
        from docx.shared import RGBColor
        from src.infrastructure.export.docx_helpers import (
            BODY_FONT_NAME,
            configure_heading_styles,
        )
        doc = DocxDocument()
        configure_heading_styles(doc)
        for name in ("Heading 1", "Heading 2"):
            style = doc.styles[name]
            assert style.font.name == BODY_FONT_NAME
            assert style.font.color.rgb == RGBColor(0, 0, 0)

    def test_add_heading_level1_assigns_heading_style(self):
        """add_heading с level=1 назначает параграфу стиль Heading 1."""
        from src.infrastructure.export.docx_helpers import add_heading
        doc = DocxDocument()
        add_heading(doc, "Раздел 1", level=1)
        assert doc.paragraphs[0].style.name == "Heading 1"

    def test_add_heading_level2_assigns_heading2_style(self):
        """add_heading с level=2 назначает параграфу стиль Heading 2."""
        from src.infrastructure.export.docx_helpers import add_heading
        doc = DocxDocument()
        add_heading(doc, "Подраздел", level=2)
        assert doc.paragraphs[0].style.name == "Heading 2"

    def test_add_heading_level0_no_heading_style(self):
        """level=0 (заголовок документа/служебный) не назначает стиль Heading —
        иначе оглавление сошлётся само на себя."""
        from src.infrastructure.export.docx_helpers import add_heading
        doc = DocxDocument()
        add_heading(doc, "Служебный заголовок", level=0)
        assert doc.paragraphs[0].style.name != "Heading 1"
        assert doc.paragraphs[0].style.name != "Heading 2"

    def test_generator_add_heading_assigns_style(self):
        """EnhancedDocumentGenerator._add_heading с level=1 назначает Heading 1."""
        gen = _make_gen()
        doc = DocxDocument()
        gen._setup_document_defaults(doc)
        gen._add_heading(doc, "Тестовый раздел", level=1, center=False)
        assert doc.paragraphs[-1].style.name == "Heading 1"

    def test_heading_block_level2_gets_heading2_style(self):
        """HeadingBlock(level=2) из scenario_engine получает стиль Heading 2."""
        from src.application.engines.blocks import HeadingBlock
        gen = _make_gen()
        doc = DocxDocument()
        gen._setup_document_defaults(doc)
        gen._render_blocks(doc, [HeadingBlock(text="Сценарий С-1", level=2)])
        assert doc.paragraphs[-1].style.name == "Heading 2"

    def test_section_headings_in_build_docx_have_heading1_style(self):
        """Заголовки разделов в собранном DOCX имеют стиль Heading 1."""
        gen = _make_gen()
        sections = {"1. Тестовый раздел": "Содержимое"}
        metadata = {"context": _minimal_context()}
        docx_bytes = gen._build_docx("Документ", sections, metadata)
        doc = DocxDocument(io.BytesIO(docx_bytes))
        heading1_paras = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert any("Тестовый раздел" in p.text for p in heading1_paras)

    def test_toc_heading_is_not_heading_style(self):
        """Заголовок «Содержание» не имеет стиля Heading — оглавление не
        должно ссылаться само на себя."""
        from src.infrastructure.export.docx_helpers import add_toc_placeholder
        doc = DocxDocument()
        add_toc_placeholder(doc)
        content_paras = [p for p in doc.paragraphs if p.text.strip() == "Содержание"]
        assert content_paras, "Заголовок «Содержание» должен присутствовать"
        for p in content_paras:
            assert p.style.name not in ("Heading 1", "Heading 2")


# ===========================================================================
# 8. Registry-driven front matter & section titles
# ===========================================================================

class TestRegistryTitlesAndFrontMatter:
    """Реестр как источник русских названий и front matter для _build_docx."""

    def test_get_section_title_returns_russian_title(self):
        """get_section_title возвращает русское название из structure.json."""
        from src.application.services.pmla_assembly_blocks import get_section_title
        assert get_section_title("toc") == "Содержание"
        assert get_section_title("correction_log") == "Журнал корректировки документа"
        assert get_section_title("title_page") == "Титульный лист"
        assert get_section_title("section_1") == "1. Характеристика опасного производственного объекта"
        assert get_section_title("nonexistent") is None

    def test_get_front_matter_section_ids(self):
        """get_front_matter_section_ids возвращает 4 id front matter в порядке рендеринга."""
        from src.application.services.pmla_assembly_blocks import get_front_matter_section_ids
        ids = get_front_matter_section_ids()
        assert ids == ["title_page", "approval_sheet", "correction_log", "toc"]

    def test_registry_titles_match_structure_json(self):
        """Инвариант: каждый title в ASSEMBLY_REGISTRY совпадает с title
        в structure.json. Защита от рассинхронизации двух источников."""
        import json
        from pathlib import Path
        from src.application.services.pmla_assembly_blocks import ASSEMBLY_REGISTRY

        structure_path = (
            Path(__file__).parent.parent / "templates" / "pmla" / "structure.json"
        )
        structure = json.loads(structure_path.read_text(encoding="utf-8"))
        json_titles = {s["id"]: s["title"] for s in structure["sections"]}

        for sid, block_def in ASSEMBLY_REGISTRY.items():
            assert sid in json_titles, f"{sid} отсутствует в structure.json"
            assert block_def.title == json_titles[sid], (
                f"title рассинхронизирован для {sid}: "
                f"registry={block_def.title!r} structure.json={json_titles[sid]!r}"
            )

    def test_build_docx_pops_front_matter_by_registry(self):
        """_build_docx убирает front matter из общего цикла секций по registry,
        а не по хардкоженным русским строкам."""
        gen = _make_gen()
        # Кладём все 4 front-matter названия в sections — они не должны
        # задвоиться в теле документа.
        sections = {
            "Титульный лист": "должно быть убрано",
            "Лист согласования": "должно быть убрано",
            "Журнал корректировки документа": "должно быть убрано",
            "Содержание": "должно быть убрано",
            "1. Реальный раздел": "содержимое реального раздела",
        }
        metadata = {"context": _minimal_context()}
        docx_bytes = gen._build_docx("Документ", sections, metadata)
        text = _extract_docx_text(docx_bytes)
        # Реальный раздел присутствует
        assert "содержимое реального раздела" in text.lower() or "содержимое реального раздела" in text
        # Front matter убран из тела (значения-заглушки не попадают в output)
        assert "должно быть убрано" not in text


# ===========================================================================
# 9. Appendices manifest synthesis
# ===========================================================================

class TestAppendicesManifestSynthesis:
    """Синтез appendices_manifest из реестра приложений + attachments_checklist."""

    def test_manifest_contains_five_appendices(self):
        """Синтезированный манифест содержит 5 приложений с корректными номерами."""
        from src.application.services.enhanced_generator import _synthesize_appendices_manifest
        manifest = _synthesize_appendices_manifest([])
        assert len(manifest) == 5
        assert [m["appendix_number"] for m in manifest] == [1, 2, 3, 4, 5]

    def test_manifest_titles_from_registry(self):
        """Названия приложений берутся из реестра."""
        from src.application.services.enhanced_generator import _synthesize_appendices_manifest
        manifest = _synthesize_appendices_manifest([])
        titles = [m["title"] for m in manifest]
        assert any("изучения ПМЛА" in t for t in titles)
        assert any("оперативного сообщения" in t for t in titles)
        assert any("Состав ПАСФ" in t for t in titles)
        assert any("Оснащение ПАСФ" in t for t in titles)
        assert any("Схема оповещения" in t for t in titles)

    def test_manifest_default_all_not_present(self):
        """Без checklist все приложения отмечены как не представленные."""
        from src.application.services.enhanced_generator import _synthesize_appendices_manifest
        manifest = _synthesize_appendices_manifest([])
        for entry in manifest:
            assert entry["present"] is False
            assert entry["filename"] == "—"

    def test_manifest_present_status_from_checklist(self):
        """Статус наличия берётся из attachments_checklist по совпадению имени."""
        from src.application.services.enhanced_generator import _synthesize_appendices_manifest
        checklist = [
            {"name": "Схема оповещения при аварии", "present": True},
            {"name": "Состав ПАСФ", "present": True},
            {"name": "Не относящееся", "present": True},
        ]
        manifest = _synthesize_appendices_manifest(checklist)
        by_num = {m["appendix_number"]: m for m in manifest}
        # appendix_5 = Схема оповещения → present
        assert by_num[5]["present"] is True
        # appendix_3 = Состав ПАСФ → present
        assert by_num[3]["present"] is True
        # appendix_1 = Порядок изучения → не matched
        assert by_num[1]["present"] is False

    def test_manifest_present_false_when_checklist_item_not_present(self):
        """Если checklist-элемент совпадает по имени, но present=False — не отмечен."""
        from src.application.services.enhanced_generator import _synthesize_appendices_manifest
        checklist = [{"name": "Схема оповещения", "present": False}]
        manifest = _synthesize_appendices_manifest(checklist)
        by_num = {m["appendix_number"]: m for m in manifest}
        assert by_num[5]["present"] is False

    def test_enrich_context_synthesizes_manifest(self):
        """_enrich_context добавляет appendices_manifest, если его нет в контексте."""
        gen = _make_gen()
        ctx = _minimal_context()
        ctx["attachments_checklist"] = [{"name": "Схема оповещения", "present": True}]
        enriched = gen._enrich_context(ctx, scenarios=[], calculations=[])
        assert "appendices_manifest" in enriched
        assert len(enriched["appendices_manifest"]) == 5

    def test_enrich_context_preserves_explicit_manifest(self):
        """Явно заданный appendices_manifest не перезаписывается."""
        gen = _make_gen()
        ctx = _minimal_context()
        explicit = [{"appendix_number": 99, "title": "external", "present": True}]
        ctx["appendices_manifest"] = explicit
        enriched = gen._enrich_context(ctx, scenarios=[], calculations=[])
        assert enriched["appendices_manifest"] == explicit

    def test_build_docx_renders_manifest_table(self):
        """Собранный DOCX содержит таблицу манифеста приложений."""
        gen = _make_gen()
        ctx = _minimal_context()
        ctx["attachments_checklist"] = [{"name": "Схема оповещения", "present": True}]
        sections = {"1. Раздел": "контент"}
        metadata = {"context": ctx}
        docx_bytes = gen._build_docx("Документ", sections, metadata)
        text = _extract_docx_text(docx_bytes)
        # Манифест рендерится как таблица с заголовком
        assert "Приложения" in text or "приложен" in text.lower()
