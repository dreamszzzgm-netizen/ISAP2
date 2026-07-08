"""Tests for PMLA document download endpoint."""
import io
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

import pytest
from fastapi.testclient import TestClient

from src.main import app


client = TestClient(app, raise_server_exceptions=False)


def _doc_id():
    return str(uuid4())


def _mock_document(docx_bytes=b"fake-docx-content", status="pending_review"):
    doc = MagicMock()
    doc.id = uuid4()
    doc.status = status
    doc.content_docx = docx_bytes
    return doc


class TestDocumentDownload:
    """Tests for GET /api/v1/pmla/{document_id}/download."""

    @patch("src.api.routers.pmla.get_document_repo")
    @patch("src.api.routers.pmla.PmlaExportService")
    def test_download_returns_docx_content(self, mock_service_cls, mock_repo_dep):
        doc = _mock_document(docx_bytes=b"PK-zip-docx-bytes")
        mock_service = MagicMock()
        mock_service.get_docx = AsyncMock(return_value=MagicMock(
            filename="pmla_test.docx",
            content=doc.content_docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ))
        mock_service_cls.return_value = mock_service

        response = client.get(f"/api/v1/pmla/{doc.id}/download")

        assert response.status_code == 200
        assert response.content == doc.content_docx
        assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in response.headers["content-type"]

    @patch("src.api.routers.pmla.get_document_repo")
    @patch("src.api.routers.pmla.PmlaExportService")
    def test_download_404_for_missing_document(self, mock_service_cls, mock_repo_dep):
        mock_service = MagicMock()
        mock_service.get_docx = AsyncMock(side_effect=ValueError("Документ не найден"))
        mock_service_cls.return_value = mock_service

        response = client.get(f"/api/v1/pmla/{_doc_id()}/download")

        assert response.status_code == 404

    @patch("src.api.routers.pmla.get_document_repo")
    @patch("src.api.routers.pmla.PmlaExportService")
    def test_download_404_for_empty_content(self, mock_service_cls, mock_repo_dep):
        mock_service = MagicMock()
        mock_service.get_docx = AsyncMock(side_effect=FileNotFoundError("content_docx пуст"))
        mock_service_cls.return_value = mock_service

        response = client.get(f"/api/v1/pmla/{_doc_id()}/download")

        assert response.status_code == 404

    @patch("src.api.routers.pmla.get_document_repo")
    @patch("src.api.routers.pmla.PmlaExportService")
    def test_download_content_disposition_header(self, mock_service_cls, mock_repo_dep):
        doc_id = uuid4()
        mock_service = MagicMock()
        mock_service.get_docx = AsyncMock(return_value=MagicMock(
            filename=f"pmla_{doc_id}.docx",
            content=b"content",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ))
        mock_service_cls.return_value = mock_service

        response = client.get(f"/api/v1/pmla/{doc_id}/download")

        assert response.status_code == 200
        assert f"pmla_{doc_id}.docx" in response.headers.get("content-disposition", "")

    @patch("src.api.routers.pmla.get_document_repo")
    @patch("src.api.routers.pmla.PmlaExportService")
    def test_download_works_for_pending_review_status(self, mock_service_cls, mock_repo_dep):
        """Download should work regardless of document status."""
        doc = _mock_document(status="pending_review")
        mock_service = MagicMock()
        mock_service.get_docx = AsyncMock(return_value=MagicMock(
            filename="pmla_test.docx",
            content=doc.content_docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ))
        mock_service_cls.return_value = mock_service

        response = client.get(f"/api/v1/pmla/{doc.id}/download")

        assert response.status_code == 200

    @patch("src.api.routers.pmla.get_document_repo")
    @patch("src.api.routers.pmla.PmlaExportService")
    def test_download_pdf_returns_pdf_content(self, mock_service_cls, mock_repo_dep):
        doc_id = uuid4()
        pdf_bytes = b"%PDF-1.7\ncontent"
        mock_service = MagicMock()
        mock_service.get_pdf = AsyncMock(return_value=MagicMock(
            filename=f"pmla_{doc_id}.pdf",
            content=pdf_bytes,
            media_type="application/pdf",
        ))
        mock_service_cls.return_value = mock_service

        response = client.get(f"/api/v1/pmla/{doc_id}/download/pdf")

        assert response.status_code == 200
        assert response.content == pdf_bytes
        assert "application/pdf" in response.headers["content-type"]
        assert f"pmla_{doc_id}.pdf" in response.headers.get("content-disposition", "")


class TestExportService:
    """Tests for PmlaExportService._get_downloadable_document."""

    @pytest.mark.asyncio
    async def test_download_requires_content_docx(self):
        from src.application.services.pmla_export_service import PmlaExportService

        mock_repo = AsyncMock()
        doc = MagicMock()
        doc.content_docx = None
        mock_repo.get = AsyncMock(return_value=doc)

        service = PmlaExportService(mock_repo)

        with pytest.raises(FileNotFoundError, match="content_docx"):
            await service._get_downloadable_document(uuid4())

    @pytest.mark.asyncio
    async def test_download_rejects_missing_document(self):
        from src.application.services.pmla_export_service import PmlaExportService

        mock_repo = AsyncMock()
        mock_repo.get = AsyncMock(return_value=None)

        service = PmlaExportService(mock_repo)

        with pytest.raises(ValueError, match="не найден"):
            await service._get_downloadable_document(uuid4())

    @pytest.mark.asyncio
    async def test_download_accepts_pending_review(self):
        """Download works for any status as long as content_docx exists."""
        from src.application.services.pmla_export_service import PmlaExportService

        mock_repo = AsyncMock()
        doc = MagicMock()
        doc.content_docx = b"docx-content"
        doc.status = "pending_review"
        mock_repo.get = AsyncMock(return_value=doc)

        service = PmlaExportService(mock_repo)
        result = await service._get_downloadable_document(uuid4())

        assert result.content_docx == b"docx-content"

    @pytest.mark.asyncio
    async def test_get_docx_returns_binary_result(self):
        from src.application.services.pmla_export_service import PmlaExportService

        mock_repo = AsyncMock()
        doc = MagicMock()
        doc.content_docx = b"real-docx-bytes"
        doc.status = "processing"
        mock_repo.get = AsyncMock(return_value=doc)

        service = PmlaExportService(mock_repo)
        result = await service.get_docx(uuid4())

        assert result.content == b"real-docx-bytes"
        assert result.media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert ".docx" in result.filename

    @pytest.mark.asyncio
    async def test_get_pdf_returns_binary_result(self):
        from src.application.services.pmla_export_service import PmlaExportService

        document_id = uuid4()
        mock_repo = AsyncMock()
        doc = MagicMock()
        doc.content_docx = b"docx-content"
        mock_repo.get = AsyncMock(return_value=doc)

        service = PmlaExportService(mock_repo)

        with patch(
            "src.infrastructure.pdf.converter.docx_bytes_to_pdf",
            return_value=b"%PDF-1.7\nconverted",
        ):
            result = await service.get_pdf(document_id)

        assert result.content.startswith(b"%PDF-")
        assert result.media_type == "application/pdf"
        assert result.filename == f"pmla_{document_id}.pdf"


class TestPdfConverter:
    """Tests for DOCX to PDF conversion fallback path."""

    def test_pdf_conversion_works_without_ms_office(self):
        from docx import Document

        from src.infrastructure.pdf import converter

        doc = Document()
        doc.add_heading("PMLA Test", level=1)
        doc.add_paragraph("Fallback PDF content")
        buffer = io.BytesIO()
        doc.save(buffer)

        with patch.object(
            converter,
            "_convert_via_libreoffice",
            side_effect=RuntimeError("soffice unavailable"),
        ):
            result = converter.docx_bytes_to_pdf(buffer.getvalue())

        assert result.startswith(b"%PDF-")
        assert len(result) > 100
