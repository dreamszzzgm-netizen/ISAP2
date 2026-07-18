"""Regression test: runtime PMLA v2 template must not contain old hardcoded service contacts.

After the template cleanup, the old hardcoded phone numbers and organization names
of the previous pilot object (Кабардино-Балкария, Чегемский район) must NOT be
present in the runtime template or rendered output. They were replaced with
Jinja placeholders that consume values from map_to_v2_context().
"""
from __future__ import annotations

import io
import json
import re
import zipfile
from pathlib import Path

import pytest

_TESTS_DIR = Path(__file__).resolve().parent  # backend/tests/
_PROJECT_ROOT = _TESTS_DIR.parent               # backend/
_FILES_DIR = _PROJECT_ROOT.parent / "files"     # isap/files
_CONTAINER_FILES_DIR = Path("/files")
_TEMPLATE_PATH = _FILES_DIR / "pmla_v2_template.docx"
if not _TEMPLATE_PATH.exists():
    _TEMPLATE_PATH = _CONTAINER_FILES_DIR / "pmla_v2_template.docx"

# Old hardcoded phone numbers that MUST NOT appear in rendered PMLA v2 output
# (notification table and scenario text — not in embedded drawing shapes, which
# are out-of-scope for Jinja templating and require redrawing the schema diagram).
OLD_HARDCODED_PHONES: list[str] = [
    "+7 (8663) 04-14-91",      # old fire department
    "+7 (86630) 4-18-68",      # old gas supplier
    "+7 (86630) 4-27-70",      # old electric networks
    "+7 (928) 307-04-62",      # old Rostekhnadzor
    "+7 (8662) 39-99-99",      # old MChS
    "+7 (86630) 7-63-99",      # old administration
    "+7 (903) 495-75-57",      # old PASF / contractor
    "+7 (903) 491-85-75",      # old PASF secondary
    "+7 928 709-95-15",        # old chairman
]

# Old hardcoded organization / location names that MUST NOT appear in rendered output
OLD_HARDCODED_NAMES: list[str] = [
    "Чегемские РЭС",
]


def _extract_template_text() -> str:
    """Extract concatenated w:t text from the runtime DOCX template."""
    with zipfile.ZipFile(_TEMPLATE_PATH) as zf:
        content = zf.read("word/document.xml").decode("utf-8", errors="replace")
    parts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", content, re.DOTALL)
    return "".join(parts)


class TestRuntimeTemplateHasNoHardcodedServiceContacts:
    """The runtime template must not contain old hardcoded service contacts."""

    def test_template_has_no_old_hardcoded_phones(self):
        """Old pilot phones must have been replaced with Jinja placeholders."""
        text = _extract_template_text()
        present = [p for p in OLD_HARDCODED_PHONES if p in text]
        assert not present, (
            f"Old hardcoded phones still in template: {present}. "
            "They should have been replaced with Jinja placeholders."
        )

    def test_template_has_no_old_hardcoded_organization_names(self):
        """Old pilot org/location names must have been replaced."""
        text = _extract_template_text()
        present = [n for n in OLD_HARDCODED_NAMES if n in text]
        assert not present, (
            f"Old hardcoded names still in template: {present}. "
            "They should have been replaced with Jinja placeholders."
        )

    def test_template_has_jinja_placeholders_for_service_phones(self):
        """After cleanup, the template must reference notification_*_phone Jinja vars."""
        text = _extract_template_text()
        required = [
            "{{ notification_fire_phone }}",
            "{{ notification_gas_phone }}",
            "{{ notification_pasf_phone }}",
            "{{ notification_rostechnadzor_phone }}",
            "{{ notification_mchs_phone }}",
        ]
        missing = [r for r in required if r not in text]
        assert not missing, (
            f"Required Jinja phone placeholders missing from template: {missing}"
        )

    def test_rendered_output_has_no_old_hardcoded_phones(self):
        """Full render through map_to_v2_context + PmlaTemplateRenderer must
        not contain any of the old hardcoded phone numbers in its output."""
        fixture_path = _TESTS_DIR / "fixtures" / "pmla_v2_pilot_case.json"
        if not fixture_path.exists():
            pytest.skip("Pilot fixture not found")
        from src.application.services.pmla_v2_context_mapper import map_to_v2_context
        from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer

        pilot = json.loads(fixture_path.read_text(encoding="utf-8"))
        ctx = map_to_v2_context(pilot)
        docx_bytes = PmlaTemplateRenderer().render(ctx)

        # Extract text from rendered DOCX
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        text = "\n".join(parts)

        present = [p for p in OLD_HARDCODED_PHONES if p in text]
        assert not present, (
            f"Old hardcoded phones found in rendered output: {present}"
        )
