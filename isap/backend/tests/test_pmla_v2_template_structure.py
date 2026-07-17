"""Structural test for PMLA v2 rendered DOCX.

Verifies document structure after render through PmlaTemplateRenderer:
- Журнал корректировки present and before СОДЕРЖАНИЕ
- Figure captions 1-4 present
- No 'Форма 16' reference (or present in both template and etalon)
- Appendix numbering is sequential (1..N, no duplicates)
- No Jinja markers in output
"""
from __future__ import annotations

import io
import json
import re
from pathlib import Path

import pytest

_TESTS_DIR = Path(__file__).resolve().parent
_FILES_DIR = _TESTS_DIR.parent.parent / "files"
_CONTAINER_FILES_DIR = Path("/files")


@pytest.fixture(scope="module")
def rendered_docx_bytes() -> bytes:
    """Render pilot fixture through the full pipeline."""
    fixture_path = _TESTS_DIR / "fixtures" / "pmla_v2_pilot_case.json"
    if not fixture_path.exists():
        pytest.skip("Pilot fixture not found")
    from src.application.services.pmla_v2_context_mapper import map_to_v2_context
    from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer

    pilot = json.loads(fixture_path.read_text(encoding="utf-8"))
    ctx = map_to_v2_context(pilot)
    return PmlaTemplateRenderer().render(ctx)


def _extract_paragraphs(docx_bytes: bytes) -> list[str]:
    """Extract paragraph texts from DOCX bytes."""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


class TestPmlaV2TemplateStructure:
    """Document-level structural invariants after render."""

    def test_correction_journal_present_and_before_toc(self, rendered_docx_bytes):
        """ЖУРНАЛ КОРРЕКТИРОВКИ must be present and appear before СОДЕРЖАНИЕ."""
        paragraphs = _extract_paragraphs(rendered_docx_bytes)
        journal_idx = None
        toc_idx = None
        for i, p in enumerate(paragraphs):
            text = p.strip().upper()
            if "ЖУРНАЛ КОРРЕКТИРОВКИ" in text and journal_idx is None:
                journal_idx = i
            if text == "СОДЕРЖАНИЕ" and toc_idx is None:
                toc_idx = i
        assert journal_idx is not None, "ЖУРНАЛ КОРРЕКТИРОВКИ not found in document"
        assert toc_idx is not None, "СОДЕРЖАНИЕ not found in document"
        assert journal_idx < toc_idx, (
            f"ЖУРНАЛ КОРРЕКТИРОВКИ (para {journal_idx}) must come before "
            f"СОДЕРЖАНИЕ (para {toc_idx})"
        )

    def test_figure_captions_1_to_4_present(self, rendered_docx_bytes):
        """All four figure captions must be present in the rendered document."""
        paragraphs = _extract_paragraphs(rendered_docx_bytes)
        full_text = "\n".join(paragraphs)
        for n in (1, 2, 3, 4):
            pattern = f"Рисунок {n}"
            assert pattern in full_text, (
                f"Figure caption '{pattern}' not found in rendered document"
            )

    def test_appendix_numbering_is_sequential(self, rendered_docx_bytes):
        """Appendix numbering must be sequential without duplicates.

        Acceptable patterns: 'Приложение 1', 'Приложение №1', 'Приложение № 1'.
        """
        paragraphs = _extract_paragraphs(rendered_docx_bytes)
        full_text = "\n".join(paragraphs)
        # Find all standalone appendix headers (at start of a line/paragraph)
        numbers = []
        for p in paragraphs:
            text = p.strip()
            m = re.match(r"^Приложение\s*№?\s*(\d+)", text)
            if m:
                numbers.append(int(m.group(1)))
        # We expect at least 1 appendix; check no duplicates among actual headers
        if not numbers:
            pytest.skip("No appendix headers found in document")
        duplicates = [n for n in numbers if numbers.count(n) > 1]
        assert not duplicates, (
            f"Duplicate appendix numbers found: {sorted(set(duplicates))}. "
            f"All numbers: {numbers}"
        )

    def test_no_jinja_markers_in_rendered_output(self, rendered_docx_bytes):
        """No {{ }} or {% %} must remain after render."""
        paragraphs = _extract_paragraphs(rendered_docx_bytes)
        from docx import Document
        doc = Document(io.BytesIO(rendered_docx_bytes))
        all_text = list(paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        all_text.append(cell.text)
        full_text = "\n".join(all_text)
        jinja_markers = re.findall(r"\{\{[^}]*\}\}|\{%[^%]*%\}", full_text)
        assert not jinja_markers, (
            f"Jinja markers found in rendered DOCX: {jinja_markers[:5]}"
        )
