"""Shared utilities for isolated PMLA template modification agents.

Each agent works on its own copy of the template and returns modified bytes.
The coordinator merges all changes.
"""
from __future__ import annotations

import copy
import os
from pathlib import Path
from lxml import etree
from docx import Document
from docx.oxml.ns import qn

# agents/ → scripts/ → backend/ → isap/ (project root)
FILES_DIR = Path(__file__).parent.parent.parent.parent / "files"
BACKUP = FILES_DIR / "pmla_v1_template.docx.bak"


def load_template_copy(agent_name: str) -> Document:
    """Load a fresh copy of the template for this agent."""
    return Document(str(BACKUP))


def save_agent_result(doc: Document, agent_name: str) -> Path:
    """Save the agent's modified template to a temp file."""
    out_dir = FILES_DIR / "agent_work"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / f"pmla_v2_{agent_name}.docx"
    doc.save(str(out_path))
    return out_path


def get_cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs)


def make_run(parent, text: str):
    """Create a w:r element with w:t child."""
    r = etree.SubElement(parent, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    t.set(qn("xml:space"), "preserve")
    return r


def remove_rows(table, start_idx: int, end_idx: int):
    """Remove rows from start_idx to end_idx (inclusive)."""
    for i in range(end_idx, start_idx - 1, -1):
        tr = table.rows[i]._tr
        tr.getparent().remove(tr)


def insert_loop_rows(table, header_row_idx: int, for_tag: str, field_tags: list[str], endfor_tag: str):
    """Insert Jinja2 for/data/endfor rows after the header row."""
    header_tr = table.rows[header_row_idx]._tr

    # Create for row
    for_tr = copy.deepcopy(header_tr)
    cells = for_tr.findall(f".//{qn('w:tc')}")
    for ci, tc in enumerate(cells):
        for p in tc.findall(f".//{qn('w:p')}"):
            for r in list(p.findall(qn("w:r"))):
                p.remove(r)
        if ci == 0:
            p = tc.findall(f".//{qn('w:p')}")[0]
            make_run(p, for_tag)

    # Create data row
    data_tr = copy.deepcopy(header_tr)
    cells = data_tr.findall(f".//{qn('w:tc')}")
    for ci, tc in enumerate(cells):
        for p in tc.findall(f".//{qn('w:p')}"):
            for r in list(p.findall(qn("w:r"))):
                p.remove(r)
        if ci < len(field_tags):
            p = tc.findall(f".//{qn('w:p')}")[0]
            make_run(p, field_tags[ci])

    # Create endfor row
    endfor_tr = copy.deepcopy(header_tr)
    cells = endfor_tr.findall(f".//{qn('w:tc')}")
    for ci, tc in enumerate(cells):
        for p in tc.findall(f".//{qn('w:p')}"):
            for r in list(p.findall(qn("w:r"))):
                p.remove(r)
        if ci == 0:
            p = tc.findall(f".//{qn('w:p')}")[0]
            make_run(p, endfor_tag)

    # Insert after header
    header_tr.addnext(endfor_tr)
    header_tr.addnext(data_tr)
    header_tr.addnext(for_tr)


def check_sectpr(doc: Document, baseline: dict) -> bool:
    """Verify sectPr baseline is intact after modification."""
    current_orient = ["landscape" if s.orientation == 1 else "portrait" for s in doc.sections]
    return (current_orient == baseline["orientations"] and
            len(doc.sections) == baseline["total_sections"])


def load_baseline() -> dict:
    """Load the sectPr baseline."""
    import json
    baseline_path = Path(__file__).parent.parent / "sectpr_baseline.json"
    with open(baseline_path) as f:
        return json.load(f)


def find_table_by_header(doc: Document, header_text: str, num_cols: int) -> int | None:
    """Find table index by header text and column count."""
    for ti, table in enumerate(doc.tables):
        if len(table.columns) != num_cols:
            continue
        header_row = table.rows[0]
        full_header = " ".join(get_cell_text(c) for c in header_row.cells)
        if header_text.lower() in full_header.lower():
            return ti
    return None
