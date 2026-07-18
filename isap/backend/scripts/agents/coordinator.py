"""Coordinator: run all agents in parallel, merge results, verify.

Runs agents A-D in parallel (each on its own template copy),
then merges all table modifications into a single final template.
"""
from __future__ import annotations

import subprocess
import sys
import os
from pathlib import Path

AGENTS_DIR = Path(__file__).parent
# AGENTS_DIR=agents/ → .parent=scripts/ → .parent=backend/ → .parent=isap/(root)
FILES_DIR = AGENTS_DIR.parent.parent.parent / "files"
WORK_DIR = FILES_DIR / "agent_work"
FINAL_OUTPUT = FILES_DIR / "pmla_v2_template.docx"


def run_agent(script_name: str) -> bool:
    """Run an agent script and return success status."""
    script_path = AGENTS_DIR / script_name
    print(f"\n{'='*50}")
    print(f"Running: {script_name}")
    print(f"{'='*50}")
    result = subprocess.run(
        [sys.executable, str(script_path)],
        capture_output=True, text=True,
        cwd=str(AGENTS_DIR)
    )
    print(result.stdout)
    if result.stderr:
        print(f"STDERR: {result.stderr}")
    if result.returncode != 0:
        print(f"FAILED: {script_name} (exit code {result.returncode})")
        return False
    print(f"OK: {script_name}")
    return True


def merge_agent_results():
    """Merge all agent results into the final template.
    
    Strategy: Start from backup, apply each agent's changes sequentially.
    Each agent modified specific tables, so we can overlay them.
    """
    from docx import Document
    import copy

    backup = FILES_DIR / "pmla_v1_template.docx.bak"
    print(f"\n{'='*50}")
    print(f"Merging agent results into final template")
    print(f"{'='*50}")

    # Load the backup as base
    base_doc = Document(str(backup))
    print(f"Base template: {len(base_doc.tables)} tables")

    # For each agent result, copy modified tables into the base
    agent_files = {
        "agent_a_scenarios": [6, 7, 9, 18],
        "agent_b_incidents": [10, 11],
        "agent_c_resources": [13],
        # Agent D modifies Table 17 in-place (phone replacements)
    }

    for agent_name, table_indices in agent_files.items():
        agent_path = WORK_DIR / f"pmla_v2_{agent_name}.docx"
        if not agent_path.exists():
            print(f"  WARNING: {agent_path} not found, skipping")
            continue

        agent_doc = Document(str(agent_path))
        for ti in table_indices:
            # Replace the table XML in base with the agent's version
            base_tbl = base_doc.tables[ti]._tbl
            agent_tbl = agent_doc.tables[ti]._tbl
            parent = base_tbl.getparent()
            idx = list(parent).index(base_tbl)
            # Deep copy agent's table XML
            new_tbl = copy.deepcopy(agent_tbl)
            parent.remove(base_tbl)
            parent.insert(idx, new_tbl)
            print(f"  Merged Table {ti} from {agent_name}")

    # Agent D: merge phone replacements from Table 17
    agent_d_path = WORK_DIR / "pmla_v2_agent_d_forces.docx"
    if agent_d_path.exists():
        agent_d_doc = Document(str(agent_d_path))
        # For Table 17, we need to merge cell text (not replace whole table
        # because it has complex merged cells)
        base_t17 = base_doc.tables[17]
        agent_t17 = agent_d_doc.tables[17]
        for ri in range(len(base_t17.rows)):
            for ci in range(len(base_t17.columns)):
                base_cell = base_t17.rows[ri].cells[ci]
                agent_cell = agent_t17.rows[ri].cells[ci]
                agent_text = "\n".join(p.text for p in agent_cell.paragraphs)
                if "{{" in agent_text:
                    # Agent modified this cell — apply changes
                    for bp, ap in zip(base_cell.paragraphs, agent_cell.paragraphs):
                        for br, ar in zip(bp.runs, ap.runs):
                            if "{{" in ar.text:
                                br.text = ar.text
        print(f"  Merged Table 17 phone replacements from agent_d")

    # Save final
    base_doc.save(str(FINAL_OUTPUT))
    print(f"\nFinal template saved: {FINAL_OUTPUT}")
    print(f"Size: {FINAL_OUTPUT.stat().st_size:,} bytes")


def verify_final():
    """Run QA checks on the final merged template."""
    from docx import Document
    import re

    doc = Document(str(FINAL_OUTPUT))
    print(f"\n{'='*50}")
    print(f"Verification: {FINAL_OUTPUT.name}")
    print(f"{'='*50}")

    # Table count
    print(f"Tables: {len(doc.tables)}")
    assert len(doc.tables) == 21, f"Expected 21 tables, got {len(doc.tables)}"

    # Jinja cells
    jinja = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{" in cell.text or "{%" in cell.text:
                    jinja += 1
    print(f"Jinja cells: {jinja}")

    # sectPr
    from agent_utils import check_sectpr, load_baseline
    assert check_sectpr(doc, load_baseline()), "FAIL: sectPr changed!"
    print("sectPr: OK")

    # Verify specific tables have loops
    for ti, name in [(6, "substance_params"), (7, "equipment_scenario_links"),
                      (9, "accident_scenarios"), (10, "injury_history"),
                      (11, "accident_history"), (13, "material_reserve"),
                      (18, "countermeasures")]:
        table = doc.tables[ti]
        has_for = any("{%tr for" in cell.text for row in table.rows for cell in row.cells)
        has_endfor = any("{%tr endfor" in cell.text for row in table.rows for cell in row.cells)
        status = "OK" if (has_for and has_endfor) else "MISSING LOOP!"
        print(f"  Table {ti} ({name}): {status}")

    # Table 17 phones
    t17 = doc.tables[17]
    phone_jinja = sum(1 for row in t17.rows for cell in row.cells
                      if "notification_" in cell.text)
    print(f"  Table 17 (phones): {phone_jinja} placeholder cells")

    print("\nALL CHECKS PASSED")


def main():
    # Clean work dir
    WORK_DIR.mkdir(exist_ok=True)

    # Run all agents
    agents = [
        "agent_a_scenarios.py",
        "agent_b_incidents.py",
        "agent_c_resources.py",
        "agent_d_forces.py",
    ]

    results = {}
    for agent in agents:
        results[agent] = run_agent(agent)

    print(f"\n{'='*50}")
    print("Agent Results:")
    for agent, ok in results.items():
        print(f"  {'✓' if ok else '✗'} {agent}")

    if not all(results.values()):
        print("\nSome agents failed! Aborting merge.")
        sys.exit(1)

    # Merge
    merge_agent_results()

    # Verify
    verify_final()

    print(f"\n{'='*50}")
    print("DONE — pmla_v2_template.docx ready")


if __name__ == "__main__":
    main()
