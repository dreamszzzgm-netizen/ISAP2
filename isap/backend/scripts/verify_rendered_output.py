"""Verify rendered PMLA v2 output."""
from docx import Document

for label, path in [
    ("Full render", r"D:\Project ISAP\isap\isap\files\pmla_v2_rendered_test.docx"),
]:
    print(f"\n=== {label} ===")
    doc = Document(path)
    
    # Check tables
    print(f"Tables: {len(doc.tables)}")
    for ti, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if rows > 0 else 0
        # Check for Jinja artifacts
        has_jinja = False
        for row in table.rows:
            for cell in row.cells:
                if "{{" in cell.text or "{%" in cell.text:
                    has_jinja = True
                    break
            if has_jinja:
                break
        status = "JINJA!" if has_jinja else "OK"
        print(f"  Table {ti}: {rows}x{cols} [{status}]")
    
    # Count Jinja artifacts
    artifacts = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{" in cell.text or "{%" in cell.text:
                    artifacts += 1
    print(f"\nTotal Jinja artifacts: {artifacts}")
    
    # Check some key values
    print("\nSample rendered values:")
    t4 = doc.tables[4]  # Organization info
    print(f"  Table 4 Row 2: {t4.rows[2].cells[1].text[:80]}")
    
    t5 = doc.tables[5]  # Equipment
    print(f"  Table 5 data rows: {len(t5.rows) - 2}")  # minus header + number row
    
    t6 = doc.tables[6]  # Substance params
    print(f"  Table 6 data rows: {len(t6.rows) - 1}")
    
    t9 = doc.tables[9]  # Scenarios
    print(f"  Table 9 data rows: {len(t9.rows) - 1}")
    
    t13 = doc.tables[13]  # Material reserve
    print(f"  Table 13 data rows: {len(t13.rows) - 1}")
    
    t18 = doc.tables[18]  # Countermeasures
    print(f"  Table 18 data rows: {len(t18.rows) - 1}")
