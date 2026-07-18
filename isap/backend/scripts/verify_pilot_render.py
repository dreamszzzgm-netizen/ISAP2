"""Verify rendered pilot PMLA v2 DOCX: content, jinja markers, phones."""
import io
import re
import zipfile
from pathlib import Path
from docx import Document

docx_path = Path('/app/pmla_v2_pilot_rendered.docx')
docx_bytes = docx_path.read_bytes()
print(f'File size: {len(docx_bytes)} bytes')
print(f'Valid ZIP: {zipfile.is_zipfile(io.BytesIO(docx_bytes))}')

with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
    names = zf.namelist()
print(f'word/document.xml present: {"word/document.xml" in names}')

doc = Document(io.BytesIO(docx_bytes))
paragraphs = [p.text for p in doc.paragraphs if p.text]
text = '\n'.join(paragraphs)
table_texts = []
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text:
                table_texts.append(cell.text)
all_text = text + '\n' + '\n'.join(table_texts)

jinja_markers = re.findall(r'\{\{[^}]*\}\}|\{%[^%]*%\}', all_text)
print(f'\nJinja markers remaining: {len(jinja_markers)}')
for m in jinja_markers[:5]:
    print(f'  {m}')

print('\n=== CONTENT CHECKS ===')
checks = [
    ('Org name', 'ТехноГазСервис'),
    ('Facility', 'Сеть газопотребления'),
    ('development_year', '2026'),
    ('insurance_amount', '10000000'),
]
for name, val in checks:
    found = val in all_text
    mark = 'OK' if found else 'MISSING'
    print(f'  {name} ({val!r}): {mark}')

print('\n=== SCENARIOS ===')
for sc in ['Утечка природного газа', 'Разрыв газопровода', 'Загазованность']:
    found = sc in all_text
    print(f'  {sc!r}: {"OK" if found else "MISSING"}')

print('\n=== EQUIPMENT ===')
for eq in ['ГРПШ-1', 'ГРУ-1', 'газопровод', 'котёл']:
    found = eq in all_text
    print(f'  {eq!r}: {"OK" if found else "MISSING"}')

print('\n=== TEST PHONES (should be present) ===')
for ph in ['+7 495', '+7 499', '112', '101']:
    found = ph in all_text
    print(f'  {ph!r}: {"OK" if found else "not found"}')

print('\n=== OLD HARDCODED PHONES (should be absent) ===')
old_phones = [
    '+7 (903) 495-75-57', '+7 (8663) 04-14-91', '+7 (86630) 4-18-68',
    '+7 (86630) 4-27-70', '+7 (928) 307-04-62', '+7 (8662) 39-99-99',
    '+7 928 709-95-15', '+7 (903) 491-85-75',
]
present_old = []
for ph in old_phones:
    if ph in all_text:
        present_old.append(ph)
        print(f'  {ph!r}: STILL PRESENT')
    else:
        print(f'  {ph!r}: removed')

print(f'\nSummary: {len(present_old)} old phones still in DOCX')
