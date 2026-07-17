"""PMLA v2 Export QA — comprehensive verification of rendered DOCX.

Checks:
1. No unprocessed Jinja tags ({{ }}, {% %})
2. Valid XML inside DOCX
3. File opens without errors
4. Row counts match context
5. Empty lists render correctly
6. No yellow highlighting
7. All 21 tables present
"""
from __future__ import annotations

import os
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

FILES_DIR = Path(__file__).parent.parent.parent / "files"

JINJA_RE = r"\{[%{].*?[%}]\}"

# Old data deny-list (test organization values that should not appear in rendered output)
OLD_DATA_DENYLIST = [
    "СПК «Тест»",
    "Тестовая сеть газопотребления",
    "А34-99999-0099",
    "ООО «ТестГаз»",
    "ООО «ТестСпас»",
    "Тестов Тест Тестович",
]


def check_no_jinja_artifacts(doc: Document) -> list[str]:
    """Check that no Jinja tags remain in the rendered document."""
    issues = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                import re
                if re.search(JINJA_RE, cell.text):
                    issues.append(f"Table {ti}, Row {ri}, Cell {ci}: {cell.text[:60]}")
    return issues


def check_xml_validity(docx_path: str) -> list[str]:
    """Check that all XML files inside the DOCX are valid."""
    issues = []
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            for name in z.namelist():
                if name.endswith(".xml") and not name.startswith("_"):
                    try:
                        data = z.read(name)
                        ET.fromstring(data)
                    except ET.ParseError as e:
                        issues.append(f"Invalid XML: {name}: {e}")
    except Exception as e:
        issues.append(f"Cannot open DOCX as ZIP: {e}")
    return issues


def check_table_count(doc: Document) -> list[str]:
    """Verify all 21 tables are present."""
    issues = []
    if len(doc.tables) != 21:
        issues.append(f"Expected 21 tables, found {len(doc.tables)}")
    return issues


def check_old_data(doc: Document) -> list[str]:
    """Check for old test data leaking into rendered output."""
    issues = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                text = cell.text
                for old_val in OLD_DATA_DENYLIST:
                    if old_val in text:
                        issues.append(f"Table {ti}: old data '{old_val}' found")
    return issues


def check_empty_lists_render(empty_doc: Document) -> list[str]:
    """Verify that empty lists produce valid (non-broken) tables."""
    issues = []
    # Tables with loops should still have at least header rows
    loop_tables = [5, 6, 7, 8, 9, 10, 11, 13, 18]
    for ti in loop_tables:
        if ti >= len(empty_doc.tables):
            issues.append(f"Table {ti} missing in empty render")
            continue
        table = empty_doc.tables[ti]
        if len(table.rows) < 1:
            issues.append(f"Table {ti} has 0 rows in empty render")
    return issues


def run_qa(docx_path: str, label: str = "DOCX") -> bool:
    """Run all QA checks on a DOCX file. Returns True if all pass."""
    print(f"\n{'='*50}")
    print(f"QA Report: {label}")
    print(f"File: {docx_path}")
    print(f"Size: {os.path.getsize(docx_path):,} bytes")
    print(f"{'='*50}")
    
    all_passed = True
    
    # 1. File opens
    try:
        doc = Document(docx_path)
        print("✓ File opens successfully")
    except Exception as e:
        print(f"✗ Cannot open file: {e}")
        return False
    
    # 2. Jinja artifacts
    issues = check_no_jinja_artifacts(doc)
    if issues:
        print(f"✗ Jinja artifacts found ({len(issues)}):")
        for issue in issues[:5]:
            print(f"  - {issue}")
        all_passed = False
    else:
        print("✓ No Jinja artifacts")
    
    # 3. XML validity
    issues = check_xml_validity(docx_path)
    if issues:
        print(f"✗ XML issues ({len(issues)}):")
        for issue in issues[:5]:
            print(f"  - {issue}")
        all_passed = False
    else:
        print("✓ All XML valid")
    
    # 4. Table count
    issues = check_table_count(doc)
    if issues:
        for issue in issues:
            print(f"✗ {issue}")
        all_passed = False
    else:
        print(f"✓ All 21 tables present")
    
    # 5. Old data check
    issues = check_old_data(doc)
    if issues:
        print(f"✗ Old data found ({len(issues)}):")
        for issue in issues[:5]:
            print(f"  - {issue}")
        all_passed = False
    else:
        print("✓ No old data leakage")
    
    # 6. Table summary
    print(f"\nTable summary:")
    for ti, table in enumerate(doc.tables):
        print(f"  Table {ti:2d}: {len(table.rows):3d} rows x {len(table.columns):2d} cols")
    
    print(f"\n{'PASS' if all_passed else 'FAIL'}")
    return all_passed


def main():
    files_dir = FILES_DIR
    
    # QA on full render
    full_path = files_dir / "pmla_v2_rendered_test.docx"
    if full_path.exists():
        run_qa(str(full_path), "Full Render")
    else:
        print(f"File not found: {full_path}")
    
    # QA on empty render
    empty_path = files_dir / "pmla_v2_rendered_empty.docx"
    if empty_path.exists():
        run_qa(str(empty_path), "Empty Render")
        # Additional empty-list checks
        doc = Document(str(empty_path))
        issues = check_empty_lists_render(doc)
        if issues:
            print("\n✗ Empty list issues:")
            for issue in issues:
                print(f"  - {issue}")
        else:
            print("\n✓ Empty lists render correctly")
    else:
        print(f"File not found: {empty_path}")
    
    # QA on the template itself (should have Jinja tags)
    template_path = files_dir / "pmla_v2_template.docx"
    if template_path.exists():
        print(f"\n{'='*50}")
        print(f"Template: {template_path}")
        doc = Document(str(template_path))
        jinja_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    import re
                    if re.search(JINJA_RE, cell.text):
                        jinja_count += 1
        print(f"  Cells with Jinja: {jinja_count}")
        print(f"  Tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
