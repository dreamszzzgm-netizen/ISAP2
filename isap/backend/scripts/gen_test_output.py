"""Генерация тестового DOCX-документа с обогащёнными справочниками."""
import asyncio
import sys
sys.path.insert(0, ".")

from src.application.engines.base import DocumentContext
from src.application.engines.data_engine import DataEngine
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

engine = DataEngine()

ctx = DocumentContext(
    organization={
        "name": 'ООО «Газпром межрегионгаз Тюмень»',
        "inn": "7736050003",
        "ogrn": "1027700196362",
        "address": "г. Тюмень, ул. Республики, д. 64",
        "phone": "+7 (3452) 53-30-00",
        "email": "info@mg-tyumen.ru",
    },
    facility={
        "name": "Сеть газопотребления города Тюмень",
        "reg_number": "А34-00000-ГТ",
        "hazard_class": "3",
        "facility_type": "Сеть газопотребления",
        "address": "г. Тюмень, ул. Мира, д. 15",
    },
    equipment=[
        {"name": "ГРПШ №1", "equipment_type": "Шкафный ГРП", "serial_number": "GRPSH-2019-001", "manufacture_year": 2019},
        {"name": "Газопровод Ду150", "equipment_type": "Газопровод", "serial_number": "GP-150-2018", "manufacture_year": 2018},
        {"name": "Регулятор РДУК-50", "equipment_type": "Регулятор давления", "serial_number": "RDUK-50-001", "manufacture_year": 2021},
        {"name": "Кран шаровой Ду150", "equipment_type": "Запорная арматура", "serial_number": "KV-150-01", "manufacture_year": 2020},
        {"name": "Сигнализатор газовый СГМ-1", "equipment_type": "Система газоанализа", "serial_number": "SGM-1-042", "manufacture_year": 2022},
    ],
    substances=[
        {"name": "Природный газ", "quantity_kg": 2500, "cas_number": "74-82-8", "threshold_quantity_kg": 1000},
    ],
    persons=[
        {"full_name": "Сидоров А.В.", "position": "Председатель", "role": "chairman", "phone": "+7-999-111-11-11"},
        {"full_name": "Козлов Д.М.", "position": "Диспетчер", "role": "dispatcher", "phone": "+7-999-222-22-22"},
    ],
    year=2026,
)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# Титульная
title = doc.add_heading("", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ПЛАН ЛОКАЛИЗАЦИИ И ЛИКВИДАЦИИ АВАРИЙ")
run.font.size = Pt(16)
run.font.name = "Times New Roman"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
fac_name = ctx.facility["name"]
run = subtitle.add_run(f"Объект: {fac_name}")
run.font.size = Pt(14)
run.font.name = "Times New Roman"

doc.add_paragraph()

# Разделы
sections = [
    ("section_1", "1. Характеристика объекта"),
    ("section_3", "3. Характеристика аварийности"),
    ("section_4", "4. Силы и средства"),
    ("section_8", "8. Управление, связь, оповещение"),
]

for sec_id, heading_text in sections:
    result = asyncio.run(engine.generate(sec_id, {"title": heading_text}, ctx))
    doc.add_heading(heading_text, level=1)

    for block in result.blocks:
        if block.__class__.__name__ == "TableBlock":
            table = doc.add_table(rows=len(block.rows) + 1, cols=len(block.headers))
            table.style = "Table Grid"

            for j, header in enumerate(block.headers):
                cell = table.rows[0].cells[j]
                cell.text = header
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
                        r.font.name = "Times New Roman"

            for i, row in enumerate(block.rows):
                for j, cell_text in enumerate(row):
                    cell = table.rows[i + 1].cells[j]
                    cell.text = str(cell_text)
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)
                            r.font.name = "Times New Roman"

            if block.caption:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(block.caption)
                r.font.size = Pt(10)
                r.font.name = "Times New Roman"
                r.italic = True

        elif block.__class__.__name__ == "ParagraphBlock":
            p = doc.add_paragraph(block.text)
            for r in p.runs:
                r.font.size = Pt(12)
                r.font.name = "Times New Roman"

    doc.add_paragraph()

output_path = "test_reference_output.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
