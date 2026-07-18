"""Modify pmla_v2_template.docx — add Jinja2 placeholders for tables 6-11, 13, 17, 18.

Uses direct XML manipulation to write correct docxtpl syntax.
"""
from __future__ import annotations

import copy
import os
from lxml import etree

from docx import Document
from docx.oxml.ns import qn

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "..", "files", "pmla_v2_template.docx"
)


def make_run_with_text(parent, text: str):
    """Create a w:r element with a w:t child containing the given text."""
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r = etree.SubElement(parent, qn("w:r"))
    # Copy run properties from existing run if any
    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    t.set(qn("xml:space"), "preserve")
    return r


def clear_cell_keep_first(cell):
    """Clear all runs from all paragraphs in a cell, keep paragraph structure."""
    for para in cell.paragraphs:
        for run in list(para.runs):
            para._element.remove(run._element)


def set_cell_first_paragraph(cell, text: str):
    """Set text on the first paragraph of a cell."""
    if cell.paragraphs:
        para = cell.paragraphs[0]
        make_run_with_text(para._element, text)


def insert_loop_rows(table, header_row_idx: int, for_tag: str, field_tags: list[str], endfor_tag: str):
    """Insert Jinja2 for/data/endfor rows after the header row.
    
    Args:
        table: The table
        header_row_idx: Index of the last header row (loop rows go after this)
        for_tag: The {%tr for ... %} tag text
        field_tags: List of {{ var.field }} texts for each column
        endfor_tag: The {%tr endfor %} tag text
    """
    header_tr = table.rows[header_row_idx]._tr
    
    # Create for row
    for_tr = copy.deepcopy(header_tr)
    cells = for_tr.findall(f".//{qn('w:tc')}")
    for ci, tc in enumerate(cells):
        for p in tc.findall(f".//{qn('w:p')}"):
            for r in list(p.findall(qn("w:r"))):
                p.remove(r)
        if ci == 0:
            p = tc.findall(f".//{qn('w:p')}")[0]
            make_run_with_text(p, for_tag)
    
    # Create data row
    data_tr = copy.deepcopy(header_tr)
    cells = data_tr.findall(f".//{qn('w:tc')}")
    for ci, tc in enumerate(cells):
        for p in tc.findall(f".//{qn('w:p')}"):
            for r in list(p.findall(qn("w:r"))):
                p.remove(r)
        if ci < len(field_tags):
            p = tc.findall(f".//{qn('w:p')}")[0]
            make_run_with_text(p, field_tags[ci])
    
    # Create endfor row
    endfor_tr = copy.deepcopy(header_tr)
    cells = endfor_tr.findall(f".//{qn('w:tc')}")
    for ci, tc in enumerate(cells):
        for p in tc.findall(f".//{qn('w:p')}"):
            for r in list(p.findall(qn("w:r"))):
                p.remove(r)
        if ci == 0:
            p = tc.findall(f".//{qn('w:p')}")[0]
            make_run_with_text(p, endfor_tag)
    
    # Insert after header
    header_tr.addnext(endfor_tr)
    header_tr.addnext(data_tr)
    header_tr.addnext(for_tr)


def remove_rows(table, start_idx: int, end_idx: int):
    """Remove rows from start_idx to end_idx (inclusive)."""
    for i in range(end_idx, start_idx - 1, -1):
        tr = table.rows[i]._tr
        tr.getparent().remove(tr)


def get_cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs)


# =========================================================================
# Table modifications
# =========================================================================

def modify_table_6(doc):
    """Table 6: Substance parameters."""
    table = doc.tables[6]
    print(f"  Table 6: {len(table.rows)} rows")
    remove_rows(table, 1, 8)  # Remove 8 data rows
    insert_loop_rows(
        table, 0,
        for_tag="{%tr for param in substance_params %}",
        field_tags=["{{ param.parameter }}", "{{ param.value }}"],
        endfor_tag="{%tr endfor %}",
    )
    print("  -> substance_params loop added")


def modify_table_7(doc):
    """Table 7: Equipment-scenario links."""
    table = doc.tables[7]
    print(f"  Table 7: {len(table.rows)} rows")
    remove_rows(table, 1, 3)  # Remove 3 data rows
    insert_loop_rows(
        table, 0,
        for_tag="{%tr for link in equipment_scenario_links %}",
        field_tags=[
            "{{ loop.index }}",
            "{{ link.equipment_name }}",
            "{{ link.scenario_codes }}",
            "{{ link.description }}",
            "{{ link.damaging_factors }}",
        ],
        endfor_tag="{%tr endfor %}",
    )
    print("  -> equipment_scenario_links loop added")


def modify_table_8(doc):
    """Table 8: Equipment defects."""
    table = doc.tables[8]
    print(f"  Table 8: {len(table.rows)} rows")
    remove_rows(table, 1, 13)  # Remove 13 data rows
    insert_loop_rows(
        table, 0,
        for_tag="{%tr for defect in equipment_defects %}",
        field_tags=[
            "{{ loop.index }}",
            "{{ defect.equipment_name }}",
            "{{ defect.defect }}",
            "{{ defect.cause }}",
            "{{ defect.source }}",
            "{{ defect.scenario }}",
        ],
        endfor_tag="{%tr endfor %}",
    )
    print("  -> equipment_defects loop added")


def modify_table_9(doc):
    """Table 9: Accident scenarios."""
    table = doc.tables[9]
    print(f"  Table 9: {len(table.rows)} rows")
    remove_rows(table, 1, 6)  # Remove 6 data rows
    insert_loop_rows(
        table, 0,
        for_tag="{%tr for scenario in accident_scenarios %}",
        field_tags=[
            "{{ scenario.code }}",
            "{{ scenario.name }}",
            "{{ scenario.source }}",
            "{{ scenario.preconditions }}",
            "{{ scenario.signs }}",
            "{{ scenario.damaging_factors }}",
        ],
        endfor_tag="{%tr endfor %}",
    )
    print("  -> accident_scenarios loop added")


def modify_table_10(doc):
    """Table 10: Injury history."""
    table = doc.tables[10]
    print(f"  Table 10: {len(table.rows)} rows")
    remove_rows(table, 1, 5)  # Remove 5 data rows
    insert_loop_rows(
        table, 0,
        for_tag="{%tr for injury in injury_history %}",
        field_tags=[
            "{{ injury.year }}",
            "{{ injury.incident_number }}",
            "{{ injury.date }}",
            "{{ injury.character }}",
            "{{ injury.trauma }}",
            "{{ injury.consequences }}",
            "{{ injury.measures_percent }}",
        ],
        endfor_tag="{%tr endfor %}",
    )
    print("  -> injury_history loop added")


def modify_table_11(doc):
    """Table 11: Accident history."""
    table = doc.tables[11]
    print(f"  Table 11: {len(table.rows)} rows")
    remove_rows(table, 1, 5)  # Remove 5 data rows
    insert_loop_rows(
        table, 0,
        for_tag="{%tr for accident in accident_history %}",
        field_tags=[
            "{{ accident.year }}",
            "{{ accident.incident_number }}",
            "{{ accident.date }}",
            "{{ accident.character }}",
            "{{ accident.trauma }}",
            "{{ accident.consequences }}",
            "{{ accident.measures_percent }}",
        ],
        endfor_tag="{%tr endfor %}",
    )
    print("  -> accident_history loop added")


def modify_table_13(doc):
    """Table 13: Material reserve with group headers."""
    table = doc.tables[13]
    print(f"  Table 13: {len(table.rows)} rows")
    remove_rows(table, 1, 21)  # Remove all data rows (including group headers)
    insert_loop_rows(
        table, 0,
        for_tag="{%tr for item in material_reserve %}",
        field_tags=[
            "{{ loop.index if not item.is_group_header else '' }}",
            "{{ item.group_name if item.is_group_header else item.name }}",
            "{{ '' if item.is_group_header else item.quantity }}",
            "{{ '' if item.is_group_header else item.location }}",
        ],
        endfor_tag="{%tr endfor %}",
    )
    print("  -> material_reserve loop added")


def modify_table_17(doc):
    """Table 17: Notification list — parameterize phone numbers."""
    table = doc.tables[17]
    print(f"  Table 17: {len(table.rows)} rows")
    
    # Phone replacements (old hardcoded -> new Jinja2 placeholder)
    phone_map = {
        "+7 928 709-95-15": "{{ notification_chairman_phone | default('', true) }}",
        "+ 7 906 881-07-07": "{{ notification_deputy_phone | default('', true) }}",
        "112 | +7 (86630) 4-00-06": "{{ notification_edds_phone | default('112', true) }}",
        "+7 (903) 495-75-57 | +7 (903) 491-85-75": "{{ notification_pasf_phone | default('', true) }}",
        "+7(8663) 04-14-91": "{{ notification_fire_phone | default('', true) }}",
        "112/03/103": "{{ notification_ambulance_phone | default('112/03/103', true) }}",
        "+7 (86630) 4-18-68; | 4-18-53": "{{ notification_gas_phone | default('', true) }}",
        "+7  (86630) 4-27-70": "{{ notification_electric_phone | default('', true) }}",
        "+7 (8662) 39-99-99": "{{ notification_mchs_phone | default('', true) }}",
        "+7 (928) 307-04-62 | +7 (8793)-34-64-24 | +7 (8662) 91-99-33": "{{ notification_rostechnadzor_phone | default('', true) }}",
        "+7 (86630) 7-63-99": "{{ notification_admin_phone | default('', true) }}",
    }
    
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in list(para.runs):
                    original = run.text
                    modified = original
                    for old, new in phone_map.items():
                        if old in modified:
                            modified = modified.replace(old, new)
                    if modified != original:
                        run.text = modified
    
    print("  -> Phone numbers parameterized")


def modify_table_18(doc):
    """Table 18: Countermeasures."""
    table = doc.tables[18]
    print(f"  Table 18: {len(table.rows)} rows")
    remove_rows(table, 1, 5)  # Remove 5 data rows
    insert_loop_rows(
        table, 0,
        for_tag="{%tr for cm in countermeasures %}",
        field_tags=[
            "{{ cm.scenario_label }}",
            "{{ cm.signs }}",
            "{{ cm.protection }}",
            "{{ cm.technical_means }}",
            "{{ cm.executors }}",
        ],
        endfor_tag="{%tr endfor %}",
    )
    print("  -> countermeasures loop added")


# =========================================================================
# Main
# =========================================================================

def main():
    print(f"Loading: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)
    print(f"Tables: {len(doc.tables)}")
    
    for name, fn in [
        ("Table 6 (Substance Params)", modify_table_6),
        ("Table 7 (Equipment-Scenario Links)", modify_table_7),
        ("Table 8 (Equipment Defects)", modify_table_8),
        ("Table 9 (Accident Scenarios)", modify_table_9),
        ("Table 10 (Injury History)", modify_table_10),
        ("Table 11 (Accident History)", modify_table_11),
        ("Table 13 (Material Reserve)", modify_table_13),
        ("Table 17 (Notification List)", modify_table_17),
        ("Table 18 (Countermeasures)", modify_table_18),
    ]:
        print(f"\n--- {name} ---")
        fn(doc)
    
    doc.save(TEMPLATE_PATH)
    print(f"\nSaved: {TEMPLATE_PATH}")
    
    # Quick verify
    doc2 = Document(TEMPLATE_PATH)
    jinja_count = 0
    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{" in cell.text or "{%" in cell.text:
                    jinja_count += 1
    print(f"Cells with Jinja: {jinja_count}")
    print("Done!")


if __name__ == "__main__":
    main()
