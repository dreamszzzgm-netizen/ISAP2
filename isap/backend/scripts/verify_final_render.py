"""Re-render with new agent-merged template and verify."""
import sys, os
sys.path.insert(0, r"D:\Project ISAP\isap\isap\backend")
from src.infrastructure.export.pmla_template_renderer import PmlaTemplateRenderer
from docx import Document

ctx = {
    'organization_full_name': 'ООО ТестПром', 'organization_short_name': 'ООО ТестПром',
    'legal_address': '123456, г. Москва', 'inn': '7701234567', 'ogrn': '1027700000123',
    'phone': '+7 (495) 123-45-67', 'email': 'info@test.ru',
    'director_position_fullname': 'Генеральный директор Иванов И.И.',
    'director_initials_surname': 'И.И. Иванов', 'director_initials_surname_full': 'Иванов Иван Иванович',
    'deputy_chairman_fullname': 'Петров П.П.', 'main_activity_description': '01.11',
    'facility_name': 'Сеть газопотребления', 'facility_reg_number': 'А34-99999-0001',
    'facility_location': 'Московская область', 'hazard_class': 'III',
    'hazardous_substances_info': 'Природный газ', 'hazard_characteristics_116fz': 'Газы',
    'contractor_organization_name': 'ООО Спас', 'contractor_organization_short_name': 'Спас',
    'contractor_agreement_date': '01.01.2026', 'gas_supplier_name': 'Газпром',
    'total_hazardous_substance_quantity': '0.5',
    'equipment_list': [
        {'location': 'Площадка ГРП', 'hazard_characteristic': 'Газ', 'device_name': 'ГРПШ-1', 'specifications': '0.6 МПа', 'process_codes': '2.1'},
        {'location': 'Газопровод ВД', 'hazard_characteristic': 'Газ', 'device_name': 'Труба Ø57', 'specifications': 'L=13.4м', 'process_codes': '2.1'},
        {'location': 'Котельная', 'hazard_characteristic': 'Газ', 'device_name': 'Котёл КГВМ-50', 'specifications': 'Q=50 Гкал/ч', 'process_codes': '2.2'},
    ],
    'substance_params': [
        {'parameter': 'Класс опасности', 'value': '4 (малоопасный)'},
        {'parameter': 'ПДК', 'value': '300 мг/м³'},
    ],
    'equipment_scenario_links': [
        {'equipment_name': 'ГРПШ-1', 'scenario_codes': 'С-1, С-2, С-4', 'description': 'утечка', 'damaging_factors': 'газ'},
    ],
    'equipment_defects': [
        {'equipment_name': 'ГРПШ-1', 'defect': 'Разгерметизация', 'cause': 'Износ', 'source': 'Фланец', 'scenario': 'С-1'},
    ],
    'accident_scenarios': [
        {'code': 'С-1', 'name': 'Выброс газа', 'source': 'ГРПШ', 'preconditions': 'Утечка', 'signs': 'Запах', 'damaging_factors': 'Токсическое'},
    ],
    'injury_history': [], 'accident_history': [],
    'material_reserve': [
        {'is_group_header': True, 'group_name': 'СИЗ'},
        {'name': 'Противогаз', 'quantity': '4 шт.', 'location': 'Шкаф'},
    ],
    'countermeasures': [
        {'scenario_label': 'С-1', 'signs': 'Запах', 'protection': 'Отсечь', 'technical_means': 'Краны', 'executors': 'Диспетчер'},
    ],
    'notification_chairman_phone': '+7 999 000-00-00', 'notification_deputy_phone': '',
    'notification_edds_phone': '112', 'notification_pasf_phone': '', 'notification_fire_phone': '',
    'notification_ambulance_phone': '103', 'notification_gas_phone': '', 'notification_electric_phone': '',
    'notification_mchs_phone': '', 'notification_rostechnadzor_phone': '', 'notification_admin_phone': '',
}

r = PmlaTemplateRenderer()

# Full render
full = r"D:\Project ISAP\isap\isap\files\pmla_v2_rendered_test.docx"
r.render_to_file(ctx, full)
doc = Document(full)
artifacts = sum(1 for t in doc.tables for row in t.rows for c in row.cells if "{{" in c.text or "{%" in c.text)
print("Full render: {:,} bytes, {} tables, {} artifacts".format(os.path.getsize(full), len(doc.tables), artifacts))

for ti, name in [(6,"substance"), (7,"eq-scenario"), (9,"scenarios"), (10,"injuries"), (11,"accidents"), (13,"reserve"), (18,"countermeasures")]:
    t = doc.tables[ti]
    has_for = any("{%tr for" in c.text for row in t.rows for c in row.cells)
    status = "OK" if has_for else "MISSING"
    print("  Table {} ({}): {} ({} rows)".format(ti, name, status, len(t.rows)))

# Empty render
empty_ctx = {k: ([] if isinstance(v, list) else v) for k, v in ctx.items()}
empty = r"D:\Project ISAP\isap\isap\files\pmla_v2_rendered_empty.docx"
r.render_to_file(empty_ctx, empty)
doc2 = Document(empty)
artifacts2 = sum(1 for t in doc2.tables for row in t.rows for c in row.cells if "{{" in c.text or "{%" in c.text)
print("Empty render: {:,} bytes, {} artifacts".format(os.path.getsize(empty), artifacts2))
print("ALL OK")
