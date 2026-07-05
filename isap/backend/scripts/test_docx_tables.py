"""Генерация тестового DOCX для проверки таблиц и форматирования."""
import sys
sys.path.insert(0, "/app")

from docx import Document as DocxDocument
from src.application.services.enhanced_generator import EnhancedDocumentGenerator
from unittest.mock import AsyncMock

# Создаём генератор
retriever = AsyncMock()
doc_repo = AsyncMock()
reg_repo = AsyncMock()
reg_repo.session = AsyncMock()
gen = EnhancedDocumentGenerator(
    llm=None, retriever=retriever,
    document_repo=doc_repo, regulatory_repo=reg_repo,
)

# Тестовый контент с markdown-таблицами
content_title = "Тестовый документ ПМЛА"
sections = {
    "1. Характеристика ОПО": (
        "Объект: Компрессорная станция «Приобье»\n"
        "Класс опасности: III\n\n"
        "| № п/п | Наименование | Характеристика | Количество |\n"
        "|-------|-------------|----------------|------------|\n"
        "| 1 | Компрессор ОС-500 | Осевой, 500 кВт | 2 шт |\n"
        "| 2 | Конденсатор КГ-200 | Трубчатый | 4 шт |\n"
        "| 3 | Холодильник ХМ-100 | Воздушный | 1 шт |"
    ),
    "2. Сценарии аварий": (
        "| № сценария | Наименование | Источник | Причины |\n"
        "|-----------|-------------|----------|--------|\n"
        "| С-1 | Выброс газа | Фланцевые соединения | Разгерметизация |\n"
        "| С-2 | Струйное горение | Газопровод | Разрыв трубы |\n"
        "| С-3 | Взрыв ГВС | Территория ГРПШ | Утечка газа |"
    ),
    "3. Состав сил и средств": (
        "| Вид сил и средств | Количество | Место дислокации |\n"
        "|------------------|-----------|------------------|\n"
        "| Пожарная охрана | 1 отделение | п. Приобье |\n"
        "| Медицинская помощь | 1 бригада | п. Приобье |"
    ),
    "4. Контакты": (
        "Ответственное лицо: Иванов И.И.\n"
        "Должность: Начальник ПБ\n"
        "Телефон: +7 (3452) 12-34-56"
    ),
}

metadata = {
    "version": "1.0",
    "generated_at": "2026-07-04",
    "status": "draft",
    "calculation_results": [],
    "validation_issues": [],
}

docx_bytes = gen._build_docx(content_title, sections, metadata)

output_path = "/app/test_output_tables.docx"
with open(output_path, "wb") as f:
    f.write(docx_bytes)

print(f"DOCX saved: {output_path} ({len(docx_bytes)} bytes)")

# Проверяем содержимое
doc = DocxDocument()
from io import BytesIO
doc = DocxDocument(BytesIO(docx_bytes))
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
for i, t in enumerate(doc.tables):
    print(f"  Table {i}: {len(t.rows)}x{len(t.columns)}")
    for j, row in enumerate(t.rows[:2]):
        cells = [c.text[:30] for c in row.cells]
        print(f"    Row {j}: {cells}")
