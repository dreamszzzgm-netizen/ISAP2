"""Modify pmla_v2_template.docx to add Jinja2 placeholders for tables 6-11, 13, 17, 18.

This script modifies the DOCX template in-place, replacing hardcoded data rows
with Jinja2 {%tr for %} loops. Uses python-docx for precise OOXML manipulation.
"""
from __future__ import annotations

import copy
import os
import sys

from docx import Document
from docx.oxml.ns import qn

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "..", "files", "pmla_v2_template.docx"
)
OUTPUT_PATH = TEMPLATE_PATH  # overwrite in place


def get_cell_text(cell) -> str:
    """Get full text from a cell including all paragraphs."""
    return "\n".join(p.text for p in cell.paragraphs)


def set_cell_text(cell, text: str) -> None:
    """Set cell text, preserving the first paragraph's formatting."""
    # Clear all existing paragraphs except the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    # Set text on the first paragraph
    if cell.paragraphs:
        cell.paragraphs[0].text = text
    else:
        # Cell has no paragraphs - add one
        from docx.oxml import OxmlElement
        tc = cell._tc
        p = OxmlElement("w:p")
        tc.append(p)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)


def clone_row(table, row_index: int) -> None:
    """Clone a table row by inserting a copy after the given index."""
    row = table.rows[row_index]
    tr = row._tr
    new_tr = copy.deepcopy(tr)
    # Insert after the current row
    tr.addnext(new_tr)


def replace_table_rows_with_loop(
    table, 
    start_data_row: int, 
    end_data_row: int,  # inclusive
    loop_variable: str,
    field_names: list[str],
    header_rows: int = 0,
) -> None:
    """Replace data rows in a table with Jinja2 {%tr for %} loop.
    
    Args:
        table: The table object
        start_data_row: Index of first data row to replace
        end_data_row: Index of last data row to replace (inclusive)
        loop_variable: Jinja2 variable name (e.g., "param", "defect")
        field_names: List of field names matching column order
        header_rows: Number of header rows to keep (default 0 = all rows before start_data_row are headers)
    """
    # Remove existing data rows (from end to start to preserve indices)
    for i in range(end_data_row, start_data_row - 1, -1):
        tr = table.rows[i]._tr
        tr.getparent().remove(tr)
    
    # Create "for" row
    for_tr = copy.deepcopy(table.rows[start_data_row - 1]._tr) if start_data_row > 0 else copy.deepcopy(table.rows[0]._tr)
    # Clear all cells and set the for tag in the first cell
    cells = for_tr.findall(f".//{qn('w:tc')}")
    for ci, tc in enumerate(cells):
        # Clear paragraph content
        for p in tc.findall(f".//{qn('w:p')}"):
            for r in list(p.findall(qn("w:r"))):
                p.remove(r)
        # Add Jinja tag
        if ci == 0:
            p = tc.find(f".//{qn('w:p')}")
            if p is None:
                p = tc.findall(qn("w:p"))[0] if tc.findall(qn("w:p")) else None
            if p is not None:
                r = p.makeelement(qn("w:r"), {})
                t = r.makeelement(qn("w:t"), {})
                t.text = "{{%tr for {v} in {arr} %}}".format(v=loop_variable, arr=loop_variable + "_array")
                t.set(qn("xml:space"), "preserve")
                r.append(t)
                p.append(r)
    
    # Create data row with field placeholders
    data_tr = copy.deepcopy(for_tr)
    cells = data_tr.findall(f".//{qn('w:tc')}")
    for ci, tc in enumerate(cells):
        for p in tc.findall(f".//{qn('w:p')}"):
            for r in list(p.findall(qn("w:r"))):
                p.remove(r)
        if ci < len(field_names):
            p = tc.findall(qn("w:p"))[0] if tc.findall(qn("w:p")) else None
            if p is not None:
                r = p.makeelement(qn("w:r"), {})
                t = r.makeelement(qn("w:t"), {})
                t.text = "{{{{ {v}.{f} }}}}".format(v=loop_variable, f=field_names[ci])
                t.set(qn("xml:space"), "preserve")
                r.append(t)
                p.append(r)
    
    # Create "endfor" row
    endfor_tr = copy.deepcopy(for_tr)
    cells = endfor_tr.findall(f".//{qn('w:tc')}")
    for ci, tc in enumerate(cells):
        for p in tc.findall(f".//{qn('w:p')}"):
            for r in list(p.findall(qn("w:r"))):
                p.remove(r)
        if ci == 0:
            p = tc.findall(qn("w:p"))[0] if tc.findall(qn("w:p")) else None
            if p is not None:
                r = p.makeelement(qn("w:r"), {})
                t = r.makeelement(qn("w:t"), {})
                t.text = "{{%tr endfor %}}"
                t.set(qn("xml:space"), "preserve")
                r.append(t)
                p.append(r)
    
    # Insert rows after the last header row
    insert_after = table.rows[start_data_row - 1]._tr if start_data_row > 0 else None
    if insert_after is not None:
        insert_after.addnext(endfor_tr)
        insert_after.addnext(data_tr)
        insert_after.addnext(for_tr)
    else:
        # Insert at the beginning of the table
        tbl = table._tbl
        first_tr = tbl.find(qn("w:tr"))
        if first_tr is not None:
            first_tr.addprevious(for_tr)
            first_tr.addprevious(data_tr)
            first_tr.addprevious(endfor_tr)


def modify_table_6(doc: Document) -> None:
    """Table 6: Substance parameters (Параметр | Значение)."""
    table = doc.tables[6]
    print(f"  Table 6: {len(table.rows)} rows, {len(table.columns)} cols")
    # Header: row 0 (Параметр | Значение)
    # Data: rows 1-8 (8 parameters)
    # Replace rows 1-8 with loop
    replace_table_rows_with_loop(
        table,
        start_data_row=1,
        end_data_row=8,
        loop_variable="param",
        field_names=["parameter", "value"],
    )
    print("  -> Added {%tr for param in substance_params %}")


def modify_table_7(doc: Document) -> None:
    """Table 7: Equipment-scenario links."""
    table = doc.tables[7]
    print(f"  Table 7: {len(table.rows)} rows, {len(table.columns)} cols")
    # Header: row 0
    # Data: rows 1-3 (3 equipment-scenario links)
    replace_table_rows_with_loop(
        table,
        start_data_row=1,
        end_data_row=3,
        loop_variable="link",
        field_names=["equipment_name", "scenario_codes", "description", "damaging_factors"],
    )
    print("  -> Added {%tr for link in equipment_scenario_links %}")


def modify_table_8(doc: Document) -> None:
    """Table 8: Equipment defects."""
    table = doc.tables[8]
    print(f"  Table 8: {len(table.rows)} rows, {len(table.columns)} cols")
    # Header: row 0
    # Data: rows 1-13 (13 defects)
    replace_table_rows_with_loop(
        table,
        start_data_row=1,
        end_data_row=13,
        loop_variable="defect",
        field_names=["equipment_name", "defect", "cause", "source", "scenario"],
    )
    print("  -> Added {%tr for defect in equipment_defects %}")


def modify_table_9(doc: Document) -> None:
    """Table 9: Accident scenarios."""
    table = doc.tables[9]
    print(f"  Table 9: {len(table.rows)} rows, {len(table.columns)} cols")
    # Header: row 0
    # Data: rows 1-6 (6 scenarios)
    replace_table_rows_with_loop(
        table,
        start_data_row=1,
        end_data_row=6,
        loop_variable="scenario",
        field_names=["code", "name", "source", "preconditions", "signs", "damaging_factors"],
    )
    print("  -> Added {%tr for scenario in accident_scenarios %}")


def modify_table_10(doc: Document) -> None:
    """Table 10: Injury history."""
    table = doc.tables[10]
    print(f"  Table 10: {len(table.rows)} rows, {len(table.columns)} cols")
    # Header: row 0
    # Data: rows 1-5 (5 years)
    replace_table_rows_with_loop(
        table,
        start_data_row=1,
        end_data_row=5,
        loop_variable="injury",
        field_names=["year", "incident_number", "date", "character", "trauma", "consequences", "measures_percent"],
    )
    print("  -> Added {%tr for injury in injury_history %}")


def modify_table_11(doc: Document) -> None:
    """Table 11: Accident history."""
    table = doc.tables[11]
    print(f"  Table 11: {len(table.rows)} rows, {len(table.columns)} cols")
    # Header: row 0
    # Data: rows 1-5 (5 years)
    replace_table_rows_with_loop(
        table,
        start_data_row=1,
        end_data_row=5,
        loop_variable="accident",
        field_names=["year", "incident_number", "date", "character", "trauma", "consequences", "measures_percent"],
    )
    print("  -> Added {%tr for accident in accident_history %}")


def modify_table_13(doc: Document) -> None:
    """Table 13: Material reserve with group headers."""
    table = doc.tables[13]
    print(f"  Table 13: {len(table.rows)} rows, {len(table.columns)} cols")
    # Header: row 0 (№ п/п | Наименование | Количество | Место расположения)
    # Data: rows 1-21 (group headers + items)
    replace_table_rows_with_loop(
        table,
        start_data_row=1,
        end_data_row=21,
        loop_variable="item",
        field_names=["name", "name", "quantity", "location"],
    )
    print("  -> Added {%tr for item in material_reserve %}")


def modify_table_17(doc: Document) -> None:
    """Table 17: Notification list — parameterize phone numbers.
    
    Table 17 has merged cells in the first 2 columns. Instead of a full loop
    (which would break merged cells), we parameterize individual phone numbers
    and names using scalar placeholders.
    """
    table = doc.tables[17]
    print(f"  Table 17: {len(table.rows)} rows, {len(table.columns)} cols")
    
    # Replace hardcoded phone numbers with placeholders
    phone_replacements = {
        "+7 928 709-95-15": "{{ notification_chairman_phone | default('+7 XXX XXX-XX-XX', true) }}",
        "+ 7 906 881-07-07": "{{ notification_deputy_phone | default('+7 XXX XXX-XX-XX', true) }}",
        "112 | +7 (86630) 4-00-06": "{{ notification_edds_phone | default('112', true) }}",
        "+7 (903) 495-75-57 | +7 (903) 491-85-75": "{{ notification_pasf_phone | default('+7 (XXX) XXX-XX-XX', true) }}",
        "+7(8663) 04-14-91": "{{ notification_fire_phone | default('+7 (XXXX) XX-XX-XX', true) }}",
        "112/03/103": "{{ notification_ambulance_phone | default('112/03/103', true) }}",
        "+7 (86630) 4-18-68; | 4-18-53": "{{ notification_gas_phone | default('+7 (XXXXXX) X-XX-XX', true) }}",
        "+7  (86630) 4-27-70": "{{ notification_electric_phone | default('+7 (XXXXXX) X-XX-XX', true) }}",
        "+7 (8662) 39-99-99": "{{ notification_mchs_phone | default('+7 (XXXX) XX-XX-XX', true) }}",
        "+7 (928) 307-04-62 | +7 (8793)-34-64-24 | +7 (8662) 91-99-33": "{{ notification_rostechnadzor_phone | default('+7 (XXX) XXX-XX-XX', true) }}",
        "+7 (86630) 7-63-99": "{{ notification_admin_phone | default('+7 (XXXXXX) X-XX-XX', true) }}",
    }
    
    # Also replace hardcoded names
    name_replacements = {
        "Тестов Тест Тестович": "{{ director_initials_surname_full }}",
    }
    
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    original = run.text
                    modified = original
                    for old, new in phone_replacements.items():
                        if old in modified:
                            modified = modified.replace(old, new)
                    for old, new in name_replacements.items():
                        if old in modified:
                            modified = modified.replace(old, new)
                    if modified != original:
                        run.text = modified
    
    print("  -> Parameterized phone numbers and names in Table 17")


def modify_table_18(doc: Document) -> None:
    """Table 18: Countermeasures."""
    table = doc.tables[18]
    print(f"  Table 18: {len(table.rows)} rows, {len(table.columns)} cols")
    # Header: row 0
    # Data: rows 1-5 (5 scenarios)
    replace_table_rows_with_loop(
        table,
        start_data_row=1,
        end_data_row=5,
        loop_variable="cm",
        field_names=["scenario_label", "signs", "protection", "technical_means", "executors"],
    )
    print("  -> Added {%tr for cm in countermeasures %}")


def main():
    print(f"Loading template: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)
    
    print(f"\nOriginal tables: {len(doc.tables)}")
    for i, t in enumerate(doc.tables):
        print(f"  Table {i}: {len(t.rows)} rows x {len(t.columns)} cols")
    
    print("\n--- Modifying Table 6 (Substance Parameters) ---")
    modify_table_6(doc)
    
    print("\n--- Modifying Table 7 (Equipment-Scenario Links) ---")
    modify_table_7(doc)
    
    print("\n--- Modifying Table 8 (Equipment Defects) ---")
    modify_table_8(doc)
    
    print("\n--- Modifying Table 9 (Accident Scenarios) ---")
    modify_table_9(doc)
    
    print("\n--- Modifying Table 10 (Injury History) ---")
    modify_table_10(doc)
    
    print("\n--- Modifying Table 11 (Accident History) ---")
    modify_table_11(doc)
    
    print("\n--- Modifying Table 13 (Material Reserve) ---")
    modify_table_13(doc)
    
    print("\n--- Modifying Table 17 (Notification List) ---")
    modify_table_17(doc)
    
    print("\n--- Modifying Table 18 (Countermeasures) ---")
    modify_table_18(doc)
    
    print(f"\nSaving modified template: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    
    # Verify
    doc2 = Document(OUTPUT_PATH)
    print(f"\nVerification — tables in saved file: {len(doc2.tables)}")
    for i, t in enumerate(doc2.tables):
        print(f"  Table {i}: {len(t.rows)} rows x {len(t.columns)} cols")
    
    # Check for Jinja tags
    jinja_count = 0
    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                text = get_cell_text(cell)
                if "{{" in text or "{%" in text:
                    jinja_count += 1
    print(f"\nCells with Jinja tags: {jinja_count}")
    print("Done!")


if __name__ == "__main__":
    main()
