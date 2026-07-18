"""Replace remaining old data with Jinja2 placeholders in paragraphs, tables, and XML."""
import zipfile, shutil, os, re
from docx import Document
from lxml import etree

TEMPLATE = r"D:\Project ISAP\isap\isap\files\pmla_v2_template.docx"

# === PART 1: Fix paragraphs and tables via python-docx ===
doc = Document(TEMPLATE)

# Paragraph replacements: (old_substring, new_placeholder)
para_replacements = [
    ("филиал в Чегемском районе", "филиал в {{ gas_supplier_branch }}"),
]

replaced = 0
for p in doc.paragraphs:
    for old, new in para_replacements:
        if old in p.text:
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    replaced += 1

# Table cell replacements
table_replacements = [
    ("филиал в Чегемском районе", "филиал в {{ gas_supplier_branch }}"),
    ("г. Чегем, Чегемский муниципальный район, КБР", "{{ dislocation_address }}"),
    ("ЕДДС Чегемского района", "{{ edds_name }}"),
    ("Чегемские Районные Электрические Сети", "{{ electric_company }}"),
    ("Местная администрация сельского поселения Чегем Второй", "{{ local_admin }}"),
]

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for old, new in table_replacements:
                if old in cell.text:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                replaced += 1

doc.save(TEMPLATE)
print(f"Part 1: {replaced} paragraph/table replacements")

# === PART 2: Fix remaining in XML (runs split across elements) ===
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def replace_in_wt(content, old, new):
    """Replace text within <w:t> elements."""
    pattern = r'(<w:t[^>]*>)([^<]*?)(' + re.escape(old) + r')([^<]*?)(</w:t>)'
    def replacer(m):
        prefix = m.group(2)
        match_text = m.group(3)
        suffix = m.group(4)
        replacement = match_text.replace(old, new)
        return f'{m.group(1)}{prefix}{replacement}{suffix}{m.group(5)}'
    return re.sub(pattern, replacer, content)

# XML-level replacements for values that might be split across runs
xml_replacements = [
    ("Чегемском районе", "{{ gas_supplier_branch }}"),
    ("Чегемский муниципальный район", "{{ dislocation_district }}"),
    ("Чегемского района", "{{ edds_district }}"),
    ("Чегемские Районные Электрические Сети", "{{ electric_company }}"),
    ("поселения Чегем Второй", "{{ local_admin }}"),
    ("г. Чегем", "{{ settlement_name }}"),
]

with zipfile.ZipFile(TEMPLATE, 'r') as z:
    xml_files = [n for n in z.namelist() if n.endswith('.xml') and not n.startswith('_')]

xml_fixed = 0
for xml_name in xml_files:
    with zipfile.ZipFile(TEMPLATE, 'r') as z:
        data = z.read(xml_name).decode('utf-8')

    original = data
    for old, new in xml_replacements:
        data = replace_in_wt(data, old, new)

    if data != original:
        xml_fixed += 1
        # Save modified XML back
        shutil.copy2(TEMPLATE, TEMPLATE + '.xmlbak')
        with zipfile.ZipFile(TEMPLATE, 'r') as zin:
            with zipfile.ZipFile(TEMPLATE + '.xmltmp', 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename == xml_name:
                        zout.writestr(item, data.encode('utf-8'))
                    else:
                        zout.writestr(item, zin.read(item.filename))
        os.remove(TEMPLATE)
        os.rename(TEMPLATE + '.xmltmp', TEMPLATE)
        print(f"  XML fixed: {xml_name}")

print(f"Part 2: {xml_fixed} XML files modified")
print(f"Saved: {TEMPLATE}")
