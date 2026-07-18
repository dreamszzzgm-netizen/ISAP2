"""Find old data locations in the v2 template."""
from docx import Document

doc = Document(r"D:\Project ISAP\isap\isap\files\pmla_v2_template.docx")

old_markers = ["СПК", "А34-99999-0099", "Чегем"]
found = []

for i, p in enumerate(doc.paragraphs):
    text = p.text
    for old in old_markers:
        if old in text:
            snippet = text[max(0, text.index(old) - 40):text.index(old) + 60]
            found.append(f"  Para {i:3d} [{p.style.name:20s}]: ...{snippet}...")

print(f"Paragraphs with old data: {len(found)}")
for f in found[:20]:
    print(f)
if len(found) > 20:
    print(f"  ... and {len(found) - 20} more")

# Also check tables (non-parameterized)
print("\nTable old data (non-parameterized tables only):")
for ti, table in enumerate(doc.tables):
    if ti in [5, 6, 7, 8, 9, 10, 11, 13, 18]:  # Skip parameterized tables
        continue
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            text = cell.text
            for old in old_markers:
                if old in text:
                    snippet = text[max(0, text.index(old) - 30):text.index(old) + 50]
                    print(f"  Table {ti}, Row {ri}: ...{snippet}...")
                    break
