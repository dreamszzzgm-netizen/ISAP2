"""Генерация полного DOCX из данных формы «Сведения об ОПО»."""
import asyncio
import sys
sys.path.insert(0, ".")

from src.application.engines.base import DocumentContext
from src.application.engines.data_engine import DataEngine
from src.application.engines.rules_engine import RulesEngine
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# === Данные формы «Сведения об ОПО» ===
FORM_DATA = {
    "f1_1": "Сеть газопотребления города Тюмень",
    "f1_2": "Сеть газопотребления",
    "f1_3": "А34-00000-ГТ",
    "f1_4": "г. Тюмень, ул. Мира, д. 15",
    "danger_class": "3",
    "processes_text": "2.1, 2.2а, 2.6",
    "classification_text": "4.1, 4.2",
    "composition": [
        {"name": "ГРПШ №1", "danger": "0.85", "substance": "", "characteristics": "Шкафный газораспределительный пункт"},
        {"name": "Газопровод Ду150", "danger": "0.72", "substance": "", "characteristics": "Стальной газопровод среднего давления"},
        {"name": "Регулятор РДУК-50", "danger": "0.45", "substance": "", "characteristics": "Регулятор давления газа"},
        {"name": "Природный газ (метан)", "danger": "0.95", "substance": "Метан", "characteristics": "Основное вещество"},
        {"name": "Кран шаровой Ду150", "danger": "0.30", "substance": "", "characteristics": "Запорная арматура"},
    ],
    "f7": "3.27",
    "applicant_type": "legal",
    "f8_1_1": "ООО «Газпром межрегионгаз Тюмень»",
    "f8_1_3": "7736050003",
    "f8_1_5": "1027700196362",
    "f8_1_6": "г. Тюмень, ул. Республики, д. 64",
    "f9_5": "+7 (3452) 53-30-00",
    "f9_6": "info@mg-tyumen.ru",
    "signDolj": "Главный инженер",
    "signPodp": "Иванов И.И.",
    "signDate": "05.07.2026",
    "signMp": "г. Тюмень",
}


def build_context(fd):
    """Маппинг формы OPO → контекст генерации."""
    org = {
        "name": fd.get("f8_1_1", ""),
        "inn": fd.get("f8_1_3", ""),
        "ogrn": fd.get("f8_1_5", ""),
        "address": fd.get("f8_1_6", ""),
        "phone": fd.get("f9_5", ""),
        "email": fd.get("f9_6", ""),
    }
    fac = {
        "name": fd.get("f1_1", ""),
        "facility_type": fd.get("f1_2", ""),
        "hazard_class": fd.get("danger_class", ""),
        "reg_number": fd.get("f1_3", ""),
        "address": fd.get("f1_4", ""),
    }
    equipment = []
    substances = []
    for row in fd.get("composition", []):
        equipment.append({
            "name": row.get("name", ""),
            "equipment_type": row.get("substance", ""),
            "serial_number": "",
            "manufacture_year": None,
        })
        sub = row.get("substance", "")
        if sub:
            substances.append({"name": sub, "quantity_kg": 0, "cas_number": ""})
    return org, fac, equipment, substances


def main():
    fd = FORM_DATA
    org, fac, equipment, substances = build_context(fd)

    data_engine = DataEngine()
    rules_engine = RulesEngine()

    ctx = DocumentContext(
        organization=org, facility=fac, equipment=equipment,
        substances=substances, persons=[], year=2026,
    )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Титульная
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ПЛАН ЛОКАЛИЗАЦИИ И ЛИКВИДАЦИИ АВАРИЙ")
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Объект: " + fac["name"])
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Организация: {org['name']}  |  ИНН: {org['inn']}  |  Класс: {fd['danger_class']}")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.italic = True

    doc.add_paragraph()

    # Сведения об ОПО
    doc.add_heading("Сведения об ОПО (из формы)", level=1)
    t = doc.add_table(rows=8, cols=2)
    t.style = "Table Grid"
    rows_data = [
        ("Полное наименование", fd["f1_1"]),
        ("Тип объекта", fd["f1_2"]),
        ("Регистрационный номер", fd["f1_3"]),
        ("Адрес", fd["f1_4"]),
        ("Класс опасности", fd["danger_class"]),
        ("Процессы", fd["processes_text"]),
        ("Классификация", fd["classification_text"]),
        ("Суммарная опасность (ОВ)", fd["f7"]),
    ]
    for i, (k, v) in enumerate(rows_data):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for cell in t.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"
    for cell in t.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()

    # Все разделы
    sections = [
        ("data", "section_1", "1. Характеристика объекта"),
        ("data", "section_3", "3. Характеристика аварийности"),
        ("data", "section_4", "4. Силы и средства"),
        ("rules", "section_10", "10. Первоочередные действия"),
        ("rules", "section_11", "11. Действия персонала"),
        ("data", "section_8", "8. Управление, связь, оповещение"),
        ("rules", "section_12", "12. Безопасность населения"),
    ]

    for engine_name, sec_id, heading_text in sections:
        engine = data_engine if engine_name == "data" else rules_engine
        result = asyncio.run(engine.generate(sec_id, {"title": heading_text}, ctx))
        doc.add_heading(heading_text, level=1)

        for block in result.blocks:
            if block.__class__.__name__ == "TableBlock":
                table = doc.add_table(rows=len(block.rows) + 1, cols=len(block.headers))
                table.style = "Table Grid"
                for j, header in enumerate(block.headers):
                    cell = table.rows[0].cells[j]
                    cell.text = header
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(9)
                            r.font.name = "Times New Roman"
                for i, row in enumerate(block.rows):
                    for j, cell_text in enumerate(row):
                        cell = table.rows[i + 1].cells[j]
                        cell.text = str(cell_text)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(9)
                                r.font.name = "Times New Roman"
                if block.caption:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(block.caption)
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"
                    r.italic = True
            elif block.__class__.__name__ == "ParagraphBlock":
                p = doc.add_paragraph()
                run = p.add_run(block.text)
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
                if getattr(block, "bold", False):
                    run.bold = True
        doc.add_paragraph()

    # Подпись
    doc.add_heading("Подпись", level=1)
    doc.add_paragraph(f"Должность: {fd['signDolj']}")
    doc.add_paragraph(f"ФИО: {fd['signPodp']}")
    doc.add_paragraph(f"Дата: {fd['signDate']}")
    doc.add_paragraph(f"М.П.: {fd['signMp']}")

    output_path = "test_opo_flow_output.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
