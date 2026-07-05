"""Inspect etalon DOCX formatting."""
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'D:\GPT PMLA\Разработка\isap_skeleton\isap\ПМЛА ООО СПК ААА.docx')

# Table 0 formatting
t = doc.tables[0]
print('=== Table 0: СОГЛАСОВАНО ===')
cell = t.cell(0, 0)
for p in cell.paragraphs:
    for run in p.runs:
        print(f'  Font: {run.font.name}, Size: {run.font.size}, Bold: {run.font.bold}')
        break
    break

tbl = t._tbl
tbl_pr = tbl.find(qn('w:tblPr'))
if tbl_pr is not None:
    tbl_style = tbl_pr.find(qn('w:tblStyle'))
    if tbl_style is not None:
        val = tbl_style.get(qn('w:val'))
        print(f'  Style: {val}')
    tbl_borders = tbl_pr.find(qn('w:tblBorders'))
    print(f'  Has borders: {tbl_borders is not None}')

# Table 9 (scenarios) formatting
t9 = doc.tables[9]
print('\n=== Table 9: Сценарии аварий ===')
cell = t9.cell(0, 0)
for p in cell.paragraphs:
    for run in p.runs:
        print(f'  Header font: {run.font.name}, Size: {run.font.size}, Bold: {run.font.bold}')
        break
    break
cell = t9.cell(1, 0)
for p in cell.paragraphs:
    for run in p.runs:
        print(f'  Data font: {run.font.name}, Size: {run.font.size}, Bold: {run.font.bold}')
        break
    break

# All tables summary
print('\n=== All tables ===')
for i, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.columns)
    first = t.cell(0, 0).text.strip().replace('\n', ' ')[:50]
    print(f'  T{i}: {rows}x{cols} — {first}')

# Page setup
section = doc.sections[0]
print(f'\n=== Page setup ===')
print(f'  Width: {section.page_width.cm:.1f} cm')
print(f'  Height: {section.page_height.cm:.1f} cm')
print(f'  Top margin: {section.top_margin.cm:.1f} cm')
print(f'  Bottom margin: {section.bottom_margin.cm:.1f} cm')
print(f'  Left margin: {section.left_margin.cm:.1f} cm')
print(f'  Right margin: {section.right_margin.cm:.1f} cm')

# Normal style
normal = doc.styles['Normal']
print(f'\n=== Normal style ===')
print(f'  Font: {normal.font.name}')
print(f'  Size: {normal.font.size}')
