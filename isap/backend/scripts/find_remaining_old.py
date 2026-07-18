"""Find all remaining old data locations in the template."""
from docx import Document

doc = Document(r"D:\Project ISAP\isap\isap\files\pmla_v2_template.docx")

markers = ["А34-99999-0099", "Чегем", "г. Чегем"]
found = []

# Check paragraphs
for i, p in enumerate(doc.paragraphs):
    for m in markers:
        if m in p.text:
            idx = p.text.index(m)
            snippet = p.text[max(0, idx-40):idx+60]
            found.append(f"  Para {i:3d}: ...{snippet}...")
            break

# Check tables
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for m in markers:
                if m in cell.text:
                    idx = cell.text.index(m)
                    snippet = cell.text[max(0, idx-30):idx+50]
                    found.append(f"  Table {ti}, Row {ri}, Col {ci}: ...{snippet}...")
                    break

# Check headers
for i, section in enumerate(doc.sections):
    header = section.header
    if header:
        for p in header.paragraphs:
            for m in markers:
                if m in p.text:
                    found.append(f"  Header sec {i}: {p.text[:80]}")
                    break

print(f"Total occurrences: {len(found)}")
for f in found:
    print(f)
