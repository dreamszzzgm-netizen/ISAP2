"""Fix headers: parameterize 'СПК «ААА»' and reg number in all section headers."""
from docx import Document
from lxml import etree
from docx.oxml.ns import qn

TEMPLATE = r"D:\Project ISAP\isap\isap\files\pmla_v2_template.docx"

doc = Document(TEMPLATE)

fixed = 0
for i, section in enumerate(doc.sections):
    header = section.header
    if not header or not header.paragraphs:
        continue
    for para in header.paragraphs:
        original = para.text
        if "СПК" in original or "ААА" in original:
            # Replace the entire header text with parameterized version
            for run in para.runs:
                if "СПК" in run.text or "ААА" in run.text:
                    run.text = run.text.replace("СПК «ААА»", "{{ organization_short_name }}")
                    run.text = run.text.replace("А34-00000-0001", "{{ facility_reg_number }}")
                    fixed += 1
                    print(f"  Section {i} header: fixed run '{run.text[:60]}'")

doc.save(TEMPLATE)
print(f"\nFixed {fixed} header runs")
print(f"Saved: {TEMPLATE}")
